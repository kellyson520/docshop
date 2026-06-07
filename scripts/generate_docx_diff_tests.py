"""
生成 DOCX diff 回归测试样本。

示例：
python scripts/generate_docx_diff_tests.py --source "C:\\path\\source.docx" --count 10 --verify
"""
from __future__ import annotations

import argparse
import json
import random
import sys
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches
from PIL import Image, ImageDraw

ROOT = Path(__file__).resolve().parents[1]
DEFAULT_OUT = ROOT / "artifacts" / "docx_diff_random_tests"

TEST_SENTENCES = [
    "【DIFF-新增文本】作物长势监测、变量施肥路径规划与闭环反馈控制描述。",
    "【DIFF-新增文本】多光谱图像特征融合模块用于识别叶绿素变化与氮素缺失区域。",
    "【DIFF-新增文本】无人机航线依据地块网格、风速与喷幅自动调整。",
    "【DIFF-新增文本】施肥决策模型新增异常值过滤策略。",
]
REPLACE_SUFFIXES = [
    "【DIFF-替换】补充：该指标将同步记录到施肥作业日志。",
    "【DIFF-替换】调整：此处参数改为动态阈值并支持人工校准。",
    "【DIFF-替换】修订：原固定采样策略改为分层抽样策略。",
]


def paragraph_texts(doc):
    return [p for p in doc.paragraphs if p.text and p.text.strip()]


def insert_paragraph_after(paragraph, text):
    new_p = paragraph._parent.add_paragraph(text)
    paragraph._element.addnext(new_p._element)
    return new_p


def delete_paragraph(paragraph):
    parent = paragraph._element.getparent()
    if parent is not None:
        parent.remove(paragraph._element)


def move_paragraph_after(src, dst):
    parent = src._element.getparent()
    if parent is None:
        return
    parent.remove(src._element)
    dst._element.addnext(src._element)


def replace_paragraph_text(paragraph, suffix):
    base = paragraph.text.strip()
    paragraph.clear()
    paragraph.add_run((base[:160] if base else "段落内容") + " " + suffix)


def force_delete_source_like_paragraph(doc, actions: list[dict[str, Any]]) -> bool:
    """删除一个未被脚本新增/替换标记污染的段落，保证 verify 能看到文本删除。"""
    candidates = [
        p for p in paragraph_texts(doc)
        if "【DIFF-" not in p.text and len(p.text.strip()) > 8
    ]
    if not candidates:
        return False
    paragraph = candidates[-1]
    actions.append({"type": "paragraph_delete_forced", "preview": paragraph.text[:80]})
    delete_paragraph(paragraph)
    return True


def add_test_image(path: Path, variant: int):
    img = Image.new("RGB", (680, 260), (245, 250, 255))
    draw = ImageDraw.Draw(img)
    draw.rectangle([8, 8, 672, 252], outline=(31, 95, 163), width=4)
    draw.rectangle([30, 42, 650, 108], fill=(31, 95, 163))
    draw.text((45, 60), f"DOCX DIFF IMAGE TEST V{variant:02d}", fill=(255, 255, 255))
    draw.text((45, 135), "新增/替换图片识别测试：变量施肥路径示意图", fill=(20, 45, 75))
    draw.text((45, 180), f"seed=20260606 variant={variant}", fill=(80, 100, 120))
    img.save(path)


def iter_image_blips(doc):
    """Yield (paragraph, blip, rel) for images visible in document paragraphs."""
    for paragraph in doc.paragraphs:
        try:
            blips = paragraph._element.xpath(".//a:blip")
        except Exception:
            blips = []
        for blip in blips:
            r_id = blip.get(qn("r:embed")) or blip.get(qn("r:link"))
            rel = doc.part.rels.get(r_id) if r_id else None
            if rel is not None and "image" in getattr(rel, "reltype", ""):
                yield paragraph, blip, rel


def remove_first_image_paragraph(doc, actions: list[dict[str, Any]]) -> bool:
    """删除第一张图片所在段落，生成图片删除样本。"""
    for paragraph, _blip, rel in iter_image_blips(doc):
        parent = paragraph._element.getparent()
        if parent is None:
            continue
        parent.remove(paragraph._element)
        actions.append({
            "type": "image_delete",
            "rId": getattr(rel, "rId", ""),
            "target": getattr(rel, "target_ref", ""),
        })
        return True
    return False


def resize_first_image(doc, scale: float, actions: list[dict[str, Any]]) -> bool:
    """调整第一张图片的显示尺寸，保留同一关系和同一图片 hash。"""
    for _paragraph, blip, rel in iter_image_blips(doc):
        node = blip
        changed = False
        old_extent = None
        new_extent = None
        while node is not None:
            tag = getattr(node, "tag", "")
            if tag.endswith("}inline") or tag.endswith("}anchor"):
                for extent in node.xpath(".//wp:extent"):
                    cx = int(extent.get("cx") or 0)
                    cy = int(extent.get("cy") or 0)
                    if cx and cy:
                        old_extent = {"cx": cx, "cy": cy}
                        new_cx = max(1, int(cx * scale))
                        new_cy = max(1, int(cy * scale))
                        extent.set("cx", str(new_cx))
                        extent.set("cy", str(new_cy))
                        new_extent = {"cx": new_cx, "cy": new_cy}
                        changed = True
                for extent in node.xpath(".//a:ext"):
                    cx = int(extent.get("cx") or 0)
                    cy = int(extent.get("cy") or 0)
                    if cx and cy:
                        extent.set("cx", str(max(1, int(cx * scale))))
                        extent.set("cy", str(max(1, int(cy * scale))))
                break
            node = node.getparent()
        if changed:
            actions.append({
                "type": "image_resize",
                "rId": getattr(rel, "rId", ""),
                "scale": scale,
                "old_extent": old_extent,
                "new_extent": new_extent,
            })
            return True
    return False


def replace_first_image_content(doc, replacement_path: Path, actions: list[dict[str, Any]]) -> bool:
    """替换第一张图片内容；外链图片改 target，内嵌图片改 blob。"""
    for _paragraph, _blip, rel in iter_image_blips(doc):
        if getattr(rel, "is_external", False):
            old_target = getattr(rel, "target_ref", "")
            rel._target = f"https://example.invalid/docdist/replaced-{replacement_path.stem}.png"
            actions.append({
                "type": "image_replace",
                "mode": "external_target",
                "rId": getattr(rel, "rId", ""),
                "old_target": old_target,
                "new_target": rel._target,
            })
            return True

        target_part = getattr(rel, "target_part", None)
        if target_part is None:
            continue
        old_blob = getattr(target_part, "blob", b"") or b""
        target_part._blob = replacement_path.read_bytes()
        actions.append({
            "type": "image_replace",
            "mode": "embedded_blob",
            "rId": getattr(rel, "rId", ""),
            "old_size": len(old_blob),
            "new_size": replacement_path.stat().st_size,
            "image": str(replacement_path),
        })
        return True
    return False


def mutate_tables(doc, rng, variant: int, actions: list[dict[str, Any]]):
    if doc.tables:
        table = rng.choice(doc.tables)
        if table.rows and table.columns:
            r = rng.randrange(len(table.rows))
            c = rng.randrange(len(table.columns))
            old = table.cell(r, c).text
            table.cell(r, c).text = (old[:80] + "\n" if old else "") + f"【DIFF-表格替换】V{variant:02d}-R{r+1}C{c+1}"
            actions.append({"type": "table_cell_replace", "row": r + 1, "col": c + 1})
        if len(table.rows) >= 2:
            new_row = table.add_row()
            for i, cell in enumerate(new_row.cells):
                cell.text = f"【DIFF-表格新增行】V{variant:02d}-{i+1}"
            actions.append({"type": "table_row_insert"})
        if len(table.rows) >= 3 and variant % 2 == 0:
            tbl = table._tbl
            moving = list(tbl.tr_lst)[1]
            tbl.remove(moving)
            tbl.append(moving)
            actions.append({"type": "table_row_reorder", "from": 2, "to": len(table.rows)})

    new_table = doc.add_table(rows=4, cols=4)
    try:
        new_table.style = "Table Grid"
    except Exception:
        pass
    for r, row in enumerate(new_table.rows):
        for c, cell in enumerate(row.cells):
            cell.text = f"V{variant:02d}-T新增-{r+1}-{c+1}"
    actions.append({"type": "table_add", "shape": "4x4"})


def make_variant(source: Path, out_dir: Path, img_dir: Path, variant: int) -> dict[str, Any]:
    rng = random.Random(20260606 + variant)
    doc = Document(str(source))
    actions: list[dict[str, Any]] = []

    paras = paragraph_texts(doc)
    if paras:
        for k in range(2):
            anchor = rng.choice(paras)
            text = f"{TEST_SENTENCES[(variant + k) % len(TEST_SENTENCES)]} [V{variant:02d}-INS{k+1}]"
            insert_paragraph_after(anchor, text)
            actions.append({"type": "paragraph_insert", "text": text})

        paras = paragraph_texts(doc)
        for k in range(min(2, len(paras))):
            p = rng.choice(paras)
            replace_paragraph_text(p, f"{REPLACE_SUFFIXES[(variant + k) % len(REPLACE_SUFFIXES)]} [V{variant:02d}-REP{k+1}]")
            actions.append({"type": "paragraph_replace"})

        paras = [p for p in paragraph_texts(doc) if len(p.text.strip()) > 12]
        for _ in range(min(1 + (variant % 2), len(paras))):
            p = rng.choice(paras)
            actions.append({"type": "paragraph_delete", "preview": p.text[:80]})
            delete_paragraph(p)
            paras = [x for x in paragraph_texts(doc) if len(x.text.strip()) > 12]

        paras = paragraph_texts(doc)
        if len(paras) >= 4:
            src = rng.choice(paras[1:-1])
            dst = rng.choice([p for p in paras if p is not src])
            move_paragraph_after(src, dst)
            actions.append({"type": "paragraph_reorder", "preview": src.text[:80]})

        force_delete_source_like_paragraph(doc, actions)

    mutate_tables(doc, rng, variant, actions)

    img_path = img_dir / f"variant_{variant:02d}_diff_image.png"
    add_test_image(img_path, variant)

    replacement_img_path = img_dir / f"variant_{variant:02d}_replace_image.png"
    add_test_image(replacement_img_path, 100 + variant)
    if variant % 3 == 1:
        remove_first_image_paragraph(doc, actions)
    elif variant % 3 == 2:
        replace_first_image_content(doc, replacement_img_path, actions)
    else:
        resize_first_image(doc, 1.35, actions)

    doc.add_paragraph(f"【DIFF-图片说明】下方为 V{variant:02d} 新增测试图片，用于检测图片新增显示。")
    doc.add_picture(str(img_path), width=Inches(4.6))
    actions.append({"type": "image_insert", "image": str(img_path)})

    # 删除源图片后，python-docx 可能复用 media/image1.png，diff 会将首张新增图
    # 判定为替换。删除型样本额外插入一张图，确保每份样本都有明确 image_added。
    if variant % 3 == 1:
        extra_img_path = img_dir / f"variant_{variant:02d}_extra_added_image.png"
        add_test_image(extra_img_path, 200 + variant)
        doc.add_paragraph(f"【DIFF-图片说明】V{variant:02d} 额外新增图片，用于区分删除与新增。")
        doc.add_picture(str(extra_img_path), width=Inches(3.8))
        actions.append({"type": "image_insert_extra", "image": str(extra_img_path)})

    if variant in {4, 8}:
        doc.add_page_break()
        doc.add_paragraph(f"【DIFF-新增分页】V{variant:02d} 独立新增测试页。")
        actions.append({"type": "page_break_insert"})

    output = out_dir / f"diff_test_variant_{variant:02d}.docx"
    doc.save(str(output))
    return {"variant": variant, "output": str(output), "output_size": output.stat().st_size, "operations": actions}


def verify(source: Path, out_dir: Path) -> dict[str, Any]:
    sys.path.insert(0, str(ROOT / "backend"))
    from app.diff_engine.docx_diff import DocxDiffEngine

    engine = DocxDiffEngine()
    results = []
    for path in sorted(out_dir.glob("diff_test_variant_*.docx")):
        diff = engine.compare(str(source), str(path))
        stats = diff.get("stats") or {}
        paragraph_changes = diff.get("paragraphs") or diff.get("text") or []
        char_segments = [
            seg
            for paragraph in paragraph_changes
            for seg in paragraph.get("char_diffs", [])
            if isinstance(seg, dict)
        ]
        checks = {
            "text_added": (
                stats.get("paragraphs_added", 0) > 0
                or any(seg.get("type") == "insert" and seg.get("text") for seg in char_segments)
            ),
            "text_deleted": (
                stats.get("paragraphs_deleted", 0) > 0
                or any(seg.get("type") == "delete" and seg.get("text") for seg in char_segments)
            ),
            "text_modified": stats.get("paragraphs_modified", 0) > 0,
            "table_changed": stats.get("tables_changed", 0) > 0,
            "image_added": stats.get("images_added", 0) > 0,
        }
        ok = all(checks.values())
        results.append({
            "file": str(path),
            "ok": ok,
            "checks": checks,
            "summary": diff.get("summary"),
            "stats": stats,
            "images": diff.get("images"),
        })
    passed = sum(1 for item in results if item["ok"])
    report = {
        "source": str(source),
        "tested": len(results),
        "passed": passed,
        "failed": len(results) - passed,
        "results": results,
    }
    (out_dir / "diff_engine_verification.json").write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    return report


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--source", type=Path, default=None)
    parser.add_argument("--count", type=int, default=10)
    parser.add_argument("--out", type=Path, default=DEFAULT_OUT)
    parser.add_argument("--verify", action="store_true")
    args = parser.parse_args()

    source = args.source or sorted(ROOT.glob("*.docx"), key=lambda p: p.stat().st_mtime, reverse=True)[0]
    out_dir = args.out
    img_dir = out_dir / "generated_images"
    out_dir.mkdir(parents=True, exist_ok=True)
    img_dir.mkdir(parents=True, exist_ok=True)

    manifest = {
        "generated_at": datetime.now().isoformat(timespec="seconds"),
        "source": str(source),
        "output_dir": str(out_dir),
        "variants": [make_variant(source, out_dir, img_dir, i) for i in range(1, args.count + 1)],
    }
    (out_dir / "manifest.json").write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")

    result = {"output_dir": str(out_dir), "manifest": str(out_dir / "manifest.json"), "count": args.count}
    if args.verify:
        result["verification"] = verify(source, out_dir)
    print(json.dumps(result, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
