"""test_conversion.py — 验证 DOCX/PDF → images HTML 转换流水线。"""

import os
import re
import zipfile
from types import SimpleNamespace
import pytest
from docx import Document
from PIL import Image

from app.services.conversion_service import convert_to_images_html, _convert_docx_to_html

@pytest.mark.slow
def test_convert_docx_to_images_html(tmp_doc_root, sample_docx):
    """DOCX → PDF（MS Word COM）→ JPEG 页面 → HTML。"""
    import time
    fid = "conv-docx-0000-0000-0000-000000000001"

    t0 = time.time()
    html = convert_to_images_html(fid, sample_docx, "docx")
    elapsed = time.time() - t0

    assert html is not None
    assert "<!DOCTYPE html>" in html
    assert "data:image/jpeg;base64," in html
    pages = len(re.findall(r'data:image/jpeg;base64,', html))
    assert pages >= 1
    print(f"\n  DOCX {pages} pages, {len(html) // 1024}KB, {elapsed:.1f}s")

def test_convert_pdf_to_images_html(tmp_doc_root, sample_pdf):
    """PDF → JPEG → HTML（跳过 Word→PDF）。"""
    import time
    fid = "conv-pdf-0000-0000-0000-000000000002"

    t0 = time.time()
    html = convert_to_images_html(fid, sample_pdf, "pdf")
    elapsed = time.time() - t0

    assert html is not None
    assert "data:image/jpeg;base64," in html
    pages = len(re.findall(r'data:image/jpeg;base64,', html))
    assert pages == 1

    # 缓存命中应该 < 2s
    t0 = time.time()
    html2 = convert_to_images_html(fid, sample_pdf, "pdf")
    cached = time.time() - t0
    assert html2 is not None
    assert cached < 2.0
    print(f"\n  PDF 1 page, cache hit {cached:.2f}s")

def test_convert_unsupported_type_returns_none(tmp_doc_root, sample_pdf):
    fid = "conv-bad-0000-0000-0000-000000000003"
    html = convert_to_images_html(fid, sample_pdf, "txt")
    assert html is None


def _png_bytes(path):
    image = Image.new("RGB", (8, 8), color=(255, 0, 0))
    image.save(path, format="PNG")
    return path.read_bytes()


def _make_docx_with_linked_image(tmp_path):
    """Create a minimal DOCX whose picture blip uses r:link instead of r:embed."""
    png_path = tmp_path / "source.png"
    _png_bytes(png_path)

    docx_path = tmp_path / "linked-image.docx"
    doc = Document()
    doc.add_paragraph("before image")
    doc.add_picture(str(png_path))
    doc.save(docx_path)

    with zipfile.ZipFile(docx_path, "r") as zin:
        files = {name: zin.read(name) for name in zin.namelist()}

    rels_name = "word/_rels/document.xml.rels"
    document_name = "word/document.xml"
    rels = files[rels_name].decode("utf-8")
    document = files[document_name].decode("utf-8")

    match = re.search(r'Id="(rId\d+)"[^>]+Target="media/image1\.png"', rels)
    assert match, rels
    rid = match.group(1)
    rels = re.sub(
        rf'(<Relationship Id="{rid}"[^>]+Target=")media/image1\.png("[^>]*/>)',
        rf'\1http://example.test/linked.png\2',
        rels,
    )
    rels = re.sub(
        rf'(<Relationship Id="{rid}"[^>]*)/>',
        rf'\1 TargetMode="External"/>',
        rels,
        count=1,
    )
    document = document.replace(f'r:embed="{rid}"', f'r:link="{rid}"')

    with zipfile.ZipFile(docx_path, "w", zipfile.ZIP_DEFLATED) as zout:
        for name, content in files.items():
            if name == rels_name:
                content = rels.encode("utf-8")
            elif name == document_name:
                content = document.encode("utf-8")
            elif name == "word/media/image1.png":
                continue
            zout.writestr(name, content)

    return docx_path, png_path.read_bytes()


def test_docx_html_embeds_linked_images(monkeypatch, tmp_path):
    docx_path, image_bytes = _make_docx_with_linked_image(tmp_path)

    def fake_get(url, timeout=0, headers=None):
        assert url == "http://example.test/linked.png"
        return SimpleNamespace(
            status_code=200,
            headers={"content-type": "image/png"},
            content=image_bytes,
            raise_for_status=lambda: None,
        )

    monkeypatch.setattr("requests.get", fake_get)

    html = _convert_docx_to_html(str(docx_path))

    assert "data:image/png;base64," in html
    assert "http://example.test/linked.png" not in html


def test_docx_html_preserves_table_position(tmp_path):
    docx_path = tmp_path / "table-order.docx"
    doc = Document()
    doc.add_paragraph("Paragraph before table")
    table = doc.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "Cell in the middle"
    doc.add_paragraph("Paragraph after table")
    doc.save(docx_path)

    html = _convert_docx_to_html(str(docx_path))

    before_idx = html.index("Paragraph before table")
    table_idx = html.index("<table")
    cell_idx = html.index("Cell in the middle")
    after_idx = html.index("Paragraph after table")
    assert before_idx < table_idx < cell_idx < after_idx
