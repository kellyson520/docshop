"""
文件格式转换服务

支持:
- DOCX/DOC → PDF（优先使用 docx2pdf / MS Word COM 原生转换）
- DOCX/DOC → HTML（增强渲染：图片、列表、页眉页脚）
- XLSX/XLS → HTML 表格 / PDF
- 30 天缓存：已转换的 PDF 缓存到磁盘，按源文件哈希索引

设计原则：
1. Windows: 使用 docx2pdf（底层调用 MS Word COM），输出像素级完美
2. 非 Windows 或无 Word: 回退到 LibreOffice headless
3. 最后回退: python-docx 生成增强 HTML
4. 所有 PDF 转换结果按源文件 SHA-256 缓存 30 天
"""

import hashlib
import json
import os
import posixpath
import shutil
import subprocess
import tempfile
import time
import zipfile
from html import escape as html_escape
from urllib.parse import quote_plus
import multiprocessing
import queue
from pathlib import Path
from typing import Optional, Tuple, List
import xml.etree.ElementTree as ET

from app.config import settings
from app.utils.logger import get_logger

logger = get_logger("services.conversion")

_REL_NS = "http://schemas.openxmlformats.org/package/2006/relationships"
_CONTENT_TYPES_NS = "http://schemas.openxmlformats.org/package/2006/content-types"
_EXTERNAL_IMAGE_PLACEHOLDER_NAME = "word/media/external_image_placeholder.png"
_EXTERNAL_IMAGE_PLACEHOLDER_BYTES = (
    b"\x89PNG\r\n\x1a\n\x00\x00\x00\rIHDR\x00\x00\x00\x01"
    b"\x00\x00\x00\x01\x08\x06\x00\x00\x00\x1f\x15\xc4\x89"
    b"\x00\x00\x00\nIDATx\x9cc`\x00\x00\x00\x02\x00\x01"
    b"\xe2!\xbc3\x00\x00\x00\x00IEND\xaeB`\x82"
)

# ── 缓存配置 ────────────────────────────────────────────────────

CACHE_DIR = os.path.join(settings.TEMP_DIR, "conversions")
CACHE_TTL_SECONDS = 30 * 24 * 3600  # 30 天


def _path_is_within_root(path: str, root: str) -> bool:
    try:
        resolved_root = Path(root).resolve()
        resolved_path = Path(path).resolve()
        return resolved_path == resolved_root or resolved_root in resolved_path.parents
    except (OSError, RuntimeError, ValueError):
        return False


def _ensure_temp_dir() -> str:
    os.makedirs(settings.TEMP_DIR, exist_ok=True)
    return settings.TEMP_DIR


def _ensure_cache_dir() -> str:
    """确保缓存目录存在。"""
    os.makedirs(CACHE_DIR, exist_ok=True)
    return CACHE_DIR


def _source_hash(file_path: str) -> str:
    """计算源文件的 SHA-256 哈希（用于缓存键）。"""
    h = hashlib.sha256()
    with open(file_path, "rb") as f:
        while chunk := f.read(65536):
            h.update(chunk)
    return h.hexdigest()


def _cache_path(source_hash: str, suffix: str) -> str:
    """返回缓存文件路径：cache_dir / {前2位} / {hash}.{suffix}"""
    prefix = source_hash[:2]
    d = os.path.join(CACHE_DIR, prefix)
    os.makedirs(d, exist_ok=True)
    return os.path.join(d, f"{source_hash}{suffix}")


def _cache_meta_path(cache_file: str) -> str:
    """返回缓存元数据文件路径。"""
    return cache_file + ".meta.json"


def _read_cache(source_hash: str, suffix: str) -> Optional[str]:
    """
    读取缓存。如果缓存存在且未过期，返回缓存文件路径；否则返回 None。
    过期缓存会被自动清理。
    """
    cache_file = _cache_path(source_hash, suffix)
    meta_file = _cache_meta_path(cache_file)

    if not os.path.exists(cache_file) or not os.path.exists(meta_file):
        return None

    try:
        with open(meta_file, "r", encoding="utf-8") as f:
            meta = json.load(f)
    except (json.JSONDecodeError, OSError):
        # 元数据损坏 → 清理
        _remove_cache(cache_file)
        return None

    created_at = meta.get("created_at", 0)
    if time.time() - created_at > CACHE_TTL_SECONDS:
        _remove_cache(cache_file)
        return None

    if os.path.getsize(cache_file) == 0:
        _remove_cache(cache_file)
        return None

    logger.info(f"Cache hit: {cache_file} (age: {(time.time() - created_at) / 3600:.1f}h)")
    return cache_file


def _write_cache(source_hash: str, source_path: str, suffix: str, source_file_type: str) -> str:
    """将源文件复制到缓存（用于 PDF 缓存的是转换结果，这里直接复制原文件作缓存）。"""
    cache_file = _cache_path(source_hash, suffix)
    shutil.copy2(source_path, cache_file)

    meta = {
        "created_at": time.time(),
        "source_hash": source_hash,
        "source_file_type": source_file_type,
        "original_size": os.path.getsize(source_path),
    }
    meta_file = _cache_meta_path(cache_file)
    with open(meta_file, "w", encoding="utf-8") as f:
        json.dump(meta, f)

    logger.info(f"Cache write: {cache_file}")
    return cache_file


def _remove_cache(cache_file: str) -> None:
    """删除缓存文件及其元数据。"""
    for f in (cache_file, _cache_meta_path(cache_file)):
        try:
            if os.path.exists(f):
                os.unlink(f)
        except OSError:
            pass


# ── 转换引擎检测 ────────────────────────────────────────────────

def _has_docx2pdf() -> bool:
    """检测 MS Word COM 是否可用（Windows 上通过 win32com 调用 Word）。"""
    try:
        import win32com.client  # noqa: F401
        return True
    except ImportError:
        return False


def _find_libreoffice() -> Optional[str]:
    """查找 LibreOffice 可执行文件路径。"""
    candidates = [
        "libreoffice", "soffice",
        "/usr/bin/libreoffice", "/usr/bin/soffice",
        "/usr/local/bin/libreoffice", "/usr/local/bin/soffice",
        r"C:\Program Files\LibreOffice\program\soffice.exe",
        r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
    ]
    for candidate in candidates:
        try:
            result = subprocess.run(
                [candidate, "--version"], capture_output=True, timeout=5
            )
            if result.returncode == 0:
                return candidate
        except (FileNotFoundError, subprocess.TimeoutExpired, OSError):
            continue
    return None


# 懒加载检测（只检测一次）
_engine = None  # 'docx2pdf' | 'libreoffice' | 'fallback'
_engine_lock = __import__("threading").Lock()


def _detect_engine() -> str:
    """检测最佳转换引擎，只执行一次。"""
    global _engine
    if _engine is not None:
        return _engine

    with _engine_lock:
        if _engine is not None:
            return _engine

        if _has_docx2pdf():
            _engine = "docx2pdf"
            logger.info("转换引擎: docx2pdf (MS Word COM)")
            return _engine

        lo = _find_libreoffice()
        if lo:
            _engine = "libreoffice"
            logger.info(f"转换引擎: LibreOffice ({lo})")
            return _engine

        _engine = "fallback"
        logger.warning("无原生转换引擎，使用 python-docx HTML 渲染")
        return _engine


def _convert_via_libreoffice(input_path: str, output_dir: str, fmt: str = "pdf") -> Optional[str]:
    """使用 LibreOffice headless 转换。"""
    lo = _find_libreoffice()
    if not lo:
        return None
    try:
        result = subprocess.run(
            [lo, "--headless", "--convert-to", fmt, "--outdir", output_dir, input_path],
            capture_output=True, timeout=120, check=False,
        )
        if result.returncode != 0:
            logger.error(f"LibreOffice 失败: {result.stderr.decode(errors='replace')}")
            return None
        base = os.path.splitext(os.path.basename(input_path))[0]
        out = os.path.join(output_dir, f"{base}.{fmt}")
        if os.path.exists(out) and os.path.getsize(out) > 0:
            return out
        return None
    except Exception as e:
        logger.error(f"LibreOffice 异常: {e}")
        return None


# ── DOCX → HTML 增强渲染 ────────────────────────────────────────

def _convert_docx_to_html(input_path: str) -> str:
    """
    将 DOCX 转为 HTML，尽量保留段落、表格、图片和分页结构。
    """
    from docx import Document as DocxDoc
    from docx.table import Table
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.text.paragraph import Paragraph
    from lxml import etree
    import base64
    import html as html_lib

    doc = DocxDoc(input_path)
    MNS = "http://schemas.openxmlformats.org/officeDocument/2006/math"

    # OMML -> MathML XSLT
    xsl_path = Path(__file__).parent.parent / "diff_engine" / "omml2mml.xsl"
    xslt_transform = None
    if xsl_path.exists():
        try:
            xslt_doc = etree.parse(str(xsl_path))
            xslt_transform = etree.XSLT(xslt_doc)
        except Exception:
            pass

    # Extract images
    image_map = {}
    external_image_map = {}
    try:
        for rel in doc.part.rels.values():
            if "image" in rel.reltype:
                try:
                    img_bytes = rel.target_part.blob
                    ext = os.path.splitext(rel.target_part.partname)[-1].lower()
                    mime_map = {".png": "image/png", ".jpg": "image/jpeg",
                                ".jpeg": "image/jpeg", ".gif": "image/gif",
                                ".bmp": "image/bmp", ".svg": "image/svg+xml"}
                    mime = mime_map.get(ext, "image/png")
                    b64 = base64.b64encode(img_bytes).decode()
                    image_map[rel.rId] = f"data:{mime};base64,{b64}"
                except Exception:
                    target_ref = getattr(rel, "target_ref", "")
                    if target_ref:
                        external_image_map[rel.rId] = target_ref
    except Exception:
        pass

    external_image_data_cache = {}
    external_image_failed_hosts = set()

    def _external_image_placeholder(url: str, reason: str = "") -> str:
        safe_url = html_lib.escape(url or "external image", quote=False)
        safe_reason = html_lib.escape(reason or "无法加载外链图片", quote=False)
        svg = (
            '<svg xmlns="http://www.w3.org/2000/svg" width="640" height="160" '
            'viewBox="0 0 640 160">'
            '<rect width="640" height="160" fill="#f8fafc" stroke="#cbd5e1"/>'
            '<text x="24" y="68" font-size="18" fill="#64748b" '
            'font-family="Microsoft YaHei, Arial">外链图片暂不可用</text>'
            f'<text x="24" y="100" font-size="13" fill="#94a3b8" '
            f'font-family="Arial">{safe_reason}</text>'
            f'<text x="24" y="126" font-size="12" fill="#94a3b8" '
            f'font-family="Arial">{safe_url[:90]}</text>'
            '</svg>'
        )
        return "data:image/svg+xml;base64," + base64.b64encode(svg.encode("utf-8")).decode()

    def _external_image_to_data_uri(url: str) -> str:
        """Fetch a linked DOCX image and embed it; fall back to an inline placeholder."""
        from urllib.parse import urlparse

        if not url:
            return ""
        if url.startswith("data:"):
            return url
        if not url.lower().startswith(("http://", "https://")):
            return url
        if url in external_image_data_cache:
            return external_image_data_cache[url]

        host = urlparse(url).netloc
        if host in external_image_failed_hosts:
            data_uri = url
            external_image_data_cache[url] = data_uri
            return data_uri

        try:
            import requests

            response = requests.get(
                url,
                timeout=(1.0, 3.0),
                headers={"User-Agent": "DocShopPreview/1.0"},
            )
            response.raise_for_status()
            content = response.content or b""
            if len(content) > 10 * 1024 * 1024:
                raise ValueError("image exceeds 10MB")
            mime = response.headers.get("content-type", "").split(";", 1)[0].strip().lower()
            if not mime.startswith("image/"):
                ext = os.path.splitext(urlparse(url).path)[-1].lower()
                mime_map = {
                    ".png": "image/png",
                    ".jpg": "image/jpeg",
                    ".jpeg": "image/jpeg",
                    ".gif": "image/gif",
                    ".bmp": "image/bmp",
                    ".svg": "image/svg+xml",
                    ".webp": "image/webp",
                }
                mime = mime_map.get(ext, "image/png")
            data_uri = f"data:{mime};base64,{base64.b64encode(content).decode()}"
        except Exception as exc:
            if host:
                external_image_failed_hosts.add(host)
            data_uri = url

        external_image_data_cache[url] = data_uri
        return data_uri

    def _image_src_from_blip(blip) -> str:
        """Return data URI or external URL for a DrawingML image blip."""
        if blip is None:
            return ""
        r_id = blip.get(qn("r:embed")) or blip.get(qn("r:link"))
        if not r_id:
            return ""
        if r_id in image_map:
            return image_map[r_id]
        return _external_image_to_data_uri(external_image_map.get(r_id, ""))

    def _render_run_images(run_elem) -> str:
        """Render all images contained in a w:r element."""
        img_parts = []
        for dw in run_elem.findall(".//" + qn("w:drawing")):
            blip = dw.find(".//" + qn("a:blip"))
            src = _image_src_from_blip(blip)
            if src:
                safe_src = html_lib.escape(src, quote=True)
                img_parts.append(f'<img src="{safe_src}" alt="图片" style="max-width:100%;height:auto">')
        return "".join(img_parts)

    # Headers / footers per section
    def _section_html(section, tag: str) -> str:
        try:
            hdr = getattr(section, tag, None)
            if hdr is None or not hdr.paragraphs:
                return ""
            parts = []
            for p in hdr.paragraphs:
                if p.text.strip():
                    parts.append(f"<p style='margin:0;font-size:9pt;color:#666'>{p.text.strip()}</p>")
            return "".join(parts)
        except Exception:
            return ""

    header_html = ""
    footer_html = ""
    try:
        for section in doc.sections:
            h = _section_html(section, "header")
            f = _section_html(section, "footer")
            if h:
                header_html += h
            if f:
                footer_html += f
    except Exception:
        pass

    # Helper: check if paragraph has a page break
    def _has_page_break(para) -> bool:
        for r in para._element.findall(qn("w:r")):
            for br in r.findall(qn("w:br")):
                if br.get(qn("w:type")) == "page":
                    return True
        # Also check in pPr directly
        pPr = para._element.find(qn("w:pPr"))
        if pPr is not None:
            for br in pPr.iter(qn("w:br")):
                if br.get(qn("w:type")) == "page":
                    return True
        return False

    # ---- CSS for immersive multi-page appearance ----
    css = (
        "html,body{margin:0;padding:0;background:#fff;font-family:\"Times New Roman\",\"SimSun\","
        "\"Microsoft YaHei\",serif;font-size:12pt;color:#111;line-height:1.6;}"
        ".doc-page{margin:0 0 12px 0;padding:0;background:transparent;overflow:visible;}"
        ".doc-page:last-child{margin-bottom:0;}"
        ".page-num{text-align:center;color:#999;font-size:10pt;padding:8px 0 12px;"
        "border-top:1px solid #ddd;margin-top:12pt;}"
        ".docx-header{border-bottom:1px solid #ddd;padding-bottom:8pt;margin-bottom:12pt;text-align:center;}"
        ".docx-footer{border-top:1px solid #ddd;padding-top:8pt;margin-top:12pt;text-align:center;}"
        "h1{font-size:22pt;margin:16pt 0 8pt;font-weight:bold;border-bottom:1px solid #000;padding-bottom:4pt;}"
        "h2{font-size:16pt;margin:14pt 0 6pt;font-weight:bold;}"
        "h3{font-size:14pt;margin:12pt 0 4pt;font-weight:bold;}"
        "p{margin:0 0 6pt 0;}"
        "img{display:block;max-width:100%;height:auto;margin:4pt auto;}"
        "table{border-collapse:collapse;width:100%;margin:8pt 0;}"
        "td,th{border:1px solid #333;padding:3pt 6pt;vertical-align:top;font-size:11pt;}"
        "ul,ol{margin:0 0 6pt 0;padding-left:24pt;}"
        "li{margin:2pt 0;}"
        ".math-block{display:block;text-align:center;margin:12pt 0;padding:4pt 0;}"
        ".math-block mjx-container,.math-block MathJax{display:inline-block!important;}"
        ".math-inline{display:inline;vertical-align:middle;margin:0 2pt;}"
        ".math-inline mjx-container,.math-inline MathJax{display:inline-block!important;vertical-align:middle;}"
        "@media print{body{background:#fff;}.doc-page{page-break-after:always;}}"
        "@media(max-width:800px){.doc-page{margin-bottom:12px;}}"
    )

    parts = [
        '<!DOCTYPE html><html lang="zh-CN"><head><meta charset="utf-8">',
        '<meta name="viewport" content="width=device-width, initial-scale=1.0">',
        '<title>文档预览</title>',
        '<script>MathJax={tex:{inlineMath:[["$","$"]]}};</script>',
        '<script src="https://cdn.jsdelivr.net/npm/mathjax@3/es5/tex-mml-chtml.js"></script>',
        f'<style>{css}</style></head><body>',
    ]

    # ---- Collect all paragraphs and tables into a flat list of "chunks" ----
    class Chunk:
        def __init__(self, kind, data):
            self.kind = kind  # "para" or "table"
            self.data = data

    chunks = []
    for child in doc.element.body.iterchildren():
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if tag == "p":
            chunks.append(Chunk("para", Paragraph(child, doc)))
        elif tag == "tbl":
            chunks.append(Chunk("table", Table(child, doc)))

    # ---- Render each chunk, splitting into pages at page breaks ----
    current_page_parts = []
    all_pages = [current_page_parts]
    page_count = 1

    for chunk in chunks:
        if chunk.kind == "para":
            para = chunk.data
            # Check for page break before rendering this paragraph
            if _has_page_break(para) and current_page_parts:
                # Start a new page
                current_page_parts = []
                all_pages.append(current_page_parts)
                page_count += 1

            pf = para.paragraph_format

            # Skip empty paragraphs
            has_math = (
                para._element.findall(f"{{{MNS}}}oMath") +
                para._element.findall(f"{{{MNS}}}oMathPara")
            )
            has_images = bool(
                para._element.findall(".//" + qn("wp:inline")) +
                para._element.findall(".//" + qn("wp:anchor"))
            )
            if not para.text.strip() and not para.runs and not has_math and not has_images:
                current_page_parts.append('<p><br></p>')
                continue

            css = []
            if pf.alignment == WD_ALIGN_PARAGRAPH.CENTER: css.append("text-align:center")
            elif pf.alignment == WD_ALIGN_PARAGRAPH.RIGHT: css.append("text-align:right")
            elif pf.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY: css.append("text-align:justify")
            if pf.first_line_indent: css.append(f"text-indent:{pf.first_line_indent.pt}pt")
            if pf.space_before: css.append(f"margin-top:{pf.space_before.pt}pt")
            if pf.space_after: css.append(f"margin-bottom:{pf.space_after.pt}pt")
            line_height = _format_docx_line_height(pf.line_spacing)
            if line_height: css.append(f"line-height:{line_height}")
            style = ";".join(css) if css else ""

            # List detection
            numPr = para._element.find(qn("w:pPr"))
            is_list_item = False
            list_tag = ""
            if numPr is not None:
                numPr_elem = numPr.find(qn("w:numPr"))
                if numPr_elem is not None:
                    is_list_item = True
                    list_tag = "ul"

            is_heading = para.style.name.startswith("Heading")
            hn = para.style.name.replace("Heading ", "") if is_heading else ""

            # Image-only paragraph
            if not para.text.strip() and not para.runs:
                images = para._element.findall(".//" + qn("wp:inline")) + para._element.findall(".//" + qn("wp:anchor"))
                for img_elem in images:
                    blip = img_elem.find(".//" + qn("a:blip"))
                    src = _image_src_from_blip(blip)
                    if src:
                        safe_src = html_lib.escape(src, quote=True)
                        current_page_parts.append(f'<p><img src="{safe_src}" alt="图片" style="max-width:100%;height:auto"></p>')
                if not images:
                    current_page_parts.append('<p><br></p>')
                continue

            if is_list_item and not is_heading:
                current_page_parts.append(f'<{list_tag}>')

            default_style = style if style else "margin:3pt 0;line-height:1.5"
            if is_heading:
                current_page_parts.append(f'<h{hn} style="{style or "margin:12pt 0 6pt"}">')
            elif is_list_item:
                style_li = style if style else "margin:2pt 0"
                current_page_parts.append(f'<li style="{style_li}">')
            else:
                current_page_parts.append(f'<p style="{default_style}">')

            # Render child elements in document order
            for child in para._element:
                tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag

                if tag == "r":
                    run_elem = child
                    current_page_parts.append(_render_run_images(run_elem))

                    t_elements = run_elem.findall(qn("w:t"))
                    run_text = "".join(t.text or "" for t in t_elements)
                    if not run_text:
                        if run_elem.find(qn("w:br")) is not None:
                            current_page_parts.append("<br>")
                        continue

                    run_text = run_text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

                    rPr = run_elem.find(qn("w:rPr"))
                    is_bold = False
                    is_italic = False
                    is_underline = False
                    font_size_pt = None
                    font_name = None
                    font_color = None
                    if rPr is not None:
                        b = rPr.find(qn("w:b"))
                        if b is not None:
                            v = b.get(qn("w:val"), "true")
                            is_bold = v not in ("false", "0", "off", "none")
                        i = rPr.find(qn("w:i"))
                        if i is not None:
                            v = i.get(qn("w:val"), "true")
                            is_italic = v not in ("false", "0", "off", "none")
                        u = rPr.find(qn("w:u"))
                        if u is not None:
                            v = u.get(qn("w:val"), "single")
                            is_underline = v not in ("false", "0", "off", "none", "nil")
                        sz = rPr.find(qn("w:sz"))
                        if sz is not None:
                            try:
                                font_size_pt = float(sz.get(qn("w:val"), "0")) / 2
                            except (ValueError, TypeError):
                                pass
                        rFonts = rPr.find(qn("w:rFonts"))
                        if rFonts is not None:
                            font_name = rFonts.get(qn("w:ascii")) or rFonts.get(qn("w:eastAsia")) or rFonts.get(qn("w:hAnsi"))
                        color = rPr.find(qn("w:color"))
                        if color is not None:
                            font_color = color.get(qn("w:val"))

                    tags_list, stys = [], []
                    if is_bold: tags_list.append("b")
                    if is_italic: tags_list.append("i")
                    if is_underline: tags_list.append("u")
                    if font_size_pt: stys.append(f"font-size:{font_size_pt}pt")
                    if font_name: stys.append(f"font-family:'{font_name}'")
                    if font_color: stys.append(f"color:#{font_color}")
                    if stys: tags_list.append(f'span style="{";".join(stys)}"')
                    for t in tags_list: current_page_parts.append(f"<{t}>")
                    current_page_parts.append(run_text)
                    for t in reversed(tags_list): current_page_parts.append(f"</{t.split()[0]}>")

                elif tag == "oMath":
                    current_page_parts.append('<span class="math-inline">')
                    try:
                        if xslt_transform is not None:
                            mathml = str(xslt_transform(child))
                            if mathml.strip():
                                mathml = mathml.replace("<math", '<math display="inline"', 1)
                                current_page_parts.append(mathml)
                            else:
                                current_page_parts.append('<span style="color:#999">[公式]</span>')
                        else:
                            current_page_parts.append('<span style="color:#999">[公式]</span>')
                    except Exception:
                        current_page_parts.append('<span style="color:#999">[公式]</span>')
                    current_page_parts.append('</span>')

                elif tag == "oMathPara":
                    current_page_parts.append('<div class="math-block">')
                    try:
                        if xslt_transform is not None:
                            mathml = str(xslt_transform(child))
                            if mathml.strip():
                                mathml = mathml.replace("<math", '<math display="block"', 1)
                                current_page_parts.append(mathml)
                            else:
                                current_page_parts.append('<span style="color:#999">[公式]</span>')
                        else:
                            current_page_parts.append('<span style="color:#999">[公式]</span>')
                    except Exception:
                        current_page_parts.append('<span style="color:#999">[公式]</span>')
                    current_page_parts.append('</div>')

            # Close paragraph container
            if is_heading:
                current_page_parts.append(f'</h{hn}>')
            elif is_list_item:
                current_page_parts.append('</li>')
            else:
                current_page_parts.append('</p>')

            if is_list_item and not is_heading:
                current_page_parts.append(f'</{list_tag}>')

        elif chunk.kind == "table":
            table = chunk.data
            current_page_parts.append('<table>')
            for row in table.rows:
                current_page_parts.append('<tr>')
                for cell in row.cells:
                    tc = cell._tc
                    gridspan = tc.find(qn("w:tcPr"))
                    colspan = 1
                    rowspan = 1
                    if gridspan is not None:
                        gm = gridspan.find(qn("w:gridSpan"))
                        if gm is not None:
                            colspan = int(gm.get(qn("w:val"), "1"))
                        vm = gridspan.find(qn("w:vMerge"))
                        if vm is not None:
                            val = vm.get(qn("w:val"), "continue")
                            if val == "restart":
                                rowspan = 2
                            elif val == "continue":
                                continue
                    attrs = []
                    if colspan > 1: attrs.append(f'colspan="{colspan}"')
                    if rowspan > 1: attrs.append(f'rowspan="{rowspan}"')
                    current_page_parts.append(f'<td {" ".join(attrs)}>')
                    for cp in cell.paragraphs:
                        has_cell_images = bool(cp._element.findall(".//" + qn("w:drawing")))
                        if cp.text.strip() or has_cell_images:
                            current_page_parts.append('<p>')
                            for r in cp.runs:
                                current_page_parts.append(_render_run_images(r._element))
                                t = r.text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                                if r.bold: current_page_parts.append(f"<b>{t}</b>")
                                elif r.italic: current_page_parts.append(f"<i>{t}</i>")
                                else: current_page_parts.append(t)
                            current_page_parts.append('</p>')
                    current_page_parts.append('</td>')
                current_page_parts.append('</tr>')
            current_page_parts.append('</table>')

    # ---- Build final HTML: render each page directly without card shell ----
    for page_idx, page_parts in enumerate(all_pages):
        page_num = page_idx + 1
        parts.append(f'<div class="doc-page" data-page="{page_num}">')
        if header_html and page_idx == 0:
            parts.append(f'<div class="docx-header">{header_html}</div>')
        parts.extend(page_parts)
        if footer_html:
            parts.append(f'<div class="docx-footer">{footer_html}</div>')
        parts.append(f'<div class="page-num">{page_num} / {page_count}</div>')
        parts.append('</div>')

    parts.append('</body></html>')
    return "".join(parts)

def _convert_xlsx_to_html(input_path: str) -> str:
    """将 XLSX/XLS 转换为 HTML 表格（限制行数防 OOM）。"""
    import openpyxl
    MAX_ROWS = 5000  # 单表最大行数，超出截断
    wb = openpyxl.load_workbook(input_path, data_only=True)
    parts = [
        '<!DOCTYPE html><html lang="zh-CN"><head><meta charset="utf-8">',
        '<title>表格预览</title>',
        '<style>',
        'body{font-family:"Microsoft YaHei","SimSun",sans-serif;padding:20px;}',
        'h2{color:#333;}',
        'table{border-collapse:collapse;width:100%;margin-bottom:20px;}',
        'td,th{border:1px solid #ccc;padding:4px 8px;font-size:11pt;}',
        'tr:nth-child(even){background:#f9f9f9;}',
        '@media print{body{padding:0;}@page{size:A4 landscape;margin:10mm;}}',
        '</style></head><body>',
    ]
    for name in wb.sheetnames:
        ws = wb[name]
        safe_name = name.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;").replace('"', "&quot;")
        parts.append(f'<h2>{safe_name}</h2><table>')
        row_count = 0
        for row in ws.iter_rows(values_only=True):
            if row_count >= MAX_ROWS:
                parts.append(f'<tr><td colspan="20" style="color:#c00">[已截断，仅展示前 {MAX_ROWS} 行]</td></tr>')
                break
            parts.append('<tr>')
            for cell in row:
                val = str(cell) if cell is not None else ""
                val = val.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                parts.append(f'<td>{val}</td>')
            parts.append('</tr>')
            row_count += 1
        parts.append('</table>')
    parts.append('</body></html>')
    return "".join(parts)


# ── DOCX → images → HTML（持久化 + 并行 + JPEG，浏览器通用）──────────

def _format_preview_title(title: Optional[str] = None) -> str:
    return title or "文档预览"


def _build_images_html(file_id: str, image_paths: List[str], total: int, title: Optional[str] = None) -> str:
    """将磁盘图片列表构建为内嵌 HTML（按页展示，优化加载速度）。"""
    import base64 as b64

    preview_title = html_escape(_format_preview_title(title), quote=False)

    # 首页同步加载，其余懒加载
    img_tags = []
    for idx, ip in enumerate(image_paths):
        try:
            with open(ip, "rb") as fh:
                data = b64.b64encode(fh.read()).decode()
            src = f"data:image/jpeg;base64,{data}"
        except OSError:
            src = ""
        loading = "eager" if idx == 0 else "lazy"
        img_tags.append(
            f'<div class="page" data-page="{idx+1}">'
            f'<img src="{src}" alt="第{idx+1}页" '
            f'loading="{loading}" decoding="async" />'
            f'<div class="page-num">{idx+1}/{total}</div></div>'
        )

    css = (
        "html,body{margin:0;padding:0;background:#f5f7fb;"
        "font-family:\"Microsoft YaHei\",\"SimSun\",sans-serif;color:#111}"
        ".preview-shell{max-width:min(100%,980px);margin:0 auto;padding:16px 0 28px}"
        ".preview-title{text-align:center;font-size:20px;line-height:1.35;font-weight:700;"
        "margin:0 0 18px;color:#111}"
        ".page{margin:0 0 12px 0;text-align:center;contain:layout style paint}"
        ".page img{display:block;max-width:100%;width:auto;height:auto;margin:0 auto}"
        ".page-num{color:#999;font-size:11px;padding:4px 0 10px}"
        "@media print{body{background:#fff}.preview-shell{max-width:none;padding:0}.page{page-break-after:always}}"
        "@media(max-width:640px){.preview-shell{padding:12px 0 20px}.preview-title{font-size:18px}.page{margin-bottom:10px}}"
    )
    return (
        '<!DOCTYPE html><html lang="zh-CN"><head><meta charset="utf-8">'
        '<meta name="viewport" content="width=device-width,initial-scale=1">'
        f'<title>{preview_title}</title><style>{css}</style></head>'
        f'<body><div class="preview-shell"><h1 class="preview-title">{preview_title}</h1>{"".join(img_tags)}</div></body></html>'
    )


def _format_docx_line_height(line_spacing) -> Optional[str]:
    """
    Convert python-docx ParagraphFormat.line_spacing to safe CSS.

    python-docx returns either:
    - a small float multiplier (e.g. 1.5)
    - a Length/EMU integer for fixed spacing (e.g. Pt(32) == 406400)

    Writing the raw integer as unitless CSS creates enormous line heights and
    makes the preview look blank.  Clamp/normalize to browser-safe values.
    """
    if not line_spacing or line_spacing == 1.0:
        return None

    try:
        value = float(line_spacing)
    except (TypeError, ValueError):
        return None

    if value <= 0:
        return None

    # python-docx Length values expose .pt.  Some Length instances are int
    # subclasses, so check this before treating the value as a multiplier.
    pt = getattr(line_spacing, "pt", None)
    if pt:
        try:
            pt_value = float(pt)
        except (TypeError, ValueError):
            pt_value = 0
        if 0 < pt_value <= 200:
            return f"{pt_value:g}pt"

    # Defensive fallback for EMU integers that arrive without a .pt property.
    # 1 point = 12700 EMU.  The problematic document produced 406400 (=32pt).
    if value > 20:
        pt_value = value / 12700
        if 0 < pt_value <= 200:
            return f"{pt_value:g}pt"
        return None

    # Unitless CSS line-height multiplier.
    return f"{min(value, 5):g}"


def _ensure_pdf(
    file_id: str,
    source_path: str,
    source_hash: str,
    *,
    timeout_seconds: Optional[float] = None,
) -> Optional[str]:
    """
    确保 document_store 中存在 PDF 版本。
    缓存命中直接返回，否则通过 MS Word COM 转换并持久化。
    """
    from app.services.document_store import get_cached_pdf, store_pdf

    cached = get_cached_pdf(file_id, source_hash)
    if cached:
        return cached

    meta = _read_meta_internal(file_id)
    if meta.get("pdf_conversion_failed_hash") == source_hash:
        failed_at = float(meta.get("pdf_conversion_failed_at", 0) or 0)
        if time.time() - failed_at < 300:
            logger.warning(f"PDF 转换近期失败，暂不重试 Word 转换: {file_id}")
            return None

    # 优先尝试 MS Word COM，失败则回退到 LibreOffice
    if timeout_seconds is None:
        pdf_path = _convert_via_docx2pdf(source_path)
    else:
        pdf_path = _convert_via_docx2pdf(source_path, timeout_seconds=timeout_seconds)

    # 如果 MS Word COM 失败，尝试 LibreOffice 回退
    _tmp_dir = None
    if pdf_path is None:
        engine = _detect_engine()
        if engine == "libreoffice":
            _tmp_dir = tempfile.mkdtemp(prefix="enspdf_", dir=_ensure_temp_dir())
            try:
                pdf_path = _convert_via_libreoffice(source_path, _tmp_dir, "pdf")
            except Exception:
                pdf_path = None

    if pdf_path is None:
        meta = _read_meta_internal(file_id)
        meta["pdf_conversion_failed_hash"] = source_hash
        meta["pdf_conversion_failed_at"] = time.time()
        _write_meta_internal(file_id, meta)
        if _tmp_dir is not None:
            try:
                shutil.rmtree(_tmp_dir, ignore_errors=True)
            except Exception:
                pass
        return None


    try:
        stored = store_pdf(file_id, pdf_path, source_hash)
        meta = _read_meta_internal(file_id)
        meta.pop("pdf_conversion_failed_hash", None)
        meta.pop("pdf_conversion_failed_at", None)
        _write_meta_internal(file_id, meta)
        return stored
    finally:
        # 清理临时 PDF（docx2pdf 输出在 tempdir 下）
        parent = os.path.dirname(pdf_path)
        if _path_is_within_root(pdf_path, tempfile.gettempdir()):
            try:
                shutil.rmtree(parent, ignore_errors=True)
            except Exception:
                pass


def _ensure_images(file_id: str, pdf_path: str, page_count: int, pdf_hash: str,
                    dpi: Optional[int] = None) -> List[str]:
    """
    确保 document_store 中存在页面图片。
    缓存命中直接返回，否则并行生成 JPEG 并持久化。
    """
    from app.services.document_store import get_cached_images, generate_images, adaptive_dpi

    cached = get_cached_images(file_id, pdf_hash, page_count)
    if cached:
        return cached

    if dpi is None:
        dpi = adaptive_dpi(page_count)

    workers = min(os.cpu_count() or 4, 10)
    return generate_images(file_id, pdf_path, page_count, pdf_hash,
                           dpi=dpi, quality=75, max_workers=workers)


# ── 大文件骨架 HTML（图片通过 API URL 按需加载）───────────────────

LARGE_FILE_THRESHOLD = 50
MULTI_PAGE_THRESHOLD = LARGE_FILE_THRESHOLD  # 超过 50 页时使用轻量 HTML，避免 base64 页面过大/过慢


def build_skeleton_html(
    file_id: str,
    page_count: int,
    total: int,
    version: Optional[int] = None,
    page_url_prefix: Optional[str] = None,
    auth_token: Optional[str] = None,
    extra_query_params: Optional[dict[str, str]] = None,
    title: Optional[str] = None,
) -> str:
    """
    构建轻量骨架 HTML：图片通过 /pages/{n} API 端点按需加载。
    多页文档 ≈ KB 级响应（vs base64 内嵌 ≈ 数十 MB）。
    """
    params = []
    if version:
        params.append(f"version={version}")
    if auth_token:
        params.append(f"auth_token={quote_plus(auth_token)}")
    for key, value in (extra_query_params or {}).items():
        if value is None or value == "":
            continue
        params.append(f"{quote_plus(str(key))}={quote_plus(str(value))}")
    version_param = ("?" + "&".join(params)) if params else ""
    if page_url_prefix is None:
        page_url_prefix = f"/api/v1/files/{file_id}/pages"
    preview_title = html_escape(_format_preview_title(title), quote=False)

    css = (
        "html,body{margin:0;padding:0;background:#f5f7fb;"
        "font-family:\"Microsoft YaHei\",\"SimSun\",sans-serif;color:#111}"
        ".preview-shell{max-width:min(100%,980px);margin:0 auto;padding:16px 0 28px}"
        ".preview-title{text-align:center;font-size:20px;line-height:1.35;font-weight:700;"
        "margin:0 0 18px;color:#111}"
        ".page{margin:0 0 12px 0;text-align:center;min-height:auto;padding:0;"
        "display:block;background:transparent}"
        ".page img{display:block;max-width:100%;width:auto;height:auto;margin:0 auto}"
        ".page-num{color:#999;font-size:11px;padding:4px 0 10px}"
        ".page-loading{color:#aaa;font-size:13px}"
        "@media print{body{background:#fff}.preview-shell{max-width:none;padding:0}.page{page-break-after:always}}"
        "@media(max-width:640px){.preview-shell{padding:12px 0 20px}.preview-title{font-size:18px}.page{margin-bottom:10px}}"
    )

    pages = []
    for i in range(page_count):
        loading = "eager" if i == 0 else "lazy"
        page_url = f"{page_url_prefix}/{i+1}{version_param}"
        pages.append(
            f'<div class="page" data-page="{i+1}">'
            f'<img src="{page_url}" '
            f'loading="{loading}" decoding="async" '
            f'alt="第{i+1}页" />'
            f'</div>'
            f'<div class="page-num">{i+1} / {total}</div>'
        )

    return (
        '<!DOCTYPE html><html lang="zh-CN"><head><meta charset="utf-8">'
        '<meta name="viewport" content="width=device-width,initial-scale=1">'
        f'<title>{preview_title}</title><style>{css}</style></head>'
        f'<body><div class="preview-shell"><h1 class="preview-title">{preview_title}</h1>{"".join(pages)}</div></body></html>'
    )


def convert_to_images_html(
    file_id: str,
    input_path: str,
    file_type: str,
    return_skeleton: bool = False,
    page_url_prefix: Optional[str] = None,
    version: Optional[int] = None,
    auth_token: Optional[str] = None,
    title: Optional[str] = None,
) -> Optional[str]:
    """
    DOCX/PDF → 持久化图片 → HTML 页面。

    PDF 文件跳过 Word→PDF 步骤，直接从 PDF 生成图片。
    所有中间产物（PDF、图片）持久化到 document_store，
    下次预览同一文件秒开。

    Args:
        return_skeleton: True 时返回骨架 HTML（图片通过 API URL 加载），
                         用于多页文档场景。

    返回 HTML 字符串，失败返回 None。
    """
    import fitz

    file_type = file_type.lower().lstrip(".")
    if file_type not in ("docx", "doc", "pdf"):
        return None  # 调用方自行处理不支持类型

    from app.services.document_store import store_original, doc_root, adaptive_dpi
    from app.exceptions import ConversionError

    # 检测引擎可用性
    if file_type in ("docx", "doc") and _detect_engine() == "fallback":
        raise ConversionError(
            "文档预览不可用：系统未安装 Microsoft Word 或 LibreOffice，无法将 Word 转为 PDF",
            reason="no_word_engine",
        )

    # 绑定上下文日志
    log = get_logger(f"services.conversion.{file_id[:8]}")
    t_start = time.time()
    log.info(f"convert_to_images_html 开始 | type={file_type}")

    # 1. 原始文件持久化 + 哈希
    store_original(file_id, input_path)
    source_hash = _source_hash(input_path)
    log.info(f"原始文件哈希: {source_hash[:16]}")

    # 2. PDF 版本
    if file_type == "pdf":
        from app.services.document_store import _ensure_dirs, dir_pdf, _write_meta
        _ensure_dirs(file_id)
        pdf_dest = os.path.join(dir_pdf(file_id), "document.pdf")
        if os.path.normcase(os.path.realpath(input_path)) != os.path.normcase(os.path.realpath(pdf_dest)):
            shutil.copy2(input_path, pdf_dest)
        pdf_path = pdf_dest
        pdf_hash = source_hash
        meta = _read_meta_internal(file_id)
        meta["pdf_source_hash"] = source_hash
        _write_meta_internal(file_id, meta)
        log.info(f"PDF 即原始文件，已复制到 pdf/")
    else:
        t_pdf = time.time()
        pdf_path = _ensure_pdf(file_id, input_path, source_hash)
        if pdf_path is None:
            log.error("Word→PDF 转换失败")
            return None
        log.info(f"Word→PDF 完成 | {time.time()-t_pdf:.1f}s")
        pdf_hash = _source_hash(pdf_path)

    # 3. 页码
    doc = fitz.open(pdf_path)
    page_count = len(doc)
    doc.close()
    log.info(f"PDF 页数: {page_count}")

    if return_skeleton and page_count > MULTI_PAGE_THRESHOLD:
        html = build_skeleton_html(
            file_id,
            page_count,
            page_count,
            version=version,
            page_url_prefix=page_url_prefix,
            auth_token=auth_token,
            title=title,
        )
        elapsed = time.time() - t_start
        log.info(f"骨架 HTML（多页文档模式）| {page_count} 页 | {elapsed:.1f}s")
        return html

    # 4. 图片生成（持久化 + 缓存，自适应 DPI）
    t_img = time.time()
    dpi = adaptive_dpi(page_count)
    image_paths = _ensure_images(file_id, pdf_path, page_count, pdf_hash, dpi=dpi)
    log.info(f"图片生成完成 | {time.time()-t_img:.1f}s | {page_count} 页 | DPI={dpi}")

    # 5. 构建 HTML
    if return_skeleton and page_count > MULTI_PAGE_THRESHOLD:
        html = build_skeleton_html(
            file_id,
            page_count,
            page_count,
            version=version,
            page_url_prefix=page_url_prefix,
            auth_token=auth_token,
            title=title,
        )
        log.info(f"骨架 HTML（多页文档模式）| {page_count} 页")
    else:
        html = _build_images_html(file_id, image_paths, page_count, title=title)

    elapsed = time.time() - t_start
    total_kb = sum(os.path.getsize(p) for p in image_paths if os.path.exists(p)) // 1024
    log.info(f"convert_to_images_html 完成 | {page_count}页 {total_kb}KB {elapsed:.1f}s")

    return html


# 内联辅助（避免循环导入）
def _read_meta_internal(file_id: str) -> dict:
    import json
    from app.services.document_store import meta_path
    mp = meta_path(file_id)
    if os.path.exists(mp):
        try:
            with open(mp, "r", encoding="utf-8") as f:
                return json.load(f)
        except (json.JSONDecodeError, OSError):
            pass
    return {}


def _write_meta_internal(file_id: str, data: dict) -> None:
    import json
    from app.services.document_store import meta_path, _ensure_dirs
    _ensure_dirs(file_id)
    mp = meta_path(file_id)
    tmp = mp + ".tmp"
    with open(tmp, "w", encoding="utf-8") as f:
        json.dump(data, f, indent=2)
        f.flush()
        os.fsync(f.fileno())
    os.replace(tmp, mp)


# ── 工具函数 ────────────────────────────────────────────────────

def schedule_cleanup(background_tasks, file_path: str) -> None:
    """在响应发送后清理临时文件（及其父目录，如果为空）。"""
    def _cleanup():
        try:
            if os.path.exists(file_path):
                os.unlink(file_path)
            parent = os.path.dirname(file_path)
            if os.path.isdir(parent) and not os.listdir(parent):
                os.rmdir(parent)
        except OSError:
            pass
    background_tasks.add_task(_cleanup)


# ── 对外 API ────────────────────────────────────────────────────

def convert_to_pdf(input_path: str, file_type: str, filename: str) -> Tuple[str, str, str, bool]:
    """
    将文件转换为 PDF，优先使用原生引擎，结果缓存 30 天。

    转换优先级：
    1. 已有 PDF → 直接返回原文件
    2. docx2pdf (MS Word COM) → 像素级完美 PDF
    3. LibreOffice headless → 原生 PDF
    4. python-docx → 增强 HTML（浏览器打印为 PDF）

    Args:
        input_path: 源文件磁盘路径
        file_type: 文件类型 (pdf/docx/doc/xlsx/xls)
        filename: 原始文件名

    Returns:
        (output_path, media_type, actual_format, needs_cleanup)
    """
    file_type = file_type.lower().lstrip(".")
    _ensure_cache_dir()
    src_hash = _source_hash(input_path)

    # 0. 文件大小保护：超过 50MB 拒绝转换，直接返回原文件
    try:
        fsize = os.path.getsize(input_path)
        if fsize > 50 * 1024 * 1024:
            logger.warning(f"文件过大 ({fsize / 1024 / 1024:.1f}MB)，跳过转换")
            return (input_path, "application/octet-stream", file_type, False)
    except OSError:
        pass

    # 1. 已是 PDF → 直接返回
    if file_type == "pdf":
        return (input_path, "application/pdf", "pdf", False)

    # 2. 检查缓存
    cached = _read_cache(src_hash, ".pdf")
    if cached:
        return (cached, "application/pdf", "pdf", False)  # cached, no cleanup

    # 3. 原生引擎转换
    engine = _detect_engine()
    output_pdf = None

    if engine == "docx2pdf" and file_type in ("docx", "doc"):
        output_pdf = _convert_via_docx2pdf(input_path)

    if output_pdf is None and engine == "libreoffice":
        tmp_dir = tempfile.mkdtemp(prefix="conv_", dir=_ensure_temp_dir())
        try:
            output_pdf = _convert_via_libreoffice(input_path, tmp_dir, "pdf")
            if output_pdf is None:
                shutil.rmtree(tmp_dir, ignore_errors=True)
        except Exception:
            shutil.rmtree(tmp_dir, ignore_errors=True)
            output_pdf = None

    # 4. 原生引擎成功 → 写入缓存
    if output_pdf and os.path.exists(output_pdf) and os.path.getsize(output_pdf) > 0:
        try:
            cached_path = _write_cache(src_hash, output_pdf, ".pdf", file_type)
            # 清理临时输出（如果是临时文件）
            if output_pdf != cached_path and _path_is_within_root(output_pdf, tempfile.gettempdir()):
                try:
                    parent = os.path.dirname(output_pdf)
                    shutil.rmtree(parent, ignore_errors=True)
                except Exception:
                    pass
            return (cached_path, "application/pdf", "pdf", False)
        except Exception as e:
            logger.warning(f"缓存写入失败: {e}，直接返回转换结果")
            return (output_pdf, "application/pdf", "pdf", True)

    # 5. 所有引擎失败 → HTML fallback
    tmp_path = None
    try:
        if file_type in ("docx", "doc"):
            html = _convert_docx_to_html(input_path)
        elif file_type in ("xlsx", "xls"):
            html = _convert_xlsx_to_html(input_path)
        else:
            return (input_path, "application/octet-stream", file_type, False)

        fd, tmp_path = tempfile.mkstemp(suffix=".html", prefix="conv_", dir=_ensure_temp_dir())
        with os.fdopen(fd, "w", encoding="utf-8") as f:
            f.write(html)
        return (tmp_path, "text/html; charset=utf-8", "html", True)
    except Exception as e:
        logger.error(f"HTML 转换失败: {e}")
        if tmp_path and os.path.exists(tmp_path):
            try:
                os.unlink(tmp_path)
            except Exception:
                pass
        return (input_path, "application/octet-stream", file_type, False)


"""
DOCX2PDF 转换超时（秒）。
多页 Word 文档转换可能需要较长时间，生产环境建议 >= 300。
可通过环境变量 DOCX2PDF_TIMEOUT_SECONDS 覆盖。
"""
DOCX2PDF_TIMEOUT_SECONDS = int(os.environ.get("DOCX2PDF_TIMEOUT_SECONDS", "300"))




# =============================================================
#  Background pre-conversion pipeline:
#  Word -> PDF -> Images on upload, near-instant preview
# =============================================================
import threading
_preconvert_lock = threading.Lock()
_preconvert_running = set()  # type: set[str]


def trigger_preconversion(file_id: str, storage_path: str, file_type: str) -> None:
    """Run Word->PDF->Images pipeline in a background thread. Call on upload."""
    file_type = file_type.lower().lstrip(".")
    if file_type not in ("docx", "doc", "pdf"):
        return

    if file_id in _preconvert_running:
        return  # already running

    with _preconvert_lock:
        if file_id in _preconvert_running:
            return
        _preconvert_running.add(file_id)

    def _run():
        try:
            log = get_logger(f"preconvert.{file_id[:8]}")
            log.info(f"Pre-convert start | type={file_type}")

            # 1) Persist original into document_store
            from app.services.document_store import store_original
            store_original(file_id, storage_path)
            source_hash = _source_hash(storage_path)

            # 2) Word -> PDF
            if file_type in ("docx", "doc"):
                pdf_path = _ensure_pdf(file_id, storage_path, source_hash)
                if pdf_path is None:
                    log.error("Pre-convert failed: cannot generate PDF")
                    return
                pdf_hash = _source_hash(pdf_path)
            else:
                # PDF: copy to pdf/ directory
                from app.services.document_store import _ensure_dirs, dir_pdf
                _ensure_dirs(file_id)
                pdf_dest = os.path.join(dir_pdf(file_id), "document.pdf")
                if os.path.normcase(os.path.realpath(storage_path)) != \
                        os.path.normcase(os.path.realpath(pdf_dest)):
                    shutil.copy2(storage_path, pdf_dest)
                pdf_path = pdf_dest
                pdf_hash = source_hash

            # 3) PDF -> Images
            import fitz
            doc = fitz.open(pdf_path)
            page_count = len(doc)
            doc.close()

            if page_count == 0:
                log.warning("PDF has 0 pages, skip image generation")
                return

            from app.services.document_store import (
                get_cached_images, generate_images, adaptive_dpi
            )
            cached = get_cached_images(file_id, pdf_hash, page_count)
            if cached:
                log.info(f"Images already cached ({page_count} pages), skip")
                return

            dpi = adaptive_dpi(page_count)
            workers = min(os.cpu_count() or 4, 10)
            generate_images(file_id, pdf_path, page_count, pdf_hash,
                            dpi=dpi, quality=75, max_workers=workers)
            log.info(f"Pre-convert done | {page_count} pages | DPI={dpi}")

        except Exception:
            logger.exception("Pre-convert error")
        finally:
            with _preconvert_lock:
                _preconvert_running.discard(file_id)

    t = threading.Thread(target=_run, daemon=True, name=f"preconvert-{file_id[:8]}")
    t.start()


def _docx2pdf_worker(input_path: str, output_dir: str, result_queue) -> None:
    """隔离执行 Word COM，避免 WINWORD 进程影响 Web 主进程。"""
    try:
        import pythoncom
        import win32com.client

        abs_input = os.path.abspath(input_path)
        base = os.path.splitext(os.path.basename(abs_input))[0]
        pdf_path = os.path.join(output_dir, f"{base}.pdf")
        abs_pdf = os.path.abspath(pdf_path)

        pythoncom.CoInitialize()
        try:
            word = win32com.client.DispatchEx("Word.Application")
            word.Visible = False
            word.DisplayAlerts = 0
            doc = None
            try:
                doc = word.Documents.Open(abs_input, ReadOnly=True, AddToRecentFiles=False, ConfirmConversions=False)
                doc.ExportAsFixedFormat(abs_pdf, 17, OptimizeFor=1)
                doc.Close(SaveChanges=0)
                doc = None
            finally:
                if doc is not None:
                    try:
                        doc.Close(SaveChanges=0)
                    except Exception:
                        pass
                try:
                    word.Quit()
                except Exception:
                    pass
        finally:
            pythoncom.CoUninitialize()

        if os.path.exists(pdf_path) and os.path.getsize(pdf_path) > 0:
            result_queue.put({"ok": True, "path": pdf_path})
        else:
            result_queue.put({"ok": False, "error": "MS Word COM did not produce PDF"})
    except ImportError:
        result_queue.put({"ok": False, "error": "win32com is not available"})
    except Exception as exc:
        result_queue.put({"ok": False, "error": str(exc)})


def _start_docx2pdf_worker(input_path: str, output_dir: str):
    result_queue = multiprocessing.Queue(maxsize=1)
    process = multiprocessing.Process(
        target=_docx2pdf_worker,
        args=(input_path, output_dir, result_queue),
        daemon=True,
    )
    process.start()
    return process, result_queue


def _rels_source_dir(rels_name: str) -> str:
    """Return the OPC source part directory for a .rels zip entry."""
    normalized = rels_name.replace("\\", "/")
    if normalized == "_rels/.rels":
        return ""
    marker = "/_rels/"
    if marker not in normalized:
        return ""
    return normalized.split(marker, 1)[0]


def _placeholder_target_for_rels(rels_name: str) -> str:
    """Return a relationship Target path from rels_name to the shared placeholder image."""
    source_dir = _rels_source_dir(rels_name) or "."
    target = posixpath.relpath(_EXTERNAL_IMAGE_PLACEHOLDER_NAME, source_dir)
    return target.replace("\\", "/")


def _sanitize_external_image_relationships(rels_name: str, xml_bytes: bytes) -> tuple[bytes, bool]:
    """Replace external image relationships with a local DOCX placeholder target."""
    try:
        ET.register_namespace("", _REL_NS)
        root = ET.fromstring(xml_bytes)
    except ET.ParseError:
        return xml_bytes, False

    changed = False
    placeholder_target = _placeholder_target_for_rels(rels_name)
    for rel in root.findall(f"{{{_REL_NS}}}Relationship"):
        rel_type = rel.get("Type", "")
        target = rel.get("Target", "")
        target_mode = rel.get("TargetMode", "")
        is_image_rel = rel_type.endswith("/image") or "/relationships/image" in rel_type
        is_external = target_mode.lower() == "external" or target.lower().startswith(("http://", "https://"))
        if is_image_rel and is_external:
            rel.set("Target", placeholder_target)
            rel.attrib.pop("TargetMode", None)
            changed = True

    if not changed:
        return xml_bytes, False

    sanitized = ET.tostring(root, encoding="utf-8", xml_declaration=True)
    return sanitized, True


def _ensure_png_content_type(xml_bytes: bytes) -> tuple[bytes, bool]:
    """Ensure [Content_Types].xml declares PNG image parts."""
    try:
        ET.register_namespace("", _CONTENT_TYPES_NS)
        root = ET.fromstring(xml_bytes)
    except ET.ParseError:
        return xml_bytes, False

    ns = ""
    if root.tag.startswith("{") and "}" in root.tag:
        ns = root.tag[1:].split("}", 1)[0]
    default_tag = f"{{{ns}}}Default" if ns else "Default"

    for default in root.findall(default_tag):
        if default.get("Extension", "").lower() == "png":
            return xml_bytes, False

    ET.SubElement(root, default_tag, Extension="png", ContentType="image/png")
    updated = ET.tostring(root, encoding="utf-8", xml_declaration=True)
    return updated, True


def _prepare_docx_for_word_conversion(input_path: str) -> tuple[str, Optional[str]]:
    """
    Create a temporary DOCX with external image links replaced by a local placeholder.

    Word COM may block for minutes when opening DOCX files containing unreachable
    external image relationships. Sanitizing those links keeps conversion local
    while preserving the document structure enough for PDF preview generation.
    Returns (path_to_use, temp_dir_to_cleanup).
    """
    if not input_path.lower().endswith(".docx") or not zipfile.is_zipfile(input_path):
        return input_path, None

    changed = False
    temp_dir = tempfile.mkdtemp(prefix="docx_word_safe_", dir=_ensure_temp_dir())
    output_path = os.path.join(temp_dir, os.path.basename(input_path))

    try:
        entries = []
        with zipfile.ZipFile(input_path, "r") as src:
            existing_names = set(src.namelist())
            for info in src.infolist():
                data = src.read(info.filename)
                if info.filename.endswith(".rels"):
                    data, rel_changed = _sanitize_external_image_relationships(info.filename, data)
                    changed = changed or rel_changed
                entries.append((info, data))

        if changed:
            updated_entries = []
            has_content_types = False
            for info, data in entries:
                if info.filename == "[Content_Types].xml":
                    has_content_types = True
                    data, _ = _ensure_png_content_type(data)
                updated_entries.append((info, data))
            entries = updated_entries

            with zipfile.ZipFile(output_path, "w", zipfile.ZIP_DEFLATED) as dst:
                for info, data in entries:
                    dst.writestr(info, data)

                if not has_content_types:
                    content_types = (
                        b'<?xml version="1.0" encoding="UTF-8"?>'
                        b'<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">'
                        b'<Default Extension="png" ContentType="image/png"/>'
                        b"</Types>"
                    )
                    dst.writestr("[Content_Types].xml", content_types)
                if _EXTERNAL_IMAGE_PLACEHOLDER_NAME not in existing_names:
                    dst.writestr(_EXTERNAL_IMAGE_PLACEHOLDER_NAME, _EXTERNAL_IMAGE_PLACEHOLDER_BYTES)

            logger.info(f"DOCX external image links sanitized for Word conversion: {input_path}")
            return output_path, temp_dir
    except Exception as exc:
        logger.warning(f"DOCX external image sanitization failed, using original file: {exc}")
    shutil.rmtree(temp_dir, ignore_errors=True)
    return input_path, None


def _convert_via_docx2pdf(input_path: str, timeout_seconds: Optional[float] = None) -> Optional[str]:
    """用 MS Word COM 将 DOCX/DOC 转为 PDF，失败返回 None。"""
    output_dir = tempfile.mkdtemp(prefix="docx2pdf_", dir=_ensure_temp_dir())
    timeout = DOCX2PDF_TIMEOUT_SECONDS if timeout_seconds is None else timeout_seconds
    process = None
    prepared_input = input_path
    prepared_cleanup_dir = None
    try:
        prepared_input, prepared_cleanup_dir = _prepare_docx_for_word_conversion(input_path)
        process, result_queue = _start_docx2pdf_worker(prepared_input, output_dir)
        process.join(timeout)
        if process.is_alive():
            logger.warning(f"MS Word COM 转换超时({timeout}s)，终止: {input_path}")
            process.terminate()
            process.join(3)
            shutil.rmtree(output_dir, ignore_errors=True)
            return None

        try:
            result = result_queue.get_nowait()
        except queue.Empty:
            result = {"ok": False, "error": "Word conversion worker returned no result"}
        except Exception as exc:
            result = {"ok": False, "error": str(exc)}

        if result.get("ok"):
            pdf_path = result.get("path")
            if pdf_path and os.path.exists(pdf_path) and os.path.getsize(pdf_path) > 0:
                logger.info(f"MS Word COM PDF 转换成功: {pdf_path}")
                return pdf_path

        logger.warning(f"MS Word COM 转换失败: {result.get('error')}")
        shutil.rmtree(output_dir, ignore_errors=True)
        return None
    except Exception as e:
        logger.warning(f"MS Word COM 转换异常: {e}")
        if process is not None and process.is_alive():
            try:
                process.terminate()
                process.join(3)
            except Exception:
                pass
        shutil.rmtree(output_dir, ignore_errors=True)
        return None
    finally:
        if prepared_cleanup_dir is not None:
            shutil.rmtree(prepared_cleanup_dir, ignore_errors=True)


def convert_to_html(input_path: str, file_type: str = "docx", title: Optional[str] = None) -> Tuple[str, str, bool]:
    """
    将 DOCX 转换为 HTML 用于预览。
    使用 python-docx 增强渲染（图片 base64 内嵌，无需外部文件）。
    
    Returns:
        (html_string, media_type, needs_cleanup)
    """
    file_type = file_type.lower().lstrip(".")
    if file_type in ("docx", "doc"):
        html = _convert_docx_to_html(input_path)
    elif file_type in ("xlsx", "xls"):
        html = _convert_xlsx_to_html(input_path)
    else:
        html = "<p>不支持此文件类型的预览</p>"
    if title:
        html = html.replace("<title>文档预览</title>", f"<title>{html_escape(title, quote=False)}</title>")
        html = html.replace("<title>表格预览</title>", f"<title>{html_escape(title, quote=False)}</title>")
    return (html, "text/html; charset=utf-8", False)


def _wrap_word_html(html_path: str, output_dir: str, title: str) -> str:
    """读取 Word 导出的 HTML，嵌入图片为 base64，包裹为带样式控制的预览页面。"""
    import re
    import base64 as b64

    with open(html_path, "r", encoding="utf-8", errors="replace") as f:
        raw = f.read()

    # ── 将外部图片引用转为 base64 data URI ──
    def _replace_img_src(match):
        src = match.group(1)
        # 跳过已经是 data: 或 http: 的
        if src.startswith("data:") or src.startswith("http"):
            return match.group(0)
        # 尝试在 output_dir 下找到图片文件
        # Word 导出图片通常在 {basename}.files/ 子目录
        img_path = os.path.join(output_dir, src)
        if not os.path.exists(img_path):
            # 也尝试直接在 output_dir 下查找
            alt_path = os.path.join(output_dir, os.path.basename(src))
            if os.path.exists(alt_path):
                img_path = alt_path
            else:
                return match.group(0)  # 找不到，保留原样
        try:
            with open(img_path, "rb") as img_f:
                img_bytes = img_f.read()
            ext = os.path.splitext(img_path)[1].lower()
            mime_map = {".png": "image/png", ".jpg": "image/jpeg", ".jpeg": "image/jpeg",
                        ".gif": "image/gif", ".bmp": "image/bmp", ".svg": "image/svg+xml",
                        ".webp": "image/webp", ".wmf": "image/wmf", ".emf": "image/emf"}
            mime = mime_map.get(ext, "image/png")
            data_uri = f"data:{mime};base64,{b64.b64encode(img_bytes).decode()}"
            return match.group(0).replace(src, data_uri)
        except Exception:
            return match.group(0)

    raw = re.sub(r'src="([^"]*)"', _replace_img_src, raw)

    # 提取 <body> 内容
    body_match = re.search(r"<body[^>]*>(.*?)</body>", raw, re.DOTALL | re.IGNORECASE)
    body = body_match.group(1) if body_match else raw

    # 提取 <style> 标签（Word 的内联样式）
    style_match = re.search(r"<style[^>]*>(.*?)</style>", raw, re.DOTALL | re.IGNORECASE)
    word_styles = style_match.group(1) if style_match else ""

    return f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="utf-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>{title}</title>
<style>
body {{
    margin: 0;
    padding: 0;
    background: #fff;
    font-family: "Times New Roman", "SimSun", "Microsoft YaHei", serif;
    line-height: 1.6;
    font-size: 12pt;
    color: #111;
}}
p {{ margin: 0 0 6pt 0; }}
h1 {{ font-size: 22pt; margin: 16pt 0 8pt; }}
h2 {{ font-size: 16pt; margin: 14pt 0 6pt; }}
h3 {{ font-size: 14pt; margin: 12pt 0 4pt; }}
table {{ border-collapse: collapse; width: 100%; margin: 8pt 0; }}
td, th {{ border: 1px solid #ccc; padding: 4pt 8pt; }}
img {{ display: block; max-width: 100%; height: auto; }}

/* Word 原始样式 */
{word_styles}

/* 打印优化 */
@media print {{
    body {{ background: #fff; }}
    @page {{ size: A4; margin: 20mm; }}
}}
</style>
</head>
<body>
{body}
</body>
</html>"""


def convert_to_word(input_path: str, file_type: str, filename: str) -> Tuple[str, str, str, bool]:
    """
    获取文件的 Word 格式版本。
    如果已是 DOCX/DOC → 直接返回；PDF → 暂不支持反向转换。
    """
    file_type = file_type.lower().lstrip(".")
    if file_type in ("docx", "doc"):
        return (input_path, "application/vnd.openxmlformats-officedocument.wordprocessingml.document", file_type, False)
    return (input_path, "application/octet-stream", file_type, False)
