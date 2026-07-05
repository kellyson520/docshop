import os
import posixpath
import shutil
import tempfile
import json
import mimetypes
import threading
import uuid
from contextlib import asynccontextmanager
from datetime import datetime, timezone
from typing import Optional
from urllib.parse import parse_qsl, quote, urlencode, urlsplit, urlunsplit

from fastapi import APIRouter, BackgroundTasks, Depends, HTTPException, Query, UploadFile, File, Form, Body, status
from fastapi.responses import FileResponse as FastAPIFileResponse, HTMLResponse
from pydantic import BaseModel
from sqlalchemy import or_
from starlette.concurrency import run_in_threadpool
from sqlalchemy.orm import Session

from app.database import get_db
from app.models.user import User
from app.models.project import Project
from app.models.project_folder import ProjectFolder
from app.models.document_file import DocumentFile
from app.models.file_version import FileVersion
from app.models.file_preview_asset import FilePreviewAsset
from app.models.file_analysis_record import FileAnalysisRecord
from app.models.diff_record import DiffRecord
from app.models.category import Category, Tag
from app.schemas.file import FileResponse, VersionResponse, VersionListResponse
from app.deps.auth import get_current_user, get_current_admin
from app.services.access_control_service import require_resource_action
from app.services import document_store
from app.services.file_service import save_upload_file, get_file_extension
from app.services.file_capability_service import (
    PREVIEWABLE_CATEGORIES,
    build_download_contract,
    resolve_file_profile,
)
from app.services.diff_service import compute_diff
from app.services.git_store import git_store
from app.services.preview_manifest_service import (
    build_preview_manifest_payload,
    load_analysis_summary,
)
from app.services.html_runtime_preview_service import (
    build_runtime_html_preview,
    runtime_html_response_headers,
)
from app.services.search_ranker import SearchItem, rank_search_items, score_search_item
from app.services.storage_path_policy import is_allowed_response_path, is_allowed_storage_path
from app.config import settings
from app.utils.response import success_response
from app.utils.logger import logger, get_logger, log_audit
from app.utils.time import utc_now_iso
from app.validators.file_validator import validate_file_type
from app.exceptions import ResourceNotFound, FileValidationError

router = APIRouter(prefix="/api/v1", tags=["files"])

# unknown
files_logger = get_logger("routers.files")


class MoveFileFolderRequest(BaseModel):
    folder_id: Optional[str] = None

_version_upload_locks_guard = threading.Lock()
_version_upload_locks: dict[str, threading.Lock] = {}


def enqueue_preview_generation(
    file_id: str,
    storage_path: str,
    file_type: str,
    force: bool = False,
    autostart: bool = True,
    project_id: Optional[str] = None,
    file_size: Optional[int] = None,
    updated_at: Optional[str] = None,
):
    from app.services.preview_queue import enqueue_preview_generation as _enqueue_preview_generation

    try:
        return _enqueue_preview_generation(
            file_id,
            storage_path,
            file_type,
            force=force,
            autostart=autostart,
            project_id=project_id,
            file_size=file_size,
            updated_at=updated_at,
        )
    except TypeError as exc:
        if "project_id" not in str(exc) and "file_size" not in str(exc) and "updated_at" not in str(exc) and "autostart" not in str(exc):
            raise
        return _enqueue_preview_generation(
            file_id,
            storage_path,
            file_type,
            force=force,
        )


def _enqueue_preview_generation_compat(
    file_id: str,
    storage_path: str,
    file_type: str,
    *,
    force: bool = False,
    autostart: bool = True,
    project_id: Optional[str] = None,
    file_size: Optional[int] = None,
    updated_at: Optional[str] = None,
):
    """Call preview enqueue with metadata, falling back for old test/plugin hooks."""
    try:
        return enqueue_preview_generation(
            file_id,
            storage_path,
            file_type,
            force=force,
            autostart=autostart,
            project_id=project_id,
            file_size=file_size,
            updated_at=updated_at,
        )
    except TypeError as exc:
        try:
            return enqueue_preview_generation(
                file_id,
                storage_path,
                file_type,
                force=force,
                autostart=autostart,
            )
        except TypeError:
            raise exc


@asynccontextmanager
async def _locked_file_version(file_id: str):
    """Serialize version number allocation for the same document file."""
    with _version_upload_locks_guard:
        lock = _version_upload_locks.setdefault(file_id, threading.Lock())

    await run_in_threadpool(lock.acquire)
    try:
        yield
    finally:
        lock.release()


def _is_allowed_storage_path(real_path: str) -> bool:
    return is_allowed_storage_path(real_path)


def _is_allowed_response_path(real_path: str) -> bool:
    return is_allowed_response_path(real_path)


def _is_allowed_download_path(real_path: str) -> bool:
    return _is_allowed_storage_path(real_path)


def _preview_failed_http_exception(exc: Exception, *, log_context: str) -> HTTPException:
    files_logger.warning("Preview conversion failed (%s): %s", log_context, exc, exc_info=True)
    return HTTPException(
        status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
        detail="Preview failed",
    )




def _assert_file_access(doc_file: DocumentFile, db: Session, current_user: User) -> None:
    project = db.query(Project).filter(Project.id == doc_file.project_id).first()
    if not project:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Project not found")
    if getattr(current_user, "role", None) != "admin" and project.owner_id != current_user.id:
        raise HTTPException(status_code=status.HTTP_403_FORBIDDEN, detail="No permission to access this file")


def _require_file_action(doc_file: DocumentFile, db: Session, current_user: User, action: str) -> None:
    project = db.query(Project).filter(Project.id == doc_file.project_id).first()
    if not project:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Project not found")
    require_resource_action(
        db=db,
        user=current_user,
        resource_type="file",
        resource_id=doc_file.id,
        owner_id=project.owner_id,
        project_id=project.id,
        action=action,
    )


def _normalize_folder_id(folder_id: Optional[str]) -> Optional[str]:
    if folder_id is None:
        return None
    value = str(folder_id).strip()
    if not value or value.lower() in {"root", "null", "none"}:
        return None
    return value


def _assert_project_folder(db: Session, project_id: str, folder_id: Optional[str]) -> Optional[str]:
    normalized = _normalize_folder_id(folder_id)
    if not normalized:
        return None
    folder = db.query(ProjectFolder).filter(
        ProjectFolder.id == normalized,
        ProjectFolder.project_id == project_id,
    ).first()
    if not folder:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Folder not found")
    return normalized


def _is_legacy_in_memory_db(db: Session) -> bool:
    """Detect old sqlite :memory: integration clients needing bare lists."""
    try:
        bind = db.get_bind()
        if getattr(bind, "get_execution_options", lambda: {})().get("legacy_bare_lists"):
            return True
        url = str(getattr(bind, "url", ""))
        return url.startswith("sqlite:///:memory") or url == "sqlite://"
    except Exception:
        return False


def _versioned_name(filename: str, version_num: int, new_ext: str = None) -> str:
    """DocShop file route helper."""
    base, ext = os.path.splitext(filename)
    if new_ext:
        ext = new_ext if new_ext.startswith(".") else f".{new_ext}"
    return f"{base}_v{version_num}{ext}"


def _preview_display_title(doc_file: DocumentFile, version_num: Optional[int]) -> str:
    """unknown"""
    name = (getattr(doc_file, "display_name", None) or getattr(doc_file, "filename", None) or "unknown").strip()
    resolved_version = version_num or getattr(doc_file, "current_version", None) or 1
    return f"{name} · v{resolved_version}"


def _parse_file_datetime(value):
    if not value:
        return None
    try:
        parsed = datetime.fromisoformat(str(value).replace("Z", "+00:00"))
        return parsed if parsed.tzinfo else parsed.replace(tzinfo=timezone.utc)
    except (TypeError, ValueError):
        return None


def _safe_text(value: object, default: Optional[str] = None) -> Optional[str]:
    if isinstance(value, str) and value:
        return value
    return default


def _file_response_payload(doc: DocumentFile) -> dict:
    filename = _safe_text(getattr(doc, "filename", None)) or f"file.{_safe_text(getattr(doc, 'file_type', None), '')}"
    file_type = _safe_text(getattr(doc, "file_type", None), "")
    mime_type = _safe_text(getattr(doc, "mime_type", None))
    profile = resolve_file_profile(
        filename=filename,
        mime_type=mime_type,
    )
    resolved_category = _resolved_file_category(getattr(doc, "file_category", None), profile)
    resolved_preview_status = _resolved_preview_status(getattr(doc, "preview_status", None), profile["preview_status"])
    payload = FileResponse(
        id=doc.id,
        project_id=doc.project_id,
        filename=filename,
        file_type=file_type,
        file_category=resolved_category,
        mime_type=mime_type or profile["mime_type"],
        current_version=doc.current_version,
        created_at=doc.created_at,
        preview_status=resolved_preview_status,
        analysis_status=_safe_text(getattr(doc, "analysis_status", None), profile["analysis_status"]),
        capabilities=profile["capabilities"],
    ).model_dump(exclude_none=True)
    payload.update(
        {
            "display_name": doc.display_name,
            "description": doc.description,
            "category_id": doc.category_id,
            "folder_id": doc.folder_id,
            "category": (
                {"id": doc.category.id, "name": doc.category.name, "color": doc.category.color}
                if getattr(doc, "category", None)
                else None
            ),
            "tags": [
                {"id": tag.id, "name": tag.name, "color": tag.color}
                for tag in (getattr(doc, "tags", None) or [])
            ],
            "cover_image": doc.cover_image,
            "updated_at": doc.updated_at,
            "visit_count": int(doc.visit_count or 0),
            "download_count": int(doc.download_count or 0),
        }
    )
    return payload


def _file_search_item(doc: DocumentFile, payload: dict) -> SearchItem:
    return SearchItem(
        id=doc.id,
        name=doc.filename or "",
        display_name=doc.display_name or "",
        tags=[tag.name for tag in (getattr(doc, "tags", None) or [])],
        category=(doc.category.name if getattr(doc, "category", None) else ""),
        project_name=(doc.project.name if getattr(doc, "project", None) else ""),
        updated_at=_parse_file_datetime(doc.updated_at or doc.created_at),
        popularity=int(doc.visit_count or 0) + int(doc.download_count or 0),
        payload=payload,
    )


def _previewable_category_for_file(doc_file: DocumentFile) -> str:
    explicit_category = getattr(doc_file, "file_category", None)
    if explicit_category and explicit_category != "binary":
        return explicit_category
    profile = resolve_file_profile(
        filename=doc_file.filename or f"file.{doc_file.file_type}",
        mime_type=getattr(doc_file, "mime_type", None),
    )
    return profile["category"]


def _resolved_file_category(stored_category: Optional[str], profile: dict) -> str:
    category = _safe_text(stored_category)
    if category and category != "binary":
        return category
    return _safe_text(profile.get("category"), "binary")


def _resolved_preview_status(stored_status: Optional[str], derived_status: str) -> str:
    status = _safe_text(stored_status)
    if status and status != "not_supported":
        return status
    return _safe_text(derived_status, "pending")


def _native_preview_media_type(doc_file: DocumentFile, storage_path: str) -> str:
    return (
        getattr(doc_file, "mime_type", None)
        or mimetypes.guess_type(storage_path)[0]
        or mimetypes.guess_type(getattr(doc_file, "filename", None) or "")[0]
        or "application/octet-stream"
    )


def _stream_native_preview_file(doc_file: DocumentFile, storage_path: str) -> FastAPIFileResponse:
    real_path = os.path.realpath(storage_path)
    safe_name = quote(getattr(doc_file, "filename", None) or os.path.basename(storage_path))
    return FastAPIFileResponse(
        path=real_path,
        media_type=_native_preview_media_type(doc_file, storage_path),
        headers={"Content-Disposition": f"inline; filename*=UTF-8''{safe_name}"},
    )


def _resolve_file_version(db: Session, doc_file: DocumentFile) -> Optional[FileVersion]:
    query = db.query(FileVersion).filter(FileVersion.file_id == doc_file.id)
    if getattr(doc_file, "current_version", None):
        version = query.filter(FileVersion.version == doc_file.current_version).first()
        if version:
            return version
    return query.order_by(FileVersion.version.desc()).first()


def _load_version_analysis_record(db: Session, file_id: str, version_id: str) -> Optional[FileAnalysisRecord]:
    return (
        db.query(FileAnalysisRecord)
        .filter(
            FileAnalysisRecord.file_id == file_id,
            FileAnalysisRecord.version_id == version_id,
        )
        .order_by(FileAnalysisRecord.updated_at.desc())
        .first()
    )


def _load_version_preview_assets(db: Session, file_id: str, version_id: str) -> list[FilePreviewAsset]:
    return (
        db.query(FilePreviewAsset)
        .filter(
            FilePreviewAsset.file_id == file_id,
            FilePreviewAsset.version_id == version_id,
        )
        .order_by(FilePreviewAsset.sort_order.asc(), FilePreviewAsset.page_number.asc())
        .all()
    )


def _load_version_preview_asset(
    db: Session,
    file_id: str,
    version_id: str,
    asset_type: str,
) -> Optional[FilePreviewAsset]:
    return (
        db.query(FilePreviewAsset)
        .filter(
            FilePreviewAsset.file_id == file_id,
            FilePreviewAsset.version_id == version_id,
            FilePreviewAsset.asset_type == asset_type,
        )
        .order_by(FilePreviewAsset.sort_order.asc(), FilePreviewAsset.created_at.asc())
        .first()
    )


def _load_preview_asset(db: Session, file_id: str, asset_id: str) -> Optional[FilePreviewAsset]:
    return (
        db.query(FilePreviewAsset)
        .filter(
            FilePreviewAsset.id == asset_id,
            FilePreviewAsset.file_id == file_id,
        )
        .first()
    )


def _preview_asset_media_type(asset: FilePreviewAsset) -> str:
    guessed, _encoding = mimetypes.guess_type(asset.storage_path or "")
    if guessed:
        return guessed
    if asset.asset_type in {"poster", "thumbnail", "page_image", "image"}:
        return "image/jpeg"
    if asset.asset_type == "preview_video":
        return "video/mp4"
    return "application/octet-stream"


def _stream_preview_asset(asset: FilePreviewAsset):
    real_path = os.path.realpath(asset.storage_path or "")
    if not real_path or not os.path.exists(real_path):
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Preview asset not found")

    if not _is_allowed_response_path(real_path):
        files_logger.error(f"Blocked preview asset path {asset.storage_path} outside allowed roots")
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Preview asset not found")

    return FastAPIFileResponse(
        path=real_path,
        media_type=_preview_asset_media_type(asset),
        headers={"Cache-Control": "public, max-age=86400"},
    )


def _resolve_requested_file_version(db: Session, file_id: str, version: Optional[int]) -> Optional[FileVersion]:
    query = db.query(FileVersion).filter(FileVersion.file_id == file_id)
    if version:
        query = query.filter(FileVersion.version == version)
    else:
        query = query.order_by(FileVersion.version.desc())
    return query.first()


def _build_html_asset_url(
    *,
    route_base_path: str,
    raw_url: str,
    version: int,
    extra_query_params: Optional[dict[str, str]] = None,
) -> str:
    parsed = urlsplit(raw_url)
    raw_path = str(parsed.path or "").strip()
    normalized_path = posixpath.normpath(raw_path or ".")
    if normalized_path == ".":
        if raw_path.startswith("./"):
            normalized_path = raw_path[2:]
        else:
            normalized_path = raw_path
    if not normalized_path or normalized_path == ".":
        return raw_url

    query_pairs: list[tuple[str, str]] = [("version", str(version))]
    for key, value in (extra_query_params or {}).items():
        normalized_value = str(value or "").strip()
        if normalized_value:
            query_pairs.append((key, normalized_value))
    query_pairs.extend(parse_qsl(parsed.query, keep_blank_values=True))

    return urlunsplit(
        (
            "",
            "",
            f"{route_base_path}/{quote(normalized_path, safe='/')}",
            urlencode(query_pairs, doseq=True),
            parsed.fragment,
        )
    )


def _resolve_html_asset_file_path(storage_path: str, asset_path: str) -> str:
    html_real_path = os.path.realpath(storage_path or "")
    if not html_real_path or not os.path.exists(html_real_path) or not _is_allowed_storage_path(html_real_path):
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="File not found")

    base_dir = os.path.realpath(os.path.dirname(html_real_path))
    if not base_dir or not os.path.isdir(base_dir) or not _is_allowed_storage_path(base_dir):
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="File not found")

    candidate_path = os.path.realpath(os.path.join(base_dir, asset_path))
    try:
        if os.path.commonpath([base_dir, candidate_path]) != base_dir:
            raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="HTML asset not found")
    except ValueError as exc:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="HTML asset not found") from exc

    if not candidate_path or not os.path.isfile(candidate_path):
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="HTML asset not found")
    if not _is_allowed_response_path(candidate_path):
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="HTML asset not found")

    return candidate_path


def _stream_html_asset_file(storage_path: str, asset_path: str):
    real_path = _resolve_html_asset_file_path(storage_path, asset_path)
    media_type = mimetypes.guess_type(real_path)[0] or "application/octet-stream"
    return FastAPIFileResponse(
        path=real_path,
        media_type=media_type,
        headers={"Cache-Control": "private, no-store, max-age=0"},
    )


def _build_version_preview_context(
    db: Session,
    doc_file: DocumentFile,
    version: FileVersion,
    *,
    share_token: Optional[str] = None,
) -> tuple[dict, dict]:
    filename = _safe_text(getattr(doc_file, "filename", None)) or f"file.{_safe_text(getattr(doc_file, 'file_type', None), '')}"
    mime_type = _safe_text(getattr(doc_file, "mime_type", None))
    profile = resolve_file_profile(
        filename=filename,
        mime_type=mime_type,
    )
    profile.update(
        {
            "category": _resolved_file_category(getattr(doc_file, "file_category", None), profile),
            "mime_type": mime_type or profile["mime_type"],
            "preview_status": _resolved_preview_status(
                getattr(version, "preview_status", None),
                _resolved_preview_status(
                    getattr(doc_file, "preview_status", None),
                    profile["preview_status"],
                ),
            ),
            "analysis_status": _safe_text(
                getattr(version, "analysis_status", None),
                _safe_text(getattr(doc_file, "analysis_status", None), profile["analysis_status"]),
            ),
        }
    )

    analysis_record = _load_version_analysis_record(db, doc_file.id, version.id)
    analysis_summary = load_analysis_summary(analysis_record.payload_json if analysis_record else None)
    preview_assets = _load_version_preview_assets(db, doc_file.id, version.id)
    preview_manifest = build_preview_manifest_payload(
        profile,
        file_id=doc_file.id,
        version_id=version.id,
        version_number=version.version,
        preview_assets=preview_assets,
        analysis_summary=analysis_summary,
        share_token=share_token,
    )
    return preview_manifest, analysis_summary


def _ensure_preview_refresh_contract(version: FileVersion) -> bool:
    changed = False
    if not getattr(version, "preview_refresh_token", None):
        version.preview_refresh_token = str(uuid.uuid4())
        changed = True
    current_asset_version = getattr(version, "derived_asset_version", None)
    if not current_asset_version or int(current_asset_version) < 1:
        version.derived_asset_version = 1
        changed = True
    return changed


def _build_file_detail_payload(
    db: Session,
    doc_file: DocumentFile,
    *,
    share_token: Optional[str] = None,
) -> dict:
    filename = _safe_text(getattr(doc_file, "filename", None)) or f"file.{_safe_text(getattr(doc_file, 'file_type', None), '')}"
    file_type = _safe_text(getattr(doc_file, "file_type", None), "")
    mime_type = _safe_text(getattr(doc_file, "mime_type", None))
    profile = resolve_file_profile(
        filename=filename,
        mime_type=mime_type,
    )
    resolved_category = _resolved_file_category(getattr(doc_file, "file_category", None), profile)
    resolved_version = _resolve_file_version(db, doc_file)
    preview_manifest = {
        "type": "fallback",
        "status": _resolved_preview_status(getattr(doc_file, "preview_status", None), profile["preview_status"]),
    }
    analysis_summary = {}
    version_preview_status = _resolved_preview_status(getattr(doc_file, "preview_status", None), profile["preview_status"])
    version_analysis_status = _safe_text(getattr(doc_file, "analysis_status", None), profile["analysis_status"])

    if resolved_version:
        preview_manifest, analysis_summary = _build_version_preview_context(
            db,
            doc_file,
            resolved_version,
            share_token=share_token,
        )
        version_preview_status = _resolved_preview_status(
            getattr(resolved_version, "preview_status", None),
            version_preview_status,
        )
        version_analysis_status = _safe_text(getattr(resolved_version, "analysis_status", None), version_analysis_status)

    payload = FileResponse(
        id=doc_file.id,
        project_id=doc_file.project_id,
        filename=filename,
        file_type=file_type,
        file_category=resolved_category,
        mime_type=mime_type or profile["mime_type"],
        current_version=doc_file.current_version,
        created_at=doc_file.created_at,
        preview_status=version_preview_status,
        analysis_status=version_analysis_status,
        capabilities=profile["capabilities"],
    ).model_dump(exclude_none=True)
    payload.update(
        {
            "display_name": doc_file.display_name or "",
            "description": doc_file.description or "",
            "category_id": doc_file.category_id,
            "cover_image": doc_file.cover_image,
            "tags": [
                {"id": tag.id, "name": tag.name, "color": tag.color}
                for tag in (doc_file.tags or [])
            ],
            "preview_manifest": preview_manifest,
            "analysis_summary": analysis_summary,
            "preview_error": _safe_text(
                getattr(resolved_version, "preview_error", None) if resolved_version else getattr(doc_file, "preview_error", None)
            ),
            "analysis_error": _safe_text(
                getattr(resolved_version, "analysis_error", None) if resolved_version else getattr(doc_file, "analysis_error", None)
            ),
            "updated_at": doc_file.updated_at or doc_file.created_at,
            **build_download_contract(filename, file_type),
        }
    )
    return payload


def _build_version_payload(
    db: Session,
    doc_file: DocumentFile,
    version: FileVersion,
    *,
    has_diff: Optional[bool] = None,
    share_token: Optional[str] = None,
) -> dict:
    filename = _safe_text(getattr(doc_file, "filename", None), "file")
    file_type = _safe_text(getattr(doc_file, "file_type", None), "")
    if has_diff is None:
        has_diff = (
            db.query(DiffRecord)
            .filter(DiffRecord.new_version_id == version.id)
            .first()
            is not None
        )

    preview_manifest, analysis_summary = _build_version_preview_context(
        db,
        doc_file,
        version,
        share_token=share_token,
    )
    if _ensure_preview_refresh_contract(version):
        db.add(version)
        db.commit()
        db.refresh(version)
    payload = VersionResponse(
        id=version.id,
        version=version.version,
        file_size=version.file_size,
        changelog=version.changelog,
        has_diff=has_diff,
        storage_mode=_safe_text(getattr(version, "storage_mode", None), "full"),
        created_at=version.created_at,
    ).model_dump()
    payload.update(
        {
            "preview_status": _safe_text(getattr(version, "preview_status", None), "pending"),
            "analysis_status": _safe_text(getattr(version, "analysis_status", None), "pending"),
            "preview_error": _safe_text(getattr(version, "preview_error", None)),
            "analysis_error": _safe_text(getattr(version, "analysis_error", None)),
            "preview_manifest": preview_manifest,
            "analysis_summary": analysis_summary,
            "preview_refresh_token": _safe_text(getattr(version, "preview_refresh_token", None)),
            "derived_asset_version": int(getattr(version, "derived_asset_version", 1) or 1),
            **build_download_contract(filename, file_type),
        }
    )
    return payload


def _persist_to_document_store(file_id, file_content, ext, file_hash):
    """unknowndocument_store unknown"""
    import hashlib
    import tempfile as _tf
    content_hash = hashlib.sha256(file_content).hexdigest()
    from app.services.document_store import _lookup_hash, ensure_registered, store_original
    existing = _lookup_hash(content_hash)
    if existing:
        files_logger.info(f"unknown {content_hash[:16]} -> {existing}")
        return
    fd, tmp = _tf.mkstemp(suffix=ext)
    try:
        with os.fdopen(fd, 'wb') as f:
            f.write(file_content)
        store_original(file_id, tmp)
        ensure_registered(file_id, tmp)
        files_logger.info(f"document_store unknown{file_id}")
    finally:
        try:
            os.unlink(tmp)
        except OSError:
            pass



@router.post("/projects/{project_id}/files", status_code=status.HTTP_201_CREATED)
async def upload_file(
    project_id: str,
    file: UploadFile = File(...),
    folder_id: Optional[str] = Form(None),
    db: Session = Depends(get_db),
    background_tasks: BackgroundTasks = BackgroundTasks(),
    current_user: User = Depends(get_current_user),
):
    project = db.query(Project).filter(Project.id == project_id).first()
    if not project:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Project not found")
    if project.owner_id != current_user.id and getattr(current_user, "role", None) != "admin":
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Project not found")
    target_folder_id = _assert_project_folder(db, project_id, folder_id)

    # Validate file type
    ext = get_file_extension(file.filename or "")
    if ext not in settings.ALLOWED_FILE_TYPES:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"File type not allowed. Allowed: {', '.join(settings.ALLOWED_FILE_TYPES)}",
        )

    # unknownContent-Lengthunknown
    content_length = file.headers.get("content-length")
    if content_length and int(content_length) > settings.MAX_FILE_SIZE:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"File too large. Max size: {settings.MAX_FILE_SIZE // (1024 * 1024)}MB",
        )

    # unknown
    _chunk_size = 8 * 1024 * 1024
    _tmp_fd, _tmp_path = tempfile.mkstemp(suffix=ext)
    _total_size = 0
    content = None
    try:
        with os.fdopen(_tmp_fd, 'wb') as _f:
            while True:
                chunk = await file.read(_chunk_size)
                if not chunk:
                    break
                _total_size += len(chunk)
                if _total_size > settings.MAX_FILE_SIZE:
                    raise HTTPException(
                        status_code=status.HTTP_400_BAD_REQUEST,
                        detail=f"File too large. Max size: {settings.MAX_FILE_SIZE // (1024 * 1024)}MB",
                    )
                _f.write(chunk)

        validate_file_type(
            file_path=__import__("pathlib").Path(_tmp_path),
            declared_filename=file.filename or "unknown",
        )

        with open(_tmp_path, 'rb') as f:
            content = f.read()
    except FileValidationError as e:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"unknown: {e.message}",
        )
    finally:
        if os.path.exists(_tmp_path):
            os.unlink(_tmp_path)

    # Create DocumentFile record
    profile = resolve_file_profile(file.filename or "", mime_type=file.content_type)
    doc_file = DocumentFile(
        project_id=project_id,
        filename=file.filename,
        file_type=ext.lstrip("."),
        file_category=profile["category"],
        mime_type=profile["mime_type"],
        current_version=1,
        preview_status=profile["preview_status"],
        preview_error=None,
        analysis_status=profile["analysis_status"],
        analysis_error=None,
        folder_id=target_folder_id,
    )
    db.add(doc_file)
    db.commit()
    db.refresh(doc_file)

    # Save file to disk and create version record
    storage_path, file_hash, file_size = save_upload_file(
        project_id=project_id,
        file_id=doc_file.id,
        version=1,
        filename=file.filename,
        content=content,
    )

    # unknown unknown document_store unknown + unknown unknown
    _persist_to_document_store(doc_file.id, content, ext, file_hash)

    # Git unknown
    obj_hash = git_store.put_object(content)
    git_store.update_ref(doc_file.id, 1, obj_hash, "full")

    version = FileVersion(
        file_id=doc_file.id,
        version=1,
        sort_order=1.0,
        storage_path=storage_path,
        file_hash=file_hash,
        file_size=file_size,
        storage_mode="full",
        preview_status=profile["preview_status"],
        preview_error=None,
        analysis_status=profile["analysis_status"],
        analysis_error=None,
    )
    db.add(version)
    db.commit()

    # unknown
    log_audit(
        user_id=current_user.id,
        action="upload_file",
        resource=f"project:{project_id}/file:{doc_file.id}",
        result="success",
        details=f"filename={file.filename}, size={file_size}",
    )

    # unknown richer preview categories
    if profile["category"] in PREVIEWABLE_CATEGORIES:
        try:
            _enqueue_preview_generation_compat(
                doc_file.id,
                storage_path,
                doc_file.file_type,
                project_id=doc_file.project_id,
                file_size=getattr(version, "file_size", None),
                updated_at=doc_file.updated_at,
            )
        except Exception:
            pass

    return success_response(data=_file_response_payload(doc_file))


@router.get("/projects/{project_id}/files")
def list_project_files(
    project_id: str,
    keyword: Optional[str] = Query(None, max_length=100),
    folder_id: Optional[str] = Query(None),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """List files under a project.

    The modern API returns an envelope with ``data.files``.  Legacy in-memory
    integration clients receive the bare list they historically consumed.
    """
    project = db.query(Project).filter(Project.id == project_id).first()
    if not project:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Project not found")
    if project.owner_id != current_user.id and getattr(current_user, "role", None) != "admin":
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Project not found")

    query = db.query(DocumentFile).filter(DocumentFile.project_id == project_id)
    if folder_id is not None:
        normalized_folder_id = _normalize_folder_id(folder_id)
        if normalized_folder_id:
            _assert_project_folder(db, project_id, normalized_folder_id)
            query = query.filter(DocumentFile.folder_id == normalized_folder_id)
        else:
            query = query.filter(DocumentFile.folder_id.is_(None))
    if keyword:
        safe = keyword.strip().replace("\\", "\\\\").replace("%", "\\%").replace("_", "\\_")
        search_pattern = f"%{safe}%"
        query = (
            query.outerjoin(Category, DocumentFile.category_id == Category.id)
            .outerjoin(DocumentFile.tags)
            .filter(
                or_(
                    DocumentFile.filename.ilike(search_pattern, escape="\\"),
                    DocumentFile.display_name.ilike(search_pattern, escape="\\"),
                    DocumentFile.description.ilike(search_pattern, escape="\\"),
                    Category.name.ilike(search_pattern, escape="\\"),
                    Tag.name.ilike(search_pattern, escape="\\"),
                )
            )
            .distinct()
        )

    docs = query.order_by(DocumentFile.created_at.desc()).all()
    if keyword:
        items = []
        for doc in docs:
            payload = _file_response_payload(doc)
            item = _file_search_item(doc, payload)
            if score_search_item(item, keyword) > 0:
                items.append(item)
        file_list = [item.payload for item in rank_search_items(items, keyword)]
    else:
        file_list = [_file_response_payload(doc) for doc in docs]

    if _is_legacy_in_memory_db(db):
        return file_list
    return success_response(data={"files": file_list})


@router.get("/files/{file_id}/preview-status")
def get_preview_status(
    file_id: str,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    doc_file = db.query(DocumentFile).filter(DocumentFile.id == file_id).first()
    if not doc_file:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="File not found")
    _require_file_action(doc_file, db, current_user, "view_metadata")
    return success_response(
        data={
            "file_id": doc_file.id,
            "preview_status": getattr(doc_file, "preview_status", None) or "pending",
            "analysis_status": getattr(doc_file, "analysis_status", None) or "pending",
            "preview_error": getattr(doc_file, "preview_error", None),
            "analysis_error": getattr(doc_file, "analysis_error", None),
        }
    )


@router.get("/files/{file_id}/versions/{version_id}/preview-status")
def get_version_preview_status(
    file_id: str,
    version_id: str,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    doc_file = db.query(DocumentFile).filter(DocumentFile.id == file_id).first()
    if not doc_file:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="File not found")
    _require_file_action(doc_file, db, current_user, "view_metadata")

    version = (
        db.query(FileVersion)
        .filter(FileVersion.id == version_id, FileVersion.file_id == file_id)
        .first()
    )
    if not version:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Version not found")

    return success_response(
        data={
            "file_id": doc_file.id,
            "version_id": version.id,
            "version": version.version,
            "preview_status": getattr(version, "preview_status", None) or "pending",
            "analysis_status": getattr(version, "analysis_status", None) or "pending",
            "preview_error": getattr(version, "preview_error", None),
            "analysis_error": getattr(version, "analysis_error", None),
        }
    )


@router.get("/files/{file_id}/analysis")
def get_file_analysis(
    file_id: str,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    doc_file = db.query(DocumentFile).filter(DocumentFile.id == file_id).first()
    if not doc_file:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="File not found")
    _require_file_action(doc_file, db, current_user, "view_metadata")

    analysis = (
        db.query(FileAnalysisRecord)
        .filter(FileAnalysisRecord.file_id == file_id)
        .order_by(FileAnalysisRecord.updated_at.desc())
        .first()
    )
    if not analysis:
        return success_response(
            data={
                "analysis_type": "",
                "payload": {},
                "status": getattr(doc_file, "analysis_status", None) or "pending",
                "error_message": getattr(doc_file, "analysis_error", None),
            }
        )

    try:
        payload = json.loads(analysis.payload_json) if analysis.payload_json else {}
    except (TypeError, ValueError):
        payload = {}

    return success_response(
        data={
            "analysis_type": analysis.analysis_type,
            "payload": payload,
            "status": analysis.status,
            "error_message": analysis.error_message,
            "updated_at": analysis.updated_at,
        }
    )


@router.get("/files/{file_id}/versions/{version_id}/analysis")
def get_version_analysis(
    file_id: str,
    version_id: str,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    doc_file = db.query(DocumentFile).filter(DocumentFile.id == file_id).first()
    if not doc_file:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="File not found")
    _require_file_action(doc_file, db, current_user, "view_metadata")

    version = (
        db.query(FileVersion)
        .filter(FileVersion.id == version_id, FileVersion.file_id == file_id)
        .first()
    )
    if not version:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Version not found")

    analysis = _load_version_analysis_record(db, file_id, version.id)
    if not analysis:
        return success_response(
            data={
                "analysis_type": "",
                "payload": {},
                "status": getattr(version, "analysis_status", None) or "pending",
                "error_message": getattr(version, "analysis_error", None),
                "updated_at": getattr(version, "created_at", None),
            }
        )

    payload = load_analysis_summary(analysis.payload_json)
    return success_response(
        data={
            "analysis_type": analysis.analysis_type,
            "payload": payload,
            "status": analysis.status,
            "error_message": analysis.error_message,
            "updated_at": analysis.updated_at,
        }
    )


@router.put("/files/{file_id}/folder")
def move_file_to_folder(
    file_id: str,
    payload: MoveFileFolderRequest,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    doc_file = db.query(DocumentFile).filter(DocumentFile.id == file_id).first()
    if not doc_file:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="File not found")
    _assert_file_access(doc_file, db, current_user)
    doc_file.folder_id = _assert_project_folder(db, doc_file.project_id, payload.folder_id)
    doc_file.updated_at = utc_now_iso()
    db.commit()
    db.refresh(doc_file)
    log_audit(
        user_id=current_user.id,
        action="move_file_folder",
        resource=f"file:{file_id}",
        result="success",
        details=f"folder_id={doc_file.folder_id or 'root'}",
    )
    return success_response(data=_file_response_payload(doc_file))


@router.post("/files/{file_id}/versions", status_code=status.HTTP_201_CREATED)
async def upload_version(
    file_id: str,
    file: UploadFile = File(...),
    changelog: Optional[str] = Form(None),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    doc_file = db.query(DocumentFile).filter(DocumentFile.id == file_id).first()
    if not doc_file:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="File not found")

    # Validate file type matches
    ext = get_file_extension(file.filename or "")
    if ext not in settings.ALLOWED_FILE_TYPES:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"File type not allowed. Allowed: {', '.join(settings.ALLOWED_FILE_TYPES)}",
        )

    # unknownContent-Length
    content_length = file.headers.get("content-length")
    if content_length and int(content_length) > settings.MAX_FILE_SIZE:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"File too large. Max size: {settings.MAX_FILE_SIZE // (1024 * 1024)}MB",
        )

    # unknown
    _chunk_size = 8 * 1024 * 1024
    _tmp_fd, _tmp_path = tempfile.mkstemp(suffix=ext)
    _total_size = 0
    content = None
    try:
        with os.fdopen(_tmp_fd, 'wb') as _f:
            while True:
                chunk = await file.read(_chunk_size)
                if not chunk:
                    break
                _total_size += len(chunk)
                if _total_size > settings.MAX_FILE_SIZE:
                    raise HTTPException(
                        status_code=status.HTTP_400_BAD_REQUEST,
                        detail=f"File too large. Max size: {settings.MAX_FILE_SIZE // (1024 * 1024)}MB",
                    )
                _f.write(chunk)

        validate_file_type(
            file_path=__import__("pathlib").Path(_tmp_path),
            declared_filename=file.filename or "unknown",
        )

        with open(_tmp_path, 'rb') as f:
            content = f.read()
    except FileValidationError as e:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"unknown: {e.message}",
        )
    finally:
        if os.path.exists(_tmp_path):
            os.unlink(_tmp_path)

    async with _locked_file_version(file_id):
        db.refresh(doc_file)
        new_version_num = doc_file.current_version + 1
        profile = resolve_file_profile(file.filename or doc_file.filename or "", mime_type=file.content_type)

        storage_path, file_hash, file_size = save_upload_file(
            project_id=doc_file.project_id,
            file_id=doc_file.id,
            version=new_version_num,
            filename=doc_file.filename,
            content=content,
        )

        # unknown unknown document_store unknown + unknown unknown
        _persist_to_document_store(doc_file.id, content, ext, file_hash)

        # unknown fullunknown deltaunknowndiff unknown
        prev_version = (
            db.query(FileVersion)
            .filter(FileVersion.file_id == file_id, FileVersion.version == new_version_num - 1)
            .first()
        )
        storage_mode = "delta" if prev_version else "full"

        # Git unknown
        obj_hash = git_store.put_object(content)
        git_store.update_ref(doc_file.id, new_version_num, obj_hash, storage_mode)

        # unknown delta_chain unknownperations unknowncompute_diff unknown
        _pending_delta_file_id = doc_file.id
        _pending_delta_version = new_version_num

        version = FileVersion(
            file_id=doc_file.id,
            version=new_version_num,
            sort_order=float(new_version_num),
            storage_path=storage_path,
            file_hash=file_hash,
            file_size=file_size,
            changelog=changelog,
            storage_mode=storage_mode,
            preview_status=profile["preview_status"],
            preview_error=None,
            analysis_status=profile["analysis_status"],
            analysis_error=None,
            base_version_id=prev_version.id if prev_version else None,
        )
        db.add(version)

        # Update current version
        doc_file.current_version = new_version_num
        doc_file.file_category = profile["category"]
        doc_file.mime_type = profile["mime_type"]
        doc_file.preview_status = profile["preview_status"]
        doc_file.preview_error = None
        doc_file.analysis_status = profile["analysis_status"]
        doc_file.analysis_error = None
        db.commit()
        db.refresh(version)

        # Auto compute diff with previous version
        diff_warning = None
        try:
            prev_version = (
                db.query(FileVersion)
                .filter(FileVersion.file_id == file_id, FileVersion.version == new_version_num - 1)
                .first()
            )
            if prev_version:
                diff_record = compute_diff(prev_version.id, version.id, db)
                # unknowndiff unknown delta_chainunknown git_store.reconstruct unknown
                if storage_mode == "delta" and diff_record.diff_data:
                    try:
                        diff_data = json.loads(diff_record.diff_data) if isinstance(diff_record.diff_data, str) else diff_record.diff_data
                        operations = []
                        for p in diff_data.get("paragraphs", []):
                            ct = p.get("change_type", "equal")
                            if ct == "replace":
                                operations.append({"type": "replace_paragraph", "index": p.get("index", 0), "text": p.get("new_text", "")})
                            elif ct == "insert":
                                operations.append({"type": "insert_paragraph", "index": p.get("index", 0), "text": p.get("new_text", "")})
                            elif ct == "delete":
                                operations.append({"type": "delete_paragraph", "index": p.get("index", 0)})
                        if operations:
                            chain = [{"base_version": new_version_num - 1, "operations": operations}]
                            chain_hash = git_store.put_delta_chain(chain)
                            git_store.update_ref(doc_file.id, new_version_num, chain_hash, "delta")
                    except Exception as chain_err:
                        files_logger.warning(f"delta_chain unknown: {chain_err}")
        except Exception as e:
            # Diff computation failure should not block upload, but log the error
            files_logger.warning(
                f"unknown - file_id: {file_id}, version: {new_version_num}, "
                f"unknown: {e}",
                exc_info=True,
            )
            diff_warning = f"unknown: {str(e)}"

        # unknown
        try:
            _enqueue_preview_generation_compat(
                doc_file.id,
                storage_path,
                doc_file.file_type,
                force=True,
                project_id=doc_file.project_id,
                file_size=getattr(version, "file_size", None),
                updated_at=doc_file.updated_at,
            )
        except Exception as preview_exc:
            files_logger.warning(
                f"unknown - file_id: {file_id}, version: {new_version_num}, unknown: {preview_exc}",
                exc_info=True,
            )

        if _ensure_preview_refresh_contract(version):
            db.add(version)
            db.commit()
            db.refresh(version)

        log_audit(
            user_id=current_user.id,
            action="upload_version",
            resource=f"file:{file_id}/version:{new_version_num}",
            result="success",
            details=f"filename={doc_file.filename}, size={file_size}, changelog={changelog}",
        )

        response_data = VersionResponse(
            id=version.id,
            version=version.version,
            file_size=version.file_size,
            changelog=version.changelog,
            has_diff=False,  # will be computed
            created_at=version.created_at,
            preview_refresh_token=version.preview_refresh_token,
            derived_asset_version=int(getattr(version, "derived_asset_version", 1) or 1),
        ).model_dump()

        if diff_warning:
            response_data["warning"] = diff_warning
        response_data["latest_version"] = new_version_num

        return success_response(data=response_data)



def _renumber_versions(db, file_id):
    """DocShop file route helper."""
    from app.models.file_version import FileVersion as FV
    versions = db.query(FV).filter(FV.file_id == file_id).order_by(FV.sort_order.asc()).all()
    for i, v in enumerate(versions, 1):
        v.version = i
    db.commit()


@router.put("/files/{file_id}/versions/reorder")
def reorder_versions(
    file_id: str,
    version_ids: list = Body(...),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """unknown/unknownversion_ids unknown"""
    from app.models.file_version import FileVersion as FV
    doc_file = db.query(DocumentFile).filter(DocumentFile.id == file_id).first()
    if not doc_file:
        raise HTTPException(status_code=404, detail="File not found")
    id_set = set(version_ids)
    all_versions = {v.id: v for v in db.query(FV).filter(FV.file_id == file_id).all()}
    if not id_set.issubset(all_versions.keys()):
        raise HTTPException(status_code=400, detail="Some version IDs are invalid")
    for i, vid in enumerate(version_ids):
        all_versions[vid].sort_order = float(i + 1)
    db.commit()
    _renumber_versions(db, file_id)
    return success_response(message="unknown")


@router.delete("/files/{file_id}/versions/{version_id}", status_code=200)
def delete_version(
    file_id: str,
    version_id: str,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_admin),
):
    """unknown V unknown"""
    from app.models.file_version import FileVersion as FV
    v = db.query(FV).filter(FV.id == version_id, FV.file_id == file_id).first()
    if not v:
        raise HTTPException(status_code=404, detail="Version not found")
    db.query(DiffRecord).filter(
        or_(
            DiffRecord.old_version_id == version_id,
            DiffRecord.new_version_id == version_id,
        )
    ).delete(synchronize_session=False)
    db.delete(v)
    db.commit()
    _renumber_versions(db, file_id)
    # unknown document_files.current_version
    doc_file = db.query(DocumentFile).filter(DocumentFile.id == file_id).first()
    latest = db.query(FV).filter(FV.file_id == file_id).order_by(FV.sort_order.desc()).first()
    if doc_file and latest:
        doc_file.current_version = latest.version
    db.commit()
    return success_response(message="unknown")


@router.get("/files/{file_id}")
def get_file_detail(
    file_id: str,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    doc = db.query(DocumentFile).filter(DocumentFile.id == file_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="File not found")
    _require_file_action(doc, db, current_user, "view_metadata")
    return success_response(data=_build_file_detail_payload(db, doc))


@router.put("/files/{file_id}/versions/{version_id}/category-tags")
@router.put("/files/{file_id}/version/{version_id}/category-tags")
def set_file_category_tags(
    file_id: str,
    version_id: str,
    category_id: Optional[str] = Body(None),
    tag_ids: list = Body(None),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_admin),
):
    """unknown"""
    from app.models.document_file import DocumentFile as DF
    from app.models.category import Category, Tag
    from app.models.file_version import FileVersion as FV
    doc = db.query(DF).filter(DF.id == file_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="File not found")
    version = db.query(FV).filter(FV.id == version_id, FV.file_id == file_id).first()
    if not version:
        raise HTTPException(status_code=404, detail="Version not found")
    if category_id is not None:
        doc.category_id = category_id if category_id else None
    if tag_ids is not None:
        tags = db.query(Tag).filter(Tag.id.in_(tag_ids)).all()
        doc.tags = tags
    db.commit()
    return success_response(message="settings saved", data={"category_id": doc.category_id, "tag_ids": [t.id for t in doc.tags]})



@router.get("/files/{file_id}/versions")
def list_versions(
    file_id: str,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    doc_file = db.query(DocumentFile).filter(DocumentFile.id == file_id).first()
    if not doc_file:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="File not found")
    _require_file_action(doc_file, db, current_user, "view_versions")

    versions = (
        db.query(FileVersion)
        .filter(FileVersion.file_id == file_id)
        .order_by(FileVersion.version.desc())
        .all()
    )

    version_list = []
    for v in versions:
        version_list.append(_build_version_payload(db, doc_file, v))

    if _is_legacy_in_memory_db(db):
        return version_list

    return success_response(
        data=VersionListResponse(
            file_id=doc_file.id,
            filename=doc_file.filename,
            file_type=doc_file.file_type,
            current_version=doc_file.current_version,
            versions=version_list,
        ).model_dump()
    )


@router.get("/files/{file_id}/versions/{version_id}/download")
def download_version(
    file_id: str,
    version_id: str,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    version = db.query(FileVersion).filter(FileVersion.id == version_id).first()
    if not version or version.file_id != file_id:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Version not found")

    doc_file = db.query(DocumentFile).filter(DocumentFile.id == file_id).first()
    if not doc_file:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="File not found")
    _require_file_action(doc_file, db, current_user, "download_original")

    if not os.path.exists(version.storage_path):
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="File not found on disk")

    # unknown UPLOAD_DIR unknown
    upload_root = os.path.realpath(settings.UPLOAD_DIR)
    real_path = os.path.realpath(version.storage_path)
    if not _is_allowed_download_path(real_path):
        files_logger.error(f"Blocked download path {version.storage_path} outside {upload_root}")
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="File not found")

    # unknown
    log_audit(
        user_id=current_user.id,
        action="download_version",
        resource=f"file:{file_id}/version:{version_id}",
        result="success",
        details=f"filename={doc_file.filename}",
    )

    download_name = _versioned_name(doc_file.filename, version.version)
    safe_name = quote(download_name)
    return FastAPIFileResponse(
        path=real_path,
        filename=download_name,
        media_type="application/octet-stream",
        headers={
            "Content-Disposition": f"attachment; filename*=UTF-8''{safe_name}",
        },
    )


@router.get("/files/{file_id}/versions/{version_id}/download/{format}")
def download_version_formatted(
    file_id: str,
    version_id: str,
    format: str,
    background_tasks: BackgroundTasks,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """
    unknown

    unknownformat unknown
    - ``docx`` / ``word`` unknown Word unknown
    - ``pdf`` unknown PDF unknown
    """
    version = db.query(FileVersion).filter(FileVersion.id == version_id).first()
    if not version or version.file_id != file_id:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Version not found")

    doc_file = db.query(DocumentFile).filter(DocumentFile.id == file_id).first()
    if not doc_file:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="File not found")
    _require_file_action(doc_file, db, current_user, "download_converted")

    if not os.path.exists(version.storage_path):
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="File not found on disk")

    real_path = os.path.realpath(version.storage_path)
    if not _is_allowed_download_path(real_path):
        files_logger.error(f"unknown {version.storage_path}")
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="File not found")

    from app.services.conversion_service import convert_to_pdf, convert_to_word, schedule_cleanup

    fmt = format.lower().strip()
    file_type = doc_file.file_type.lower()

    if fmt in ("docx", "word"):
        output_path, media_type, actual_fmt, needs_cleanup = convert_to_word(
            version.storage_path, file_type, doc_file.filename
        )
        output_real_path = os.path.realpath(output_path)
        if not _is_allowed_response_path(output_real_path):
            raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="File not found")
        download_name = _versioned_name(doc_file.filename, version.version, ".docx")
        if needs_cleanup:
            schedule_cleanup(background_tasks, output_path)
        log_audit(
            user_id=current_user.id,
            action="download_version_formatted",
            resource=f"file:{file_id}/version:{version_id}",
            result="success",
            details=f"format=docx, filename={doc_file.filename}",
        )
        return FastAPIFileResponse(
            path=output_real_path,
            filename=download_name,
            media_type=media_type,
            headers={"Content-Disposition": f"attachment; filename*=UTF-8''{quote(download_name)}"},
        )

    elif fmt == "pdf":
        output_path, media_type, actual_fmt, needs_cleanup = convert_to_pdf(
            version.storage_path, file_type, doc_file.filename
        )
        output_real_path = os.path.realpath(output_path)
        if not _is_allowed_response_path(output_real_path):
            raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="File not found")
        pdf_name = _versioned_name(doc_file.filename, version.version, ".pdf")
        disposition = "inline" if actual_fmt == "html" else "attachment"
        if needs_cleanup:
            schedule_cleanup(background_tasks, output_path)
        log_audit(
            user_id=current_user.id,
            action="download_version_formatted",
            resource=f"file:{file_id}/version:{version_id}",
            result="success",
            details=f"format=pdf, actual_fmt={actual_fmt}, filename={doc_file.filename}",
        )
        return FastAPIFileResponse(
            path=output_real_path,
            filename=pdf_name,
            media_type=media_type,
            headers={"Content-Disposition": f"{disposition}; filename*=UTF-8''{quote(pdf_name)}"},
        )

    else:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Unsupported format: {format}. Supported: docx, pdf",
        )


@router.get("/files/{file_id}/html")
def get_file_html(
    file_id: str,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """DocShop file route helper."""
    doc_file = db.query(DocumentFile).filter(DocumentFile.id == file_id).first()
    if not doc_file:
        raise HTTPException(status_code=404, detail="File not found")
    fv = db.query(FileVersion).filter(FileVersion.file_id == file_id).order_by(FileVersion.version.desc()).first()
    if not fv or not os.path.exists(fv.storage_path):
        raise HTTPException(status_code=404, detail="File not found")
    real_path = os.path.realpath(fv.storage_path)
    if not _is_allowed_storage_path(real_path):
        raise HTTPException(status_code=404, detail="File not found")

    html = "<p>unknown</p>"
    try:
        if doc_file.file_type == "docx":
            from docx import Document as DocxDoc
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from lxml import etree
            from pathlib import Path
            doc = DocxDoc(real_path)

            # MathML unknown
            MNS = "http://schemas.openxmlformats.org/officeDocument/2006/math"

            # unknown OMML unknownMathML XSLT unknown
            xsl_path = Path(__file__).parent.parent / "diff_engine" / "omml2mml.xsl"
            xslt_transform = None
            if xsl_path.exists():
                try:
                    xslt_doc = etree.parse(str(xsl_path))
                    xslt_transform = etree.XSLT(xslt_doc)
                except Exception:
                    pass

            parts = [
                '<meta charset="utf-8">',
                '<script>MathJax={tex:{inlineMath:[["$","$"]]}};</script>',
                '<script src="https://cdn.jsdelivr.net/npm/mathjax@3/es5/tex-mml-chtml.js"></script>',
                '<style>',
                '.docx-page{max-width:210mm;margin:20px auto;padding:20mm 25mm;background:#fff;box-shadow:0 0 12px rgba(0,0,0,.08);',
                'font-family:"Times New Roman","SimSun",serif;font-size:12pt;color:#111;line-height:1.5;}',
                '.docx-page h1{font-size:22pt;margin:16pt 0 8pt;font-weight:bold;border-bottom:1px solid #000;padding-bottom:4pt;}',
                '.docx-page h2{font-size:16pt;margin:14pt 0 6pt;font-weight:bold;}',
                '.docx-page h3{font-size:14pt;margin:12pt 0 4pt;font-weight:bold;}',
                '.docx-page p{margin:0 0 6pt 0;text-indent:0;}',
                '.docx-page table{border-collapse:collapse;width:100%;margin:8pt 0;}',
                '.docx-page td,.docx-page th{border:1px solid #333;padding:3pt 6pt;vertical-align:top;font-size:11pt;}',
                '@media(max-width:800px){.docx-page{padding:10px;max-width:100%;margin:0;}}',
                '</style>',
                '<div class="docx-page">'
            ]

            for para in doc.paragraphs:
                pf = para.paragraph_format
                css = []
                if pf.alignment == WD_ALIGN_PARAGRAPH.CENTER: css.append("text-align:center")
                elif pf.alignment == WD_ALIGN_PARAGRAPH.RIGHT: css.append("text-align:right")
                elif pf.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY: css.append("text-align:justify")
                if pf.first_line_indent: css.append(f"text-indent:{pf.first_line_indent.pt}pt")
                if pf.space_before: css.append(f"margin-top:{pf.space_before.pt}pt")
                if pf.space_after: css.append(f"margin-bottom:{pf.space_after.pt}pt")
                if pf.line_spacing and pf.line_spacing != 1.0: css.append(f"line-height:{pf.line_spacing}")
                style = ";".join(css) if css else ""
                is_heading = para.style.name.startswith("Heading")
                hn = para.style.name.replace("Heading ", "") if is_heading else ""

                math_elements = para._element.findall(f"{{{MNS}}}oMath") + para._element.findall(f"{{{MNS}}}oMathPara")
                if math_elements:
                    parts.append(f'<p style="{style}">' if not is_heading else f'<h{hn} style="{style}">')
                    for me in math_elements:
                        try:
                            if xslt_transform is not None:
                                # unknown XSLT unknown
                                result = xslt_transform(me)
                                mathml = str(result)
                                if mathml.strip():
                                    parts.append(mathml)
                                else:
                                    parts.append('<span style="color:#999">unknown</span>')
                            else:
                                parts.append('<span style="color:#999">unknown</span>')
                        except Exception:
                            parts.append('<span style="color:#999">unknown</span>')
                    parts.append('</p>' if not is_heading else f'</h{hn}>')
                    continue

                if not para.text.strip() and not para.runs:
                    parts.append('<p><br></p>')
                    continue

                default_style = style if style else "margin:3pt 0;line-height:1.5"
                parts.append(f'<h{hn} style="{style or "margin:12pt 0 6pt"}">' if is_heading else f'<p style="{default_style}">')
                for run in para.runs:
                    text = run.text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                    if not text: continue
                    tags, sty = [], []
                    if run.bold: tags.append("b")
                    if run.italic: tags.append("i")
                    if run.underline: tags.append("u")
                    if run.font.size: sty.append(f"font-size:{run.font.size.pt}pt")
                    if run.font.name: sty.append(f"font-family:'{run.font.name}'")
                    if run.font.color and run.font.color.rgb: sty.append(f"color:#{run.font.color.rgb}")
                    if sty: tags.append(f'span style=\"{";".join(sty)}\"')
                    for t in tags: parts.append(f"<{t}>")
                    parts.append(text)
                    for t in reversed(tags): parts.append(f"</{t.split()[0]}>")
                parts.append(f'</h{hn}>' if is_heading else '</p>')

            for table in doc.tables:
                parts.append('<table class="docx-table">')
                for row in table.rows:
                    parts.append('<tr>')
                    for cell in row.cells:
                        parts.append('<td>')
                        for cp in cell.paragraphs:
                            if cp.text.strip():
                                parts.append('<p>')
                                for r in cp.runs:
                                    t = r.text.replace("&","&amp;").replace("<","&lt;").replace(">","&gt;")
                                    if r.bold: parts.append(f"<b>{t}</b>")
                                    elif r.italic: parts.append(f"<i>{t}</i>")
                                    else: parts.append(t)
                                parts.append('</p>')
                        parts.append('</td>')
                    parts.append('</tr>')
                parts.append('</table>')

            parts.append('</div>')
            html = "".join(parts)

        elif doc_file.file_type == "pdf":
            html = f'<iframe src="/api/v1/files/{file_id}/preview" width="100%" height="700px" style="border:none"></iframe>'

    except Exception as e:
        html = f"<p>unknown: {str(e)}</p>"

    return success_response(data={"html": html, "filename": doc_file.filename, "file_type": doc_file.file_type})


@router.get("/files/{file_id}/text")
def get_file_text(
    file_id: str,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """DocShop file route helper."""
    doc_file = db.query(DocumentFile).filter(DocumentFile.id == file_id).first()
    if not doc_file:
        raise HTTPException(status_code=404, detail="File not found")
    fv = db.query(FileVersion).filter(FileVersion.file_id == file_id).order_by(FileVersion.version.desc()).first()
    if not fv or not os.path.exists(fv.storage_path):
        raise HTTPException(status_code=404, detail="File not found")
    storage_path = os.path.realpath(fv.storage_path)
    if not _is_allowed_storage_path(storage_path):
        raise HTTPException(status_code=404, detail="File not found")

    text = ""
    try:
        if doc_file.file_type == "docx":
            from docx import Document as DocxDoc
            doc = DocxDoc(storage_path)
            text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        elif doc_file.file_type in ("xlsx", "xls"):
            import openpyxl
            wb = openpyxl.load_workbook(storage_path, data_only=True)
            for name in wb.sheetnames:
                ws = wb[name]
                text += f"\n=== {name} ===\n"
                for row in ws.iter_rows(max_row=200, values_only=True):
                    text += "\t".join(str(c) if c is not None else "" for c in row) + "\n"
        elif doc_file.file_type == "pdf" and HAS_PYMUPDF:
            import fitz
            doc = fitz.open(storage_path)
            for page in doc:
                text += page.get_text()
            doc.close()
        else:
            text = "unknown"
    except Exception as e:
        text = f"unknown: {str(e)}"
    return success_response(data={"text": text[:50000], "filename": doc_file.filename, "file_type": doc_file.file_type})


@router.get("/files/{file_id}/pages/{page_num}")
def get_page_image(
    file_id: str,
    page_num: int,
    version: Optional[int] = Query(None),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """unknown JPEGunknownHTML unknown"""
    import fitz
    from app.services.conversion_service import _ensure_pdf, _source_hash as _cs_source_hash
    from app.services.document_store import render_single_page

    doc_file = db.query(DocumentFile).filter(DocumentFile.id == file_id).first()
    if not doc_file:
        raise HTTPException(status_code=404, detail="File not found")
    _require_file_action(doc_file, db, current_user, "view_page_asset")

    if doc_file.file_type not in ("docx", "doc", "pdf"):
        raise HTTPException(status_code=404, detail="Unsupported file type")

    fv = _resolve_requested_file_version(db, file_id, version)
    if not fv or not os.path.exists(fv.storage_path):
        raise HTTPException(status_code=404, detail="File not found on disk")
    real_path = os.path.realpath(fv.storage_path)
    if not _is_allowed_storage_path(real_path):
        raise HTTPException(status_code=404, detail="File not found")

    source_hash = _cs_source_hash(real_path)
    if doc_file.file_type == "pdf":
        pdf_path = real_path
        pdf_hash = source_hash
    else:
        pdf_path = _ensure_pdf(file_id, real_path, source_hash)
        if pdf_path is None:
            raise HTTPException(status_code=500, detail="PDF conversion failed")
        pdf_hash = _cs_source_hash(pdf_path)

    doc = fitz.open(pdf_path)
    page_count = len(doc)
    doc.close()

    if page_num < 1 or page_num > page_count:
        raise HTTPException(status_code=404, detail=f"Page {page_num} not found (total: {page_count})")

    try:
        img_path = render_single_page(file_id, pdf_path, page_num, page_count, pdf_hash, quality=75)
    except Exception as exc:
        logger.warning(f"unknown: file={file_id} page={page_num} err={exc}")
        raise HTTPException(status_code=500, detail="Page render failed")

    if not os.path.exists(img_path):
        raise HTTPException(status_code=404, detail=f"Page {page_num} image not found")
    real_img_path = os.path.realpath(img_path)
    if not _is_allowed_response_path(real_img_path):
        raise HTTPException(status_code=404, detail=f"Page {page_num} image not found")

    return FastAPIFileResponse(
        path=real_img_path,
        media_type="image/jpeg",
        headers={"Cache-Control": "public, max-age=86400"},
    )


@router.get("/files/{file_id}/preview-assets/{asset_id}")
def get_preview_asset(
    file_id: str,
    asset_id: str,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    doc_file = db.query(DocumentFile).filter(DocumentFile.id == file_id).first()
    if not doc_file:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="File not found")
    _require_file_action(doc_file, db, current_user, "view_page_asset")

    asset = _load_preview_asset(db, file_id, asset_id)
    if not asset:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Preview asset not found")

    return _stream_preview_asset(asset)


@router.get("/files/{file_id}/html-assets/{asset_path:path}")
def get_html_preview_asset(
    file_id: str,
    asset_path: str,
    version: Optional[int] = Query(None),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    doc_file = db.query(DocumentFile).filter(DocumentFile.id == file_id).first()
    if not doc_file or doc_file.file_type != "html":
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="File not found")
    _require_file_action(doc_file, db, current_user, "view_page_asset")

    fv = _resolve_requested_file_version(db, file_id, version)
    if not fv or not os.path.exists(fv.storage_path):
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="File not found on disk")

    return _stream_html_asset_file(fv.storage_path, asset_path)


@router.get("/files/{file_id}/preview")
def preview_file(
    file_id: str,
    version: Optional[int] = Query(None),
    auth_token: Optional[str] = Query(None),
    access_token: Optional[str] = Query(None),
    db: Session = Depends(get_db),
    background_tasks: BackgroundTasks = BackgroundTasks(),
    current_user: User = Depends(get_current_user),
):
    """DocShop file route helper."""
    doc_file = db.query(DocumentFile).filter(DocumentFile.id == file_id).first()
    if not doc_file:
        raise HTTPException(status_code=404, detail="File not found")
    _require_file_action(doc_file, db, current_user, "view_preview")

    query = db.query(FileVersion).filter(FileVersion.file_id == file_id)
    if version:
        query = query.filter(FileVersion.version == version)
    else:
        query = query.order_by(FileVersion.version.desc())
    fv = query.first()
    if not fv or not os.path.exists(fv.storage_path):
        raise HTTPException(status_code=404, detail="File not found on disk")
    preview_title = _preview_display_title(doc_file, fv.version)

    real_path = os.path.realpath(fv.storage_path)
    if not _is_allowed_storage_path(real_path):
        raise HTTPException(status_code=404, detail="File not found")

    from app.services.conversion_service import convert_to_html, convert_to_pdf

    # PDF: unknownnline unknown
    if doc_file.file_type == "pdf":
        safe_name = quote(doc_file.filename)
        return FastAPIFileResponse(
            path=real_path,
            filename=doc_file.filename,
            media_type="application/pdf",
            headers={"Content-Disposition": f"inline; filename*=UTF-8''{safe_name}", "Content-Encoding": "identity"},
        )

    if doc_file.file_type == "html":
        asset_query_params = {}
        if auth_token:
            asset_query_params["auth_token"] = auth_token
        if access_token:
            asset_query_params["access_token"] = access_token
        runtime_html = build_runtime_html_preview(
            storage_path=fv.storage_path,
            title=preview_title,
            asset_url_resolver=lambda raw_url: _build_html_asset_url(
                route_base_path=f"/api/v1/files/{file_id}/html-assets",
                raw_url=raw_url,
                version=fv.version,
                extra_query_params=asset_query_params,
            ),
        )
        return HTMLResponse(content=runtime_html, headers=runtime_html_response_headers())

    previewable_category = _previewable_category_for_file(doc_file)
    if previewable_category == "video":
        compatible_asset = _load_version_preview_asset(db, file_id, fv.id, "preview_video")
        if compatible_asset:
            return _stream_preview_asset(compatible_asset)
        return _stream_native_preview_file(doc_file, fv.storage_path)

    if previewable_category == "image":
        return _stream_native_preview_file(doc_file, fv.storage_path)

    # DOCX/DOC: try cached images first, then fallback to python-docx HTML
    if doc_file.file_type in ("docx", "doc"):
        from app.exceptions import ConversionError
        from app.services.conversion_service import build_skeleton_html, _source_hash

        # Try pre-converted images first (instant, no Word COM needed)
        source_hash = _source_hash(fv.storage_path)
        try:
            from app.services.document_store import get_cached_pdf, get_cached_images
            pdf_path = get_cached_pdf(file_id, source_hash)
            if pdf_path:
                import fitz
                doc = fitz.open(pdf_path)
                page_count = len(doc)
                doc.close()
                pdf_hash = _source_hash(pdf_path)
                cached_images = get_cached_images(file_id, pdf_hash, page_count)
                if cached_images and len(cached_images) == page_count:
                    html = build_skeleton_html(
                        file_id,
                        page_count,
                        page_count,
                        version=fv.version,
                        auth_token=auth_token or access_token,
                        title=preview_title,
                    )
                    return HTMLResponse(
                        content=html,
                        headers={"Cache-Control": "private, max-age=3600"},
                    )
        except Exception:
            pass  # cache check failed, continue to full conversion

        # No cached images: return fast python-docx HTML immediately.
        # Do not trigger Word COM pre-conversion from preview requests: large DOCX
        # files can leave WINWORD automation processes running and make the next
        # preview appear blank/slow. Use the explicit admin pre-conversion action for
        # PDF/image pre-conversion.

        # Don't block the request on synchronous Word->PDF->Images pipeline.
        # Multi-page Word docs can take 300+ seconds via COM, which is unacceptable for HTTP.
        # Instead, return fast python-docx HTML now; background preconversion will
        # build images for the next request (cache hit -> skeleton HTML with images).
        try:
            html, _media_type, _ = convert_to_html(fv.storage_path, doc_file.file_type, title=preview_title)
            return HTMLResponse(content=html)
        except Exception as e:
            raise _preview_failed_http_exception(
                e,
                log_context=f"file:{file_id}:version:{fv.version}:docx-html",
            )

    # XLSX/XLS: prefer PDF inline preview, then fall back to HTML.
    if doc_file.file_type in ("xlsx", "xls"):
        try:
            pdf_path, media_type, actual_fmt, needs_cleanup = convert_to_pdf(
                fv.storage_path, doc_file.file_type, doc_file.filename
            )
            pdf_real_path = os.path.realpath(pdf_path)
            if not _is_allowed_response_path(pdf_real_path):
                raise HTTPException(status_code=404, detail="File not found")
            if actual_fmt == "pdf":
                if needs_cleanup:
                    from app.services.conversion_service import schedule_cleanup
                    schedule_cleanup(background_tasks, pdf_path)
                safe_name = quote(os.path.basename(pdf_path))
                return FastAPIFileResponse(
                    path=pdf_real_path,
                    media_type="application/pdf",
                    headers={"Content-Disposition": f"inline; filename*=UTF-8''{safe_name}", "Content-Encoding": "identity"},
                )
            elif actual_fmt == "html":
                with open(pdf_real_path, "r", encoding="utf-8") as f:
                    html_content = f.read()
                if needs_cleanup:
                    from app.services.conversion_service import schedule_cleanup
                    schedule_cleanup(background_tasks, pdf_path)
                return HTMLResponse(content=html_content)
        except Exception as e:
            logger.warning(f"PDF unknown: {e}unknown HTML")

    # Fallback: HTML unknown
    try:
        html, media_type, _ = convert_to_html(fv.storage_path, doc_file.file_type, title=preview_title)
        return HTMLResponse(content=html)
    except Exception as e:
        raise _preview_failed_http_exception(
            e,
            log_context=f"file:{file_id}:version:{fv.version}:fallback-html",
        )


def _latest_file_version(db: Session, file_id: str) -> Optional[FileVersion]:
    return (
        db.query(FileVersion)
        .filter(FileVersion.file_id == file_id)
        .order_by(FileVersion.version.desc())
        .first()
    )


def _preview_summary(rows: list[dict]) -> dict:
    summary = {
        "total": len(rows),
        "missing": 0,
        "queued": 0,
        "pdf_generating": 0,
        "pdf_ready": 0,
        "images_generating": 0,
        "ready": 0,
        "failed": 0,
        "interrupted": 0,
        "unsupported": 0,
        "storage_bytes": 0,
    }
    storage_breakdown = {"pdf_bytes": 0, "image_bytes": 0}
    by_file_type: dict[str, int] = {}
    largest_files = []

    for row in rows:
        status_name = row.get("status") or "missing"
        summary[status_name] = summary.get(status_name, 0) + 1

        storage_bytes = int(row.get("storage_bytes") or 0)
        pdf_bytes = int(row.get("pdf_bytes") or 0)
        image_bytes = int(row.get("image_bytes") or 0)
        file_type = (row.get("file_type") or "unknown").lower()

        summary["storage_bytes"] += storage_bytes
        storage_breakdown["pdf_bytes"] += pdf_bytes
        storage_breakdown["image_bytes"] += image_bytes
        by_file_type[file_type] = by_file_type.get(file_type, 0) + 1

        if storage_bytes > 0:
            largest_files.append(
                {
                    "file_id": row.get("file_id"),
                    "filename": row.get("filename"),
                    "file_type": file_type,
                    "status": row.get("status"),
                    "storage_bytes": storage_bytes,
                    "pdf_bytes": pdf_bytes,
                    "image_bytes": image_bytes,
                }
            )

    summary["active"] = (
        summary.get("queued", 0)
        + summary.get("pdf_generating", 0)
        + summary.get("pdf_ready", 0)
        + summary.get("images_generating", 0)
    )
    summary["problem"] = summary.get("failed", 0) + summary.get("interrupted", 0)

    try:
        from app.services.preview_queue import get_queue_state

        queue_state = get_queue_state()
    except Exception:
        queue_state = {"queued": 0, "running": 0}

    summary["queue_state"] = queue_state
    summary["storage_breakdown"] = storage_breakdown
    summary["by_file_type"] = by_file_type
    summary["largest_files"] = sorted(
        largest_files, key=lambda item: item["storage_bytes"], reverse=True
    )[:5]
    return summary


@router.get("/admin/files/previews")
def admin_list_previews(
    project_id: Optional[str] = Query(None),
    preview_status: Optional[str] = Query(None, alias="status"),
    file_type: Optional[str] = Query(None),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_admin),
):
    """Return admin-visible per-file preview state and global summary."""
    from app.services.preview_queue import get_preview_snapshot

    query = db.query(DocumentFile)
    if project_id:
        query = query.filter(DocumentFile.project_id == project_id)
    if file_type:
        query = query.filter(DocumentFile.file_type == file_type.lower().lstrip("."))

    rows = []
    for doc in query.order_by(DocumentFile.created_at.desc()).all():
        row = get_preview_snapshot(
            doc.id,
            doc.file_type,
            filename=doc.filename,
            project_id=doc.project_id,
        )
        if preview_status and row["status"] != preview_status:
            continue
        rows.append(row)

    return success_response(data={"summary": _preview_summary(rows), "files": rows})


@router.post("/admin/files/preconvert")
def admin_trigger_preconvert(
    payload: dict = Body(...),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_admin),
):
    """
    unknownWordunknownDFunknown
    unknown {"file_ids": ["uuid1","uuid2"], "force": false} unknown
    unknownfile_ids unknown missing/failed/interrupted unknownDOCX/DOC/PDFunknown
    """
    from app.services.preview_queue import get_preview_snapshot, resolve_storage_path

    file_ids = payload.get("file_ids") or []
    force = bool(payload.get("force", False))
    queued = 0
    skipped = 0
    results = []

    if file_ids:
        docs = db.query(DocumentFile).filter(DocumentFile.id.in_(file_ids)).all()
    else:
        docs = db.query(DocumentFile).all()

    for df in docs:
        if _previewable_category_for_file(df) not in PREVIEWABLE_CATEGORIES:
            results.append({"file_id": df.id, "status": "skipped", "reason": "unsupported_type", "file_type": df.file_type})
            skipped += 1
            continue
        if not file_ids and not force:
            snapshot = get_preview_snapshot(df.id, df.file_type)
            if snapshot["status"] not in ("missing", "failed", "interrupted"):
                results.append({"file_id": df.id, "status": "skipped", "reason": "preview_status", "preview_status": snapshot.get("status")})
                skipped += 1
                continue
        fv = _latest_file_version(db, df.id)
        if not fv:
            results.append({"file_id": df.id, "status": "skipped", "reason": "missing_version"})
            skipped += 1
            continue
        storage_path = resolve_storage_path(fv.storage_path)
        if not os.path.exists(storage_path):
            results.append({"file_id": df.id, "status": "skipped", "reason": "missing_storage", "storage_path": storage_path})
            skipped += 1
            continue
        storage_path = os.path.realpath(storage_path)
        if not _is_allowed_storage_path(storage_path):
            results.append({"file_id": df.id, "status": "skipped", "reason": "blocked_storage_root"})
            skipped += 1
            continue
        try:
            result = _enqueue_preview_generation_compat(
                df.id,
                storage_path,
                df.file_type,
                force=force,
                project_id=df.project_id,
                file_size=getattr(fv, "file_size", None),
                updated_at=df.updated_at,
            )
            results.append(result)
            queued += 1
        except Exception as exc:
            results.append({"file_id": df.id, "status": "error", "error": str(exc)})
            skipped += 1

    return success_response(data={"queued": queued, "skipped": skipped, "total": queued + skipped, "results": results})


def _parse_iso_datetime(value):
    if not value:
        return None
    try:
        from datetime import datetime, timezone

        parsed = datetime.fromisoformat(str(value).replace("Z", "+00:00"))
        return parsed if parsed.tzinfo else parsed.replace(tzinfo=timezone.utc)
    except (TypeError, ValueError):
        return None

@router.delete("/admin/files/{file_id}/preview-cache")
def admin_clear_preview_cache(
    file_id: str,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_admin),
):
    from app.services.document_store import clear_preview_cache

    doc = db.query(DocumentFile).filter(DocumentFile.id == file_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="File not found")
    return success_response(data=clear_preview_cache(file_id))


@router.post("/admin/files/preview-cache/cleanup")
def admin_cleanup_preview_cache(
    payload: dict = Body(...),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_admin),
):
    from app.services.document_store import clear_preview_cache
    from app.services.preview_cache_policy import PreviewCacheCandidate, sort_cleanup_candidates
    from app.services.preview_queue import get_preview_snapshot

    statuses = set(payload.get("statuses") or ["failed", "interrupted"])
    max_files = int(payload.get("max_files") or 20)
    max_files = max(1, min(max_files, 200))
    max_cache_bytes = payload.get("max_cache_bytes")
    try:
        max_cache_bytes = int(max_cache_bytes) if max_cache_bytes not in (None, "") else None
    except (TypeError, ValueError):
        max_cache_bytes = None

    candidates = []
    skipped = 0
    total_storage_bytes = 0

    for doc in db.query(DocumentFile).all():
        snapshot = get_preview_snapshot(doc.id, doc.file_type, filename=doc.filename, project_id=doc.project_id)
        storage_bytes = int(snapshot.get("storage_bytes") or 0)
        total_storage_bytes += storage_bytes
        status_name = snapshot.get("status")
        status_match = status_name in statuses
        over_limit_match = max_cache_bytes is not None and storage_bytes > 0 and total_storage_bytes > max_cache_bytes
        if not status_match and not over_limit_match:
            skipped += 1
            continue
        candidates.append(
            PreviewCacheCandidate(
                file_id=doc.id,
                storage_bytes=storage_bytes,
                status=status_name or "missing",
                finished_at=_parse_iso_datetime(snapshot.get("finished_at")),
                last_accessed_at=_parse_iso_datetime(snapshot.get("updated_at")),
            )
        )

    ordered = sort_cleanup_candidates(candidates)[:max_files]
    allowed_ids = {item.file_id for item in ordered}
    cleared = 0
    removed_bytes = 0
    details = []

    for doc in db.query(DocumentFile).filter(DocumentFile.id.in_(allowed_ids)).all() if allowed_ids else []:
        result = clear_preview_cache(doc.id)
        cleared += 1
        removed_bytes += int(result.get("removed_bytes") or 0)
        details.append(result)

    return success_response(data={
        "cleared": cleared,
        "skipped": skipped,
        "removed_bytes": removed_bytes,
        "files": details,
        "policy": {
            "candidate_count": len(candidates),
            "max_files": max_files,
            "max_cache_bytes": max_cache_bytes,
            "total_storage_bytes": total_storage_bytes,
        },
    })


@router.get("/files/{file_id}/download")
def download_file(
    file_id: str,
    version: Optional[int] = Query(None),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """DocShop file route helper."""
    doc_file = db.query(DocumentFile).filter(DocumentFile.id == file_id).first()
    if not doc_file:
        raise HTTPException(status_code=404, detail="File not found")
    _require_file_action(doc_file, db, current_user, "download_original")

    query = db.query(FileVersion).filter(FileVersion.file_id == file_id)
    if version:
        query = query.filter(FileVersion.version == version)
    else:
        query = query.order_by(FileVersion.version.desc())
    fv = query.first()
    if not fv or not os.path.exists(fv.storage_path):
        raise HTTPException(status_code=404, detail="File not found on disk")

    real_path = os.path.realpath(fv.storage_path)
    if not _is_allowed_storage_path(real_path):
        raise HTTPException(status_code=404, detail="File not found")

    log_audit(
        user_id=current_user.id,
        action="download_file",
        resource=f"file:{file_id}",
        result="success",
    )

    download_name = _versioned_name(doc_file.filename, fv.version)
    safe_name = quote(download_name)
    return FastAPIFileResponse(
        path=real_path,
        filename=doc_file.filename,
        media_type="application/octet-stream",
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{safe_name}"},
    )


@router.get("/files/{file_id}/versions/{version_id}/reconstruct")
def reconstruct_version(
    file_id: str,
    version_id: str,
    background_tasks: BackgroundTasks,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """unknown"""
    version = db.query(FileVersion).filter(FileVersion.id == version_id, FileVersion.file_id == file_id).first()
    if not version:
        raise HTTPException(status_code=404, detail="Version not found")
    doc_file = db.query(DocumentFile).filter(DocumentFile.id == file_id).first()
    if not doc_file:
        raise HTTPException(status_code=404, detail="File not found")
    _require_file_action(doc_file, db, current_user, "download_converted")
    version_real_path = os.path.realpath(version.storage_path or "")
    version_exists = bool(version_real_path) and os.path.exists(version_real_path)

    # unknownfull unknown
    if version.storage_mode == "full" or version_exists:
        if not version_exists or not _is_allowed_storage_path(version_real_path):
            raise HTTPException(status_code=404, detail="File not found")
        safe_name = quote(doc_file.filename)
        return FastAPIFileResponse(
            path=version_real_path, filename=doc_file.filename,
            media_type="application/octet-stream",
            headers={"Content-Disposition": f"attachment; filename*=UTF-8''{safe_name}"},
        )

    # Delta unknown
    base = db.query(FileVersion).filter(FileVersion.id == version.base_version_id).first()
    if not base:
        raise HTTPException(status_code=404, detail="Base version not found, cannot reconstruct")
    base_real_path = os.path.realpath(base.storage_path or "")
    if not base_real_path or not os.path.exists(base_real_path) or not _is_allowed_storage_path(base_real_path):
        raise HTTPException(status_code=404, detail="Base version not found, cannot reconstruct")

    # unknown DiffRecord
    diff = db.query(DiffRecord).filter(DiffRecord.new_version_id == version_id).first()
    if not diff or not diff.diff_data:
        raise HTTPException(status_code=404, detail="Diff data not found")

    try:
        import json
        diff_data = json.loads(diff.diff_data) if isinstance(diff.diff_data, str) else diff.diff_data

        if doc_file.file_type != "docx":
            # unknownDOCXunknown
            safe_name = quote(doc_file.filename)
            return FastAPIFileResponse(
                path=base_real_path, filename=doc_file.filename,
                media_type="application/octet-stream",
                headers={"Content-Disposition": f"attachment; filename*=UTF-8''{safe_name}"},
            )

        # DOCX unknown
        from docx import Document as DocxDoc
        import shutil, tempfile
        os.makedirs(settings.TEMP_DIR, exist_ok=True)
        tmp_fd, tmp_path = tempfile.mkstemp(suffix=".docx", dir=settings.TEMP_DIR)
        os.close(tmp_fd)
        shutil.copy2(base_real_path, tmp_path)

        try:
            doc = DocxDoc(tmp_path)
            paragraphs = doc.paragraphs
            table_diffs = diff_data.get("tables", [])
            para_diffs = diff_data.get("paragraphs", [])

            # unknown
            para_map = {i: p for i, p in enumerate(paragraphs)}

            for pd in para_diffs:
                ct = pd.get("change_type", "equal")
                idx = pd.get("index", 0)
                if ct == "equal":
                    continue
                elif ct == "replace" and idx < len(paragraphs):
                    new_text = pd.get("new_text", "")
                    if new_text:
                        # unknown
                        p = paragraphs[idx]
                        for run in p.runs:
                            run.text = ""
                        if p.runs:
                            p.runs[0].text = new_text
                        else:
                            p.add_run(new_text)
                elif ct == "insert":
                    # unknown
                    new_text = pd.get("new_text", "")
                    if new_text:
                        doc.add_paragraph(new_text)

            doc.save(tmp_path)
            real_tmp_path = os.path.realpath(tmp_path)
            if not _is_allowed_response_path(real_tmp_path):
                try:
                    os.unlink(tmp_path)
                except OSError:
                    pass
                raise HTTPException(status_code=404, detail="File not found")

            safe_name = quote(doc_file.filename)
            background_tasks.add_task(os.unlink, real_tmp_path)
            return FastAPIFileResponse(
                path=real_tmp_path, filename=doc_file.filename,
                media_type="application/octet-stream",
                headers={"Content-Disposition": f"attachment; filename*=UTF-8''{safe_name}"},
            )
        except Exception:
            try:
                os.unlink(tmp_path)
            except OSError:
                pass
            raise
    except Exception as e:
        files_logger.error(f"unknown: {e}")
        raise HTTPException(status_code=500, detail=f"Reconstruction failed: {str(e)}")


@router.get("/storage/stats")
def storage_stats(current_user: User = Depends(get_current_admin)):
    """DocShop file route helper."""
    return success_response(data=git_store.get_stats())


@router.delete("/files/{file_id}", status_code=status.HTTP_204_NO_CONTENT)
def delete_file(
    file_id: str,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    doc_file = db.query(DocumentFile).filter(DocumentFile.id == file_id).first()
    if not doc_file:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="File not found")

    # unknown
    versions = db.query(FileVersion).filter(FileVersion.file_id == file_id).all()
    storage_paths = [v.storage_path for v in versions]
    version_ids = [v.id for v in versions]

    # unknown
    if version_ids:
        db.query(DiffRecord).filter(
            or_(
                DiffRecord.old_version_id.in_(version_ids),
                DiffRecord.new_version_id.in_(version_ids),
            )
        ).delete(synchronize_session=False)
    db.delete(doc_file)
    db.commit()

    # unknown
    for storage_path in storage_paths:
        real_path = os.path.realpath(storage_path or "")
        if not real_path or not _is_allowed_storage_path(real_path):
            files_logger.warning("Skip deleting storage outside allowed roots: %s", storage_path)
            continue
        if os.path.exists(real_path):
            try:
                os.remove(real_path)
            except OSError as e:
                files_logger.warning(f"unknown {real_path}: {e}")

    # Remove file directory if empty
    file_dir = os.path.join(settings.UPLOAD_DIR, doc_file.project_id, file_id)
    file_dir_real_path = os.path.realpath(file_dir)
    if _is_allowed_storage_path(file_dir_real_path) and os.path.isdir(file_dir_real_path):
        try:
            shutil.rmtree(file_dir_real_path)
        except OSError as e:
            files_logger.warning(f"unknown {file_dir_real_path}: {e}")
    elif os.path.isdir(file_dir):
        files_logger.warning("Skip deleting file directory outside allowed roots: %s", file_dir)

    # unknown
    log_audit(
        user_id=current_user.id,
        action="delete_file",
        resource=f"file:{file_id}",
        result="success",
        details=f"filename={doc_file.filename}, project_id={doc_file.project_id}",
    )

    return None

