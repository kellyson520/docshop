"""
DOCX 文档差异引擎

功能特性：
1. 段落级差异检测：识别新增、删除、修改的段落
2. 字符级精细对比：对修改的段落进行字符级差异分析
3. 表格对比：检测表格结构变化和内容变化
4. 格式检测：可选的格式变化检测
5. 性能优化：大文档分片处理、缓存机制
"""

import base64
import difflib
import hashlib
from io import BytesIO
from typing import List, Dict, Any, Optional, Tuple
from dataclasses import dataclass, field
from enum import Enum

from docx import Document
from docx.oxml.ns import qn
from diff_match_patch import diff_match_patch
from PIL import Image as PILImage

from app.diff_engine.base import BaseDiffEngine
from app.utils.logger import logger


class ChangeType(Enum):
    """变更类型"""
    EQUAL = "equal"       # 未变化
    INSERT = "insert"     # 新增
    DELETE = "delete"     # 删除
    REPLACE = "replace"   # 替换（修改）
    MOVE = "move"         # 移动/调序


@dataclass
class CharDiff:
    """字符级差异"""
    type: str  # equal, insert, delete
    text: str
    
    def to_dict(self) -> Dict[str, Any]:
        """转换为字典格式"""
        return {"type": self.type, "text": self.text}


@dataclass
class ParagraphDiff:
    """段落级差异"""
    index: int                    # 段落索引
    change_type: ChangeType       # 变更类型
    old_text: Optional[str] = None
    new_text: Optional[str] = None
    char_diffs: List[CharDiff] = field(default_factory=list)
    old_index: Optional[int] = None
    new_index: Optional[int] = None
    metadata: Dict[str, Any] = field(default_factory=dict)
    
    def to_dict(self) -> Dict[str, Any]:
        """转换为字典格式"""
        return {
            "index": self.index,
            "change_type": self.change_type.value,
            "old_text": self.old_text,
            "new_text": self.new_text,
            "char_diffs": [cd.to_dict() for cd in self.char_diffs],
            "old_index": self.old_index,
            "new_index": self.new_index,
            "metadata": self.metadata,
        }


@dataclass
class TableCellDiff:
    """表格单元格差异"""
    row: int
    col: int
    old_value: Optional[str]
    new_value: Optional[str]
    change_type: ChangeType
    
    def to_dict(self) -> Dict[str, Any]:
        """转换为字典格式"""
        return {
            "row": self.row,
            "col": self.col,
            "old_value": self.old_value,
            "new_value": self.new_value,
            "change_type": self.change_type.value
        }


@dataclass
class TableDiff:
    """表格差异"""
    table_index: int
    old_shape: Tuple[int, int]  # (rows, cols)
    new_shape: Tuple[int, int]
    cell_changes: List[TableCellDiff] = field(default_factory=list)
    structure_changed: bool = False
    added_rows: List[int] = field(default_factory=list)
    deleted_rows: List[int] = field(default_factory=list)
    added_cols: List[int] = field(default_factory=list)
    deleted_cols: List[int] = field(default_factory=list)
    row_moves: List[Dict[str, int]] = field(default_factory=list)
    col_moves: List[Dict[str, int]] = field(default_factory=list)
    old_rows: List[List[str]] = field(default_factory=list)
    new_rows: List[List[str]] = field(default_factory=list)
    
    def to_dict(self) -> Dict[str, Any]:
        """转换为字典格式"""
        return {
            "table_index": self.table_index,
            "old_shape": self.old_shape,
            "new_shape": self.new_shape,
            "structure_changed": self.structure_changed,
            "cell_changes": [cc.to_dict() for cc in self.cell_changes],
            "added_rows": self.added_rows,
            "deleted_rows": self.deleted_rows,
            "added_cols": self.added_cols,
            "deleted_cols": self.deleted_cols,
            "row_moves": self.row_moves,
            "col_moves": self.col_moves,
            "old_rows": self.old_rows,
            "new_rows": self.new_rows,
        }


class DocxDiffEngine(BaseDiffEngine):
    """DOCX 差异引擎"""
    
    def __init__(self):
        self.dmp = diff_match_patch()
        # 性能配置
        self.max_paragraphs = 10000  # 最大处理段落数
        self.chunk_size = 1000       # 分片大小
        
    def compare(self, old_path: str, new_path: str) -> Dict[str, Any]:
        """
        对比两个 DOCX 文件
        
        Args:
            old_path: 旧版本文件路径
            new_path: 新版本文件路径
            
        Returns:
            差异结果字典
            
        Raises:
            FileNotFoundError: 文件不存在
            Exception: 文档解析失败
        """
        logger.info(f"Starting DOCX diff: {old_path} vs {new_path}")
        
        try:
            # 加载文档
            old_doc = Document(old_path)
            new_doc = Document(new_path)
            
            # 提取内容
            old_paragraphs = self._extract_paragraphs(old_doc)
            new_paragraphs = self._extract_paragraphs(new_doc)
            
            logger.debug(f"Extracted paragraphs: old={len(old_paragraphs)}, new={len(new_paragraphs)}")
            
            # 检查文档大小
            if len(old_paragraphs) > self.max_paragraphs or len(new_paragraphs) > self.max_paragraphs:
                logger.warning(f"Large document detected: old={len(old_paragraphs)}, new={len(new_paragraphs)}")
                return self._compare_large(old_paragraphs, new_paragraphs)
            
            # 段落级对比
            paragraph_diffs = self._compare_paragraphs(old_paragraphs, new_paragraphs)
            
            # 表格对比
            old_tables = self._extract_tables(old_doc)
            new_tables = self._extract_tables(new_doc)
            table_diffs = self._compare_tables(old_tables, new_tables)
            
            # 图片对比
            old_images = self._extract_images(old_doc)
            new_images = self._extract_images(new_doc)
            image_diffs = self._compare_images(old_images, new_images)
            
            # 生成摘要
            summary = self._generate_summary(paragraph_diffs, table_diffs, image_diffs)
            
            text_changes = [pd.to_dict() for pd in paragraph_diffs]
            table_changes = [td.to_dict() for td in table_diffs]
            stats = {
                "paragraphs_added": sum(1 for p in paragraph_diffs if p.change_type == ChangeType.INSERT),
                "paragraphs_deleted": sum(1 for p in paragraph_diffs if p.change_type == ChangeType.DELETE),
                "paragraphs_modified": sum(1 for p in paragraph_diffs if p.change_type == ChangeType.REPLACE),
                "paragraphs_moved": self._count_move_pairs(paragraph_diffs),
                "tables_changed": len(table_diffs),
                "table_rows_moved": sum(len(t.row_moves) for t in table_diffs),
                "table_cols_moved": sum(len(t.col_moves) for t in table_diffs),
                "images_added": image_diffs.get("added", 0),
                "images_deleted": image_diffs.get("deleted", 0),
                "images_replaced": image_diffs.get("replaced", 0),
                "images_resized": image_diffs.get("resized", 0),
            }
            metadata = {
                "old_paragraph_count": len(old_paragraphs),
                "new_paragraph_count": len(new_paragraphs),
                "old_table_count": len(old_tables),
                "new_table_count": len(new_tables),
                "old_image_count": len(old_images),
                "new_image_count": len(new_images),
            }

            result = {
                "type": "docx_diff",
                "text": text_changes,
                "paragraphs": text_changes,
                "tables": table_changes,
                "images": image_diffs,
                "metadata": metadata,
                "summary": summary,
                "stats": stats,
                # 统一结构，保留 paragraphs/tables/images 旧字段给现有前端兼容。
                "changes": {
                    "text": text_changes,
                    "tables": table_changes,
                    "images": image_diffs,
                    "metadata": metadata,
                    "summary": summary,
                    "stats": stats,
                },
            }
            
            logger.info(f"DOCX diff completed: {summary}")
            return result
            
        except FileNotFoundError as e:
            logger.error(f"File not found during DOCX diff: {str(e)}")
            raise
        except Exception as e:
            logger.error(f"DOCX diff failed: {str(e)}", exc_info=True)
            raise
    
    def _extract_paragraphs(self, doc: Document) -> List[str]:
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text or para.runs:
                paragraphs.append(text)
        return paragraphs

    def _extract_images(self, doc: Document) -> List[Dict[str, Any]]:
        """Extract DOCX images with hash, paragraph slot, dimensions and thumbnail.

        Relationship-only extraction can count images, but cannot explain where an
        image was deleted or replaced.  Walking paragraph drawings gives us a stable
        `position_key` (`p{paragraph}:img{index}`) that is used by `_compare_images`.
        External image relationships are kept as first-class records, using URL as
        their content identity.
        """
        images: List[Dict[str, Any]] = []
        seen: set[str] = set()

        if not hasattr(doc, "part"):
            return images

        def append_from_rel(rel, **position):
            if not rel or "image" not in getattr(rel, "reltype", ""):
                return
            unique_key = position.get("position_key") or f"rel:{rel.rId}"
            if unique_key in seen:
                return
            seen.add(unique_key)
            images.append(self._build_image_info(rel, **position))

        for paragraph_index, paragraph in enumerate(doc.paragraphs):
            try:
                blips = paragraph._element.xpath(".//a:blip")
            except Exception:
                blips = []

            for image_index, blip in enumerate(blips):
                r_id = blip.get(qn("r:embed")) or blip.get(qn("r:link"))
                rel = doc.part.rels.get(r_id) if r_id else None
                width_emu, height_emu = self._drawing_extent(blip)
                append_from_rel(
                    rel,
                    paragraph_index=paragraph_index,
                    paragraph_text=(paragraph.text or "").strip()[:160],
                    image_index=image_index,
                    position_key=f"p{paragraph_index}:img{image_index}",
                    width_emu=width_emu,
                    height_emu=height_emu,
                )

        # Keep orphan image relationships visible for relationship-level diffing.
        for rel in doc.part.rels.values():
            if "image" not in getattr(rel, "reltype", ""):
                continue
            if any(item.get("rId") == rel.rId for item in images):
                continue
            append_from_rel(
                rel,
                paragraph_index=None,
                paragraph_text="",
                image_index=None,
                position_key=f"rel:{rel.rId}",
                width_emu=None,
                height_emu=None,
            )
        return images

    def _drawing_extent(self, blip) -> Tuple[Optional[int], Optional[int]]:
        node = blip
        while node is not None:
            tag = getattr(node, "tag", "")
            if tag.endswith("}inline") or tag.endswith("}anchor"):
                try:
                    extents = node.xpath(".//wp:extent")
                except Exception:
                    extents = []
                if extents:
                    extent = extents[0]
                    return self._safe_int(extent.get("cx")), self._safe_int(extent.get("cy"))
            node = node.getparent()
        return None, None

    def _build_image_info(self, rel, **position) -> Dict[str, Any]:
        width_emu = position.get("width_emu")
        height_emu = position.get("height_emu")
        filename = getattr(rel, "target_ref", None) or rel.rId
        base = {
            "rId": rel.rId,
            "filename": filename,
            "display_name": self._image_display_name({"filename": filename, "rId": rel.rId}),
            "paragraph_index": position.get("paragraph_index"),
            "paragraph_text": position.get("paragraph_text") or "",
            "image_index": position.get("image_index"),
            "position_key": position.get("position_key") or f"rel:{rel.rId}",
            "width_emu": width_emu,
            "height_emu": height_emu,
            "width_cm": self._emu_to_cm(width_emu),
            "height_cm": self._emu_to_cm(height_emu),
            "external": bool(getattr(rel, "is_external", False)),
        }

        if getattr(rel, "is_external", False):
            base.update({
                "content_type": "external",
                "size": 0,
                "sha256": "",
                "short_hash": "",
                "external_url": filename,
                "data_uri": "",
                "pixel_width": None,
                "pixel_height": None,
            })
            return base

        try:
            target_part = rel.target_part
            blob = getattr(target_part, "blob", b"") or b""
            sha256 = hashlib.sha256(blob).hexdigest() if blob else ""
            pixel_width, pixel_height = self._image_pixel_size(blob)
            content_type = getattr(target_part, "content_type", "") or "application/octet-stream"
            base.update({
                "content_type": content_type,
                "size": len(blob),
                "sha256": sha256,
                "short_hash": sha256[:12],
                "external_url": "",
                "data_uri": self._image_data_uri(blob, content_type),
                "pixel_width": pixel_width,
                "pixel_height": pixel_height,
            })
        except Exception as exc:
            logger.warning(f"skip unreadable image relationship {getattr(rel, 'rId', '-')}: {exc}")
            base.update({
                "content_type": "unknown",
                "size": 0,
                "sha256": "",
                "short_hash": "",
                "external_url": "",
                "data_uri": "",
                "pixel_width": None,
                "pixel_height": None,
            })
        return base

    def _image_pixel_size(self, blob: bytes) -> Tuple[Optional[int], Optional[int]]:
        if not blob:
            return None, None
        try:
            with PILImage.open(BytesIO(blob)) as image:
                return int(image.width), int(image.height)
        except Exception:
            return None, None

    def _image_data_uri(self, blob: bytes, content_type: str) -> str:
        # Keep API payloads bounded while still showing thumbnails for test images.
        if not blob or len(blob) > 768 * 1024:
            return ""
        return f"data:{content_type};base64,{base64.b64encode(blob).decode('ascii')}"

    def _safe_int(self, value) -> Optional[int]:
        try:
            return int(value) if value is not None else None
        except (TypeError, ValueError):
            return None

    def _emu_to_cm(self, value) -> Optional[float]:
        if value is None:
            return None
        return round(value / 360000, 2)

    def _extract_tables(self, doc: Document) -> List[List[List[str]]]:
        tables = []
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
            tables.append(table_data)
        return tables
    
    def _compare_paragraphs(
        self,
        old_paragraphs: List[str],
        new_paragraphs: List[str]
    ) -> List[ParagraphDiff]:
        """
        段落级对比
        
        Args:
            old_paragraphs: 旧文档段落列表
            new_paragraphs: 新文档段落列表
            
        Returns:
            段落差异列表
        """
        old_keys = [self._normalize_text(p) for p in old_paragraphs]
        new_keys = [self._normalize_text(p) for p in new_paragraphs]
        sm = difflib.SequenceMatcher(None, old_keys, new_keys, autojunk=False)
        
        diffs = []
        para_idx = 0
        
        for tag, i1, i2, j1, j2 in sm.get_opcodes():
            if tag == 'equal':
                # 未变化的段落
                for i in range(i1, i2):
                    diffs.append(ParagraphDiff(
                        index=para_idx,
                        change_type=ChangeType.EQUAL,
                        old_text=old_paragraphs[i],
                        new_text=old_paragraphs[i],
                        old_index=i,
                        new_index=j1 + (i - i1),
                    ))
                    para_idx += 1
                    
            elif tag == 'delete':
                # 删除的段落
                for i in range(i1, i2):
                    diffs.append(ParagraphDiff(
                        index=para_idx,
                        change_type=ChangeType.DELETE,
                        old_text=old_paragraphs[i],
                        new_text=None,
                        old_index=i,
                    ))
                    para_idx += 1
                    
            elif tag == 'insert':
                # 新增的段落
                for j in range(j1, j2):
                    diffs.append(ParagraphDiff(
                        index=para_idx,
                        change_type=ChangeType.INSERT,
                        old_text=None,
                        new_text=new_paragraphs[j],
                        new_index=j,
                    ))
                    para_idx += 1
                    
            elif tag == 'replace':
                block_diffs = self._compare_replace_block(
                    old_paragraphs[i1:i2],
                    new_paragraphs[j1:j2],
                    para_idx,
                    old_start=i1,
                    new_start=j1,
                )
                diffs.extend(block_diffs)
                para_idx += len(block_diffs)
        
        self._mark_moved_paragraphs(diffs)
        return diffs

    def _normalize_text(self, text: str) -> str:
        return " ".join((text or "").split()).casefold()

    def _compare_replace_block(
        self,
        old_block: List[str],
        new_block: List[str],
        start_index: int,
        old_start: int = 0,
        new_start: int = 0,
    ) -> List[ParagraphDiff]:
        old_keys = [self._normalize_text(p) for p in old_block]
        new_keys = [self._normalize_text(p) for p in new_block]
        sm = difflib.SequenceMatcher(None, old_keys, new_keys, autojunk=False)

        diffs: List[ParagraphDiff] = []
        idx = start_index

        for tag, i1, i2, j1, j2 in sm.get_opcodes():
            if tag == "equal":
                for offset in range(i2 - i1):
                    text = old_block[i1 + offset]
                    diffs.append(ParagraphDiff(
                        index=idx,
                        change_type=ChangeType.EQUAL,
                        old_text=text,
                        new_text=text,
                        old_index=old_start + i1 + offset,
                        new_index=new_start + j1 + offset,
                    ))
                    idx += 1
            elif tag == "delete":
                for offset, old_text in enumerate(old_block[i1:i2]):
                    diffs.append(ParagraphDiff(
                        index=idx,
                        change_type=ChangeType.DELETE,
                        old_text=old_text,
                        new_text=None,
                        old_index=old_start + i1 + offset,
                    ))
                    idx += 1
            elif tag == "insert":
                for offset, new_text in enumerate(new_block[j1:j2]):
                    diffs.append(ParagraphDiff(
                        index=idx,
                        change_type=ChangeType.INSERT,
                        old_text=None,
                        new_text=new_text,
                        new_index=new_start + j1 + offset,
                    ))
                    idx += 1
            else:
                old_items = old_block[i1:i2]
                new_items = new_block[j1:j2]
                pair_count = min(len(old_items), len(new_items))

                for offset in range(pair_count):
                    old_text = old_items[offset]
                    new_text = new_items[offset]
                    diffs.append(ParagraphDiff(
                        index=idx,
                        change_type=ChangeType.REPLACE,
                        old_text=old_text,
                        new_text=new_text,
                        char_diffs=self._compute_char_diffs(old_text, new_text),
                        old_index=old_start + i1 + offset,
                        new_index=new_start + j1 + offset,
                    ))
                    idx += 1

                for offset, old_text in enumerate(old_items[pair_count:]):
                    diffs.append(ParagraphDiff(
                        index=idx,
                        change_type=ChangeType.DELETE,
                        old_text=old_text,
                        new_text=None,
                        old_index=old_start + i1 + pair_count + offset,
                    ))
                    idx += 1

                for offset, new_text in enumerate(new_items[pair_count:]):
                    diffs.append(ParagraphDiff(
                        index=idx,
                        change_type=ChangeType.INSERT,
                        old_text=None,
                        new_text=new_text,
                        new_index=new_start + j1 + pair_count + offset,
                    ))
                    idx += 1

        return diffs

    def _mark_moved_paragraphs(self, diffs: List[ParagraphDiff]) -> None:
        """把同内容的 delete+insert 标记为 move，避免调序被误报为纯增删。"""
        deleted: Dict[str, List[ParagraphDiff]] = {}
        inserted: Dict[str, List[ParagraphDiff]] = {}

        for diff in diffs:
            if diff.change_type == ChangeType.DELETE:
                key = self._normalize_text(diff.old_text or "")
                if len(key) >= 4:
                    deleted.setdefault(key, []).append(diff)
            elif diff.change_type == ChangeType.INSERT:
                key = self._normalize_text(diff.new_text or "")
                if len(key) >= 4:
                    inserted.setdefault(key, []).append(diff)

        move_id = 0
        for key, old_items in deleted.items():
            new_items = inserted.get(key) or []
            pair_count = min(len(old_items), len(new_items))
            for offset in range(pair_count):
                move_id += 1
                old_diff = old_items[offset]
                new_diff = new_items[offset]
                from_index = old_diff.old_index
                to_index = new_diff.new_index
                meta = {
                    "move_id": move_id,
                    "from": from_index,
                    "to": to_index,
                    "description": f"第 {from_index + 1 if from_index is not None else '?'} 段移动到第 {to_index + 1 if to_index is not None else '?'} 段之后",
                }
                old_diff.change_type = ChangeType.MOVE
                old_diff.new_text = old_diff.old_text
                old_diff.new_index = new_diff.new_index
                old_diff.metadata.update(meta)

                new_diff.change_type = ChangeType.MOVE
                new_diff.old_text = new_diff.new_text
                new_diff.old_index = old_diff.old_index
                new_diff.metadata.update(meta)

    def _count_move_pairs(self, diffs: List[ParagraphDiff]) -> int:
        move_ids = {
            diff.metadata.get("move_id")
            for diff in diffs
            if diff.change_type == ChangeType.MOVE and diff.metadata.get("move_id") is not None
        }
        return len(move_ids)
    
    def _compute_char_diffs(self, old_text: str, new_text: str) -> List[CharDiff]:
        """
        计算字符级差异
        
        Args:
            old_text: 旧文本
            new_text: 新文本
            
        Returns:
            字符差异列表
        """
        diffs = self.dmp.diff_main(old_text, new_text)
        self.dmp.diff_cleanupSemantic(diffs)  # 语义清理
        
        char_diffs = []
        for op, text in diffs:
            if op == 0:
                char_diffs.append(CharDiff(type="equal", text=text))
            elif op == -1:
                char_diffs.append(CharDiff(type="delete", text=text))
            elif op == 1:
                char_diffs.append(CharDiff(type="insert", text=text))
        
        return char_diffs
    
    def _compare_tables(
        self,
        old_tables: List[List[List[str]]],
        new_tables: List[List[List[str]]]
    ) -> List[TableDiff]:
        """
        表格对比
        
        Args:
            old_tables: 旧文档表格列表
            new_tables: 新文档表格列表
            
        Returns:
            表格差异列表
        """
        diffs = []
        max_tables = max(len(old_tables), len(new_tables))
        
        for idx in range(max_tables):
            old_table = old_tables[idx] if idx < len(old_tables) else None
            new_table = new_tables[idx] if idx < len(new_tables) else None
            
            if old_table is None and new_table is not None:
                # 新增表格
                diffs.append(TableDiff(
                    table_index=idx,
                    old_shape=(0, 0),
                    new_shape=(len(new_table), len(new_table[0]) if new_table else 0),
                    structure_changed=True,
                    old_rows=[],
                    new_rows=new_table,
                ))
            elif old_table is not None and new_table is None:
                # 删除表格
                diffs.append(TableDiff(
                    table_index=idx,
                    old_shape=(len(old_table), len(old_table[0]) if old_table else 0),
                    new_shape=(0, 0),
                    structure_changed=True,
                    old_rows=old_table,
                    new_rows=[],
                ))
            elif old_table is not None and new_table is not None:
                # 对比表格内容
                table_diff = self._compare_single_table(idx, old_table, new_table)
                if table_diff:
                    diffs.append(table_diff)
        
        return diffs
    
    def _compare_single_table(
        self,
        idx: int,
        old_table: List[List[str]],
        new_table: List[List[str]]
    ) -> Optional[TableDiff]:
        """
        ??????
        
        Args:
            idx: ????
            old_table: ?????
            new_table: ?????
            
        Returns:
            ?????????????? None
        """
        old_rows, old_cols = len(old_table), len(old_table[0]) if old_table else 0
        new_rows, new_cols = len(new_table), len(new_table[0]) if new_table else 0
        
        structure_changed = (old_rows != new_rows or old_cols != new_cols)
        cell_changes: List[TableCellDiff] = []

        if structure_changed:
            row_new_to_old, added_rows, deleted_rows = self._sequence_alignment(
                [self._row_key(row) for row in old_table],
                [self._row_key(row) for row in new_table],
            )
            col_new_to_old, added_cols, deleted_cols = self._sequence_alignment(
                [self._row_key(col) for col in self._table_columns(old_table)],
                [self._row_key(col) for col in self._table_columns(new_table)],
            )
            # Insert/delete shifts are not semantic row/column moves. Only use
            # row_moves/col_moves for same-shape reorder detection below.
            row_moves: List[Dict[str, int]] = []
            col_moves: List[Dict[str, int]] = []
        else:
            row_moves = self._detect_row_moves(old_table, new_table)
            col_moves = self._detect_col_moves(old_table, new_table)
            added_rows = []
            deleted_rows = []
            added_cols = []
            deleted_cols = []
            row_new_to_old = {i: i for i in range(new_rows)}
            col_new_to_old = {i: i for i in range(new_cols)}
            for move in row_moves:
                row_new_to_old[move["to"]] = move["from"]
            for move in col_moves:
                col_new_to_old[move["to"]] = move["from"]

        # Compare cells through the row/column alignment. For inserted rows or
        # columns the old coordinate is None, so only the new cells are marked as
        # inserts; existing shifted rows/columns are not misreported as replace.
        for r in range(new_rows):
            old_r = row_new_to_old.get(r)
            for c in range(new_cols):
                old_c = col_new_to_old.get(c)
                old_val = (
                    old_table[old_r][old_c]
                    if old_r is not None and old_c is not None and old_r < old_rows and old_c < old_cols
                    else None
                )
                new_val = new_table[r][c] if r < new_rows and c < new_cols else None

                if old_val != new_val:
                    if old_val is None:
                        change_type = ChangeType.INSERT
                    elif new_val is None:
                        change_type = ChangeType.DELETE
                    else:
                        change_type = ChangeType.REPLACE

                    cell_changes.append(TableCellDiff(
                        row=r, col=c,
                        old_value=old_val,
                        new_value=new_val,
                        change_type=change_type
                    ))

        # Deleted rows/columns have no new coordinate, so emit explicit delete
        # cells in their old coordinate space for the front-end old-table pane.
        deleted_cell_keys = set()
        for old_r in deleted_rows:
            for old_c in range(old_cols):
                old_val = old_table[old_r][old_c] if old_r < old_rows and old_c < old_cols else None
                if old_val is None:
                    continue
                deleted_cell_keys.add((old_r, old_c))
                cell_changes.append(TableCellDiff(
                    row=old_r,
                    col=old_c,
                    old_value=old_val,
                    new_value=None,
                    change_type=ChangeType.DELETE,
                ))

        for old_c in deleted_cols:
            for old_r in range(old_rows):
                if (old_r, old_c) in deleted_cell_keys:
                    continue
                old_val = old_table[old_r][old_c] if old_r < old_rows and old_c < old_cols else None
                if old_val is None:
                    continue
                cell_changes.append(TableCellDiff(
                    row=old_r,
                    col=old_c,
                    old_value=old_val,
                    new_value=None,
                    change_type=ChangeType.DELETE,
                ))

        if cell_changes or structure_changed or row_moves or col_moves:
            return TableDiff(
                table_index=idx,
                old_shape=(old_rows, old_cols),
                new_shape=(new_rows, new_cols),
                cell_changes=cell_changes,
                structure_changed=structure_changed,
                added_rows=added_rows,
                deleted_rows=deleted_rows,
                added_cols=added_cols,
                deleted_cols=deleted_cols,
                row_moves=row_moves,
                col_moves=col_moves,
                old_rows=old_table,
                new_rows=new_table,
            )
        return None

    def _row_key(self, row: List[str]) -> str:
        return hashlib.sha256("\u241f".join(self._normalize_text(cell) for cell in row).encode("utf-8")).hexdigest()

    def _table_columns(self, table: List[List[str]]) -> List[List[str]]:
        if not table:
            return []
        width = max((len(row) for row in table), default=0)
        return [[row[col] if col < len(row) else "" for row in table] for col in range(width)]

    def _sequence_alignment(self, old_keys: List[str], new_keys: List[str]) -> Tuple[Dict[int, int], List[int], List[int]]:
        """Map new indexes to old indexes and expose true insert/delete slots."""
        new_to_old: Dict[int, int] = {}
        added: List[int] = []
        deleted: List[int] = []
        matcher = difflib.SequenceMatcher(None, old_keys, new_keys, autojunk=False)

        for tag, i1, i2, j1, j2 in matcher.get_opcodes():
            if tag == "equal":
                for offset in range(i2 - i1):
                    new_to_old[j1 + offset] = i1 + offset
            elif tag == "insert":
                added.extend(range(j1, j2))
            elif tag == "delete":
                deleted.extend(range(i1, i2))
            else:
                pair_count = min(i2 - i1, j2 - j1)
                for offset in range(pair_count):
                    new_to_old[j1 + offset] = i1 + offset
                if i2 - i1 > pair_count:
                    deleted.extend(range(i1 + pair_count, i2))
                if j2 - j1 > pair_count:
                    added.extend(range(j1 + pair_count, j2))

        return new_to_old, added, deleted

    def _detect_row_moves(self, old_table: List[List[str]], new_table: List[List[str]]) -> List[Dict[str, int]]:
        old_positions: Dict[str, List[int]] = {}
        new_positions: Dict[str, List[int]] = {}
        for idx, row in enumerate(old_table):
            old_positions.setdefault(self._row_key(row), []).append(idx)
        for idx, row in enumerate(new_table):
            new_positions.setdefault(self._row_key(row), []).append(idx)

        moves = []
        for key, old_indexes in old_positions.items():
            new_indexes = new_positions.get(key) or []
            if len(old_indexes) == 1 and len(new_indexes) == 1 and old_indexes[0] != new_indexes[0]:
                moves.append({"from": old_indexes[0], "to": new_indexes[0]})
        return moves

    def _detect_col_moves(self, old_table: List[List[str]], new_table: List[List[str]]) -> List[Dict[str, int]]:
        old_cols = self._table_columns(old_table)
        new_cols = self._table_columns(new_table)
        old_positions: Dict[str, List[int]] = {}
        new_positions: Dict[str, List[int]] = {}
        for idx, col in enumerate(old_cols):
            old_positions.setdefault(self._row_key(col), []).append(idx)
        for idx, col in enumerate(new_cols):
            new_positions.setdefault(self._row_key(col), []).append(idx)

        moves = []
        for key, old_indexes in old_positions.items():
            new_indexes = new_positions.get(key) or []
            if len(old_indexes) == 1 and len(new_indexes) == 1 and old_indexes[0] != new_indexes[0]:
                moves.append({"from": old_indexes[0], "to": new_indexes[0]})
        return moves
    
    def _image_display_name(self, image: Dict[str, Any]) -> str:
        raw = image.get("display_name") or image.get("filename") or image.get("rId") or image.get("sha256") or "unknown"
        raw = str(raw)
        return raw.rsplit("/", 1)[-1] if "/" in raw else raw

    def _image_identity(self, image: Dict[str, Any]) -> str:
        if image.get("sha256"):
            return f"sha256:{image['sha256']}"
        if image.get("external_url") or image.get("filename"):
            return f"external:{image.get('external_url') or image.get('filename')}"
        return f"rel:{image.get('rId') or image.get('position_key') or id(image)}"

    def _image_dimensions(self, image: Dict[str, Any]) -> Tuple[Any, Any, Any, Any]:
        return (
            image.get("width_emu"),
            image.get("height_emu"),
            image.get("pixel_width"),
            image.get("pixel_height"),
        )

    def _public_image(self, image: Dict[str, Any]) -> Dict[str, Any]:
        keys = [
            "rId", "filename", "display_name", "content_type", "size", "sha256", "short_hash",
            "external", "external_url", "paragraph_index", "paragraph_text", "image_index",
            "position_key", "width_emu", "height_emu", "width_cm", "height_cm",
            "pixel_width", "pixel_height", "data_uri",
        ]
        item = {key: image.get(key) for key in keys if key in image}
        item.setdefault("display_name", self._image_display_name(image))
        item.setdefault("short_hash", (image.get("sha256") or "")[:12])
        return item

    def _compare_images(self, old_images, new_images):
        """Compare image occurrences, not only unique image blobs.

        DOCX packages often reuse names such as ``word/media/image1.png`` and
        may reference the same binary image from multiple paragraphs. A pure
        dict keyed by hash/name collapses those occurrences and turns one
        deleted copy into "no change". This matcher consumes old/new records by
        list index so deletion, insertion, replacement and resize are counted at
        the paragraph slot where they actually happened.
        """
        replaced_items: List[Dict[str, Any]] = []
        resized_items: List[Dict[str, Any]] = []
        consumed_old: set[int] = set()
        consumed_new: set[int] = set()

        def consume(old_idx: int, new_idx: int) -> None:
            consumed_old.add(old_idx)
            consumed_new.add(new_idx)

        def append_resize(old_img: Dict[str, Any], new_img: Dict[str, Any], reason: str) -> None:
            resized_items.append({
                "filename": self._image_display_name(new_img),
                "old": self._public_image(old_img),
                "new": self._public_image(new_img),
                "old_width_cm": old_img.get("width_cm"),
                "old_height_cm": old_img.get("height_cm"),
                "new_width_cm": new_img.get("width_cm"),
                "new_height_cm": new_img.get("height_cm"),
                "paragraph_index": new_img.get("paragraph_index"),
                "reason": reason,
            })

        def group_remaining(images: List[Dict[str, Any]], consumed: set[int], key_fn) -> Dict[str, List[int]]:
            grouped: Dict[str, List[int]] = {}
            for idx, image in enumerate(images):
                if idx in consumed:
                    continue
                key = key_fn(image)
                if not key:
                    continue
                grouped.setdefault(str(key), []).append(idx)
            return grouped

        old_slots = self._group_image_indexes(old_images, lambda image: image.get("position_key"))
        new_slots = self._group_image_indexes(new_images, lambda image: image.get("position_key"))

        # 1) Same paragraph/image slot: identity change means replacement; same
        # identity with extent change means resize; identical records are
        # consumed as unchanged so duplicate occurrences stay accurate.
        for slot in sorted(set(old_slots) & set(new_slots)):
            for old_idx, new_idx in self._pair_image_indexes(old_slots[slot], new_slots[slot], old_images, new_images):
                if old_idx in consumed_old or new_idx in consumed_new:
                    continue
                old_img = old_images[old_idx]
                new_img = new_images[new_idx]
                old_identity = self._image_identity(old_img)
                new_identity = self._image_identity(new_img)
                if old_identity != new_identity:
                    replaced_items.append(self._image_replacement(old_img, new_img, reason="same_position"))
                    consume(old_idx, new_idx)
                elif self._image_dimensions(old_img) != self._image_dimensions(new_img):
                    append_resize(old_img, new_img, reason="same_image_resized")
                    consume(old_idx, new_idx)
                else:
                    consume(old_idx, new_idx)

        # 2) Same content identity at another position: treat as unchanged/moved
        # occurrence. If the binary is the same but the drawing extent changed,
        # report resize rather than delete+add.
        old_by_identity = group_remaining(old_images, consumed_old, self._image_identity)
        new_by_identity = group_remaining(new_images, consumed_new, self._image_identity)
        for identity in sorted(set(old_by_identity) & set(new_by_identity)):
            old_candidates = old_by_identity[identity]
            new_candidates = new_by_identity[identity]
            for old_idx, new_idx in self._pair_image_indexes(old_candidates, new_candidates, old_images, new_images):
                if old_idx in consumed_old or new_idx in consumed_new:
                    continue
                old_img = old_images[old_idx]
                new_img = new_images[new_idx]
                if self._image_dimensions(old_img) != self._image_dimensions(new_img):
                    append_resize(old_img, new_img, reason="same_image_resized")
                consume(old_idx, new_idx)

        # 3) Same display name fallback. This catches real replacements when a
        # paragraph insert shifted the slot, but deliberately refuses far apart
        # matches because DOCX auto media names are frequently reused.
        old_by_name = group_remaining(old_images, consumed_old, self._image_display_name)
        new_by_name = group_remaining(new_images, consumed_new, self._image_display_name)
        for name in sorted(set(old_by_name) & set(new_by_name)):
            old_candidates = old_by_name[name]
            new_candidates = new_by_name[name]
            for old_idx, new_idx in self._pair_image_indexes(old_candidates, new_candidates, old_images, new_images):
                if old_idx in consumed_old or new_idx in consumed_new:
                    continue
                old_img = old_images[old_idx]
                new_img = new_images[new_idx]
                if self._image_identity(old_img) == self._image_identity(new_img):
                    continue
                if not self._is_probable_image_replacement(old_img, new_img):
                    continue
                replaced_items.append(self._image_replacement(old_img, new_img, reason="same_name_nearby"))
                consume(old_idx, new_idx)

        added_items = [
            self._public_image(image)
            for idx, image in enumerate(new_images)
            if idx not in consumed_new
        ]
        deleted_items = [
            self._public_image(image)
            for idx, image in enumerate(old_images)
            if idx not in consumed_old
        ]

        changes = []
        changes.extend({"type": "added", "image": item} for item in added_items)
        changes.extend({"type": "deleted", "image": item} for item in deleted_items)
        changes.extend({"type": "replaced", **item} for item in replaced_items)
        changes.extend({"type": "resized", **item} for item in resized_items)

        return {
            "added": len(added_items),
            "deleted": len(deleted_items),
            "replaced": len(replaced_items),
            "resized": len(resized_items),
            "old_count": len(old_images),
            "new_count": len(new_images),
            "added_list": [item.get("display_name") for item in added_items],
            "deleted_list": [item.get("display_name") for item in deleted_items],
            "replaced_list": replaced_items,
            "resized_list": resized_items,
            "added_items": added_items,
            "deleted_items": deleted_items,
            "changes": changes,
        }

    def _group_image_indexes(self, images: List[Dict[str, Any]], key_fn) -> Dict[str, List[int]]:
        grouped: Dict[str, List[int]] = {}
        for idx, image in enumerate(images):
            key = key_fn(image)
            if not key:
                continue
            grouped.setdefault(str(key), []).append(idx)
        return grouped

    def _pair_image_indexes(
        self,
        old_indexes: List[int],
        new_indexes: List[int],
        old_images: List[Dict[str, Any]],
        new_images: List[Dict[str, Any]],
    ) -> List[Tuple[int, int]]:
        pairs: List[Tuple[int, int]] = []
        remaining_new = list(new_indexes)
        for old_idx in sorted(old_indexes, key=lambda idx: self._image_sort_key(old_images[idx])):
            if not remaining_new:
                break
            new_idx = min(
                remaining_new,
                key=lambda idx: (
                    self._image_paragraph_distance(old_images[old_idx], new_images[idx]),
                    self._image_sort_key(new_images[idx]),
                ),
            )
            pairs.append((old_idx, new_idx))
            remaining_new.remove(new_idx)
        return pairs

    def _image_sort_key(self, image: Dict[str, Any]) -> Tuple[int, int, str]:
        paragraph = self._optional_int(image.get("paragraph_index"))
        image_index = self._optional_int(image.get("image_index"))
        return (
            paragraph if paragraph is not None else 10**9,
            image_index if image_index is not None else 10**9,
            str(image.get("position_key") or image.get("rId") or ""),
        )

    def _image_paragraph_distance(self, old_img: Dict[str, Any], new_img: Dict[str, Any]) -> int:
        old_paragraph = self._optional_int(old_img.get("paragraph_index"))
        new_paragraph = self._optional_int(new_img.get("paragraph_index"))
        if old_paragraph is None or new_paragraph is None:
            return 10**9
        return abs(old_paragraph - new_paragraph)

    def _optional_int(self, value) -> Optional[int]:
        try:
            return int(value) if value is not None else None
        except (TypeError, ValueError):
            return None

    def _is_probable_image_replacement(self, old_img: Dict[str, Any], new_img: Dict[str, Any]) -> bool:
        if old_img.get("position_key") and old_img.get("position_key") == new_img.get("position_key"):
            return True

        old_paragraph = self._optional_int(old_img.get("paragraph_index"))
        new_paragraph = self._optional_int(new_img.get("paragraph_index"))
        if old_paragraph is not None and new_paragraph is not None:
            if abs(old_paragraph - new_paragraph) > 2:
                return False
            old_image_index = self._optional_int(old_img.get("image_index"))
            new_image_index = self._optional_int(new_img.get("image_index"))
            return (
                old_image_index is None
                or new_image_index is None
                or old_image_index == new_image_index
            )

        # Relationship-only/orphan images have no usable paragraph slot. In
        # that case a stable display name is the strongest available signal.
        return old_paragraph is None and new_paragraph is None

    def _image_replacement(self, old_img: Dict[str, Any], new_img: Dict[str, Any], reason: str) -> Dict[str, Any]:
        return {
            "filename": self._image_display_name(new_img) or self._image_display_name(old_img),
            "old_hash": old_img.get("sha256") or old_img.get("external_url") or old_img.get("filename") or "",
            "new_hash": new_img.get("sha256") or new_img.get("external_url") or new_img.get("filename") or "",
            "old_short_hash": (old_img.get("sha256") or "")[:12],
            "new_short_hash": (new_img.get("sha256") or "")[:12],
            "old_size": old_img.get("size", 0),
            "new_size": new_img.get("size", 0),
            "old": self._public_image(old_img),
            "new": self._public_image(new_img),
            "reason": reason,
            "paragraph_index": new_img.get("paragraph_index", old_img.get("paragraph_index")),
        }

    def _compare_large(
        self,
        old_paragraphs: List[str],
        new_paragraphs: List[str]
    ) -> Dict[str, Any]:
        """
        大文档分片对比
        
        Args:
            old_paragraphs: 旧文档段落列表
            new_paragraphs: 新文档段落列表
            
        Returns:
            简化的差异结果字典
        """
        logger.info("Using chunked comparison for large document")
        
        # 简化的对比策略：只比较段落数和内容哈希
        old_hash = hashlib.sha256('\n'.join(old_paragraphs).encode()).hexdigest()[:16]
        new_hash = hashlib.sha256('\n'.join(new_paragraphs).encode()).hexdigest()[:16]
        
        image_diffs = self._empty_image_diff()
        stats = {
            "paragraphs_added": max(0, len(new_paragraphs) - len(old_paragraphs)),
            "paragraphs_deleted": max(0, len(old_paragraphs) - len(new_paragraphs)),
            "paragraphs_modified": 0,
            "paragraphs_moved": 0,
            "tables_changed": 0,
            "table_rows_moved": 0,
            "table_cols_moved": 0,
            "images_added": 0,
            "images_deleted": 0,
            "images_replaced": 0,
            "images_resized": 0,
            "is_large_document": True,
            "old_hash": old_hash,
            "new_hash": new_hash
        }
        metadata = {
            "is_large_document": True,
            "old_paragraph_count": len(old_paragraphs),
            "new_paragraph_count": len(new_paragraphs),
            "old_hash": old_hash,
            "new_hash": new_hash,
        }
        summary = f"大文档对比完成，段落数: {len(old_paragraphs)} -> {len(new_paragraphs)}"

        return {
            "type": "docx_diff",
            "text": [],
            "paragraphs": [],  # 大文档不返回详细差异
            "tables": [],
            "images": image_diffs,
            "metadata": metadata,
            "summary": summary,
            "stats": stats,
            "changes": {
                "text": [],
                "tables": [],
                "images": image_diffs,
                "metadata": metadata,
                "summary": summary,
                "stats": stats,
            },
        }

    def _empty_image_diff(self) -> Dict[str, Any]:
        return {
            "added": 0,
            "deleted": 0,
            "replaced": 0,
            "resized": 0,
            "old_count": 0,
            "new_count": 0,
            "added_list": [],
            "deleted_list": [],
            "replaced_list": [],
            "resized_list": [],
            "added_items": [],
            "deleted_items": [],
            "changes": [],
        }
    
    def _generate_summary(
        self,
        paragraph_diffs: List[ParagraphDiff],
        table_diffs: List[TableDiff],
        image_diffs: dict = None,
    ) -> str:
        added = sum(1 for p in paragraph_diffs if p.change_type == ChangeType.INSERT)
        deleted = sum(1 for p in paragraph_diffs if p.change_type == ChangeType.DELETE)
        modified = sum(1 for p in paragraph_diffs if p.change_type == ChangeType.REPLACE)
        moved = self._count_move_pairs(paragraph_diffs)
        
        parts = []
        if added: parts.append(f"新增 {added} 段")
        if deleted: parts.append(f"删除 {deleted} 段")
        if modified: parts.append(f"修改 {modified} 段")
        if moved: parts.append(f"移动 {moved} 段")
        if table_diffs: parts.append(f"{len(table_diffs)} 个表格变化")
        if image_diffs:
            if image_diffs.get("added"): parts.append(f"新增 {image_diffs['added']} 张图片")
            if image_diffs.get("deleted"): parts.append(f"删除 {image_diffs['deleted']} 张图片")
            if image_diffs.get("replaced"): parts.append(f"替换 {image_diffs['replaced']} 张图片")
            if image_diffs.get("resized"): parts.append(f"尺寸调整 {image_diffs['resized']} 张图片")
        
        if parts: return "；".join(parts)
        return "文档内容无变化"
    
    def generate_summary(self, diff_data: Dict[str, Any]) -> str:
        """
        生成摘要（接口要求）
        
        Args:
            diff_data: 差异结果数据
            
        Returns:
            摘要字符串
        """
        return diff_data.get("summary", "")
