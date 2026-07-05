"""DOCX preview image rendering regressions."""

import base64
import zipfile
from pathlib import Path


PNG_1X1 = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAwMCAO+/p9sAAAAASUVORK5CYII="
)


def _make_docx_with_picture(tmp_path: Path, in_table: bool = False) -> Path:
    from docx import Document

    img = tmp_path / "pixel.png"
    img.write_bytes(PNG_1X1)

    doc = Document()
    if in_table:
        cell = doc.add_table(rows=1, cols=1).cell(0, 0)
        run = cell.paragraphs[0].add_run()
    else:
        run = doc.add_paragraph().add_run()
    run.add_picture(str(img))

    out = tmp_path / ("table_image.docx" if in_table else "external_image.docx")
    doc.save(out)
    return out


def _rewrite_first_picture_as_external_link(docx_path: Path, url: str) -> None:
    """Turn a normal embedded test image into an external r:link image."""
    with zipfile.ZipFile(docx_path, "r") as zin:
        files = {name: zin.read(name) for name in zin.namelist()}

    rels_name = "word/_rels/document.xml.rels"
    doc_name = "word/document.xml"
    rels = files[rels_name].decode("utf-8")
    document = files[doc_name].decode("utf-8")

    import re

    rid = re.search(r'Id="([^"]+)"[^>]+Type="[^"]+/image"', rels).group(1)
    rels = re.sub(
        rf'(<Relationship[^>]+Id="{rid}"[^>]+Type="[^"]+/image"[^>]+)Target="[^"]+"([^>]*/>)',
        rf'\1Target="{url}" TargetMode="External"\2',
        rels,
        count=1,
    )
    document = document.replace(f'r:embed="{rid}"', f'r:link="{rid}"')

    with zipfile.ZipFile(docx_path, "w", zipfile.ZIP_DEFLATED) as zout:
        for name, data in files.items():
            if name.startswith("word/media/"):
                continue
            if name == rels_name:
                data = rels.encode("utf-8")
            elif name == doc_name:
                data = document.encode("utf-8")
            zout.writestr(name, data)


def test_docx_preview_renders_external_linked_images(tmp_path):
    from app.services.conversion_service import convert_to_html

    docx = _make_docx_with_picture(tmp_path)
    _rewrite_first_picture_as_external_link(docx, "https://example.test/doc-image.png")

    html, _, _ = convert_to_html(str(docx), "docx")

    assert '<img src="https://example.test/doc-image.png"' in html


def test_docx_preview_renders_images_inside_table_cells(tmp_path):
    from app.services.conversion_service import convert_to_html

    docx = _make_docx_with_picture(tmp_path, in_table=True)

    html, _, _ = convert_to_html(str(docx), "docx")

    assert html.count("<img") >= 1
    assert "data:image/png;base64," in html


def test_convert_to_html_includes_visible_preview_title_and_page_shell(tmp_path):
    from docx import Document
    from app.services.conversion_service import convert_to_html

    docx = tmp_path / "title_shell.docx"
    document = Document()
    document.add_paragraph("第一页正文")
    document.save(docx)

    html, _, _ = convert_to_html(str(docx), "docx", title="汽车服务 - protable.docx · v3")

    assert '<h1 class="preview-title">汽车服务 - protable.docx · v3</h1>' in html
    assert 'class="preview-shell"' in html
    assert 'class="page-num">1 / 1</div>' in html
