"""
DOCX diff 样本生成脚本回归测试。

覆盖脚本是否固化了图片删除/替换/尺寸变化样本，以及 --verify 是否真的做断言。
"""

import shutil
import sys
import tempfile
import unittest
from pathlib import Path

from docx import Document
from docx.shared import Inches
from PIL import Image

ROOT = Path(__file__).resolve().parents[2]
sys.path.insert(0, str(ROOT))

from scripts.generate_docx_diff_tests import make_variant, verify  # noqa: E402


class TestGenerateDocxDiffTests(unittest.TestCase):
    def setUp(self):
        self.temp_dir = Path(tempfile.mkdtemp())
        self.source = self.temp_dir / "source.docx"
        self.out_dir = self.temp_dir / "out"
        self.img_dir = self.out_dir / "generated_images"
        self.out_dir.mkdir(parents=True, exist_ok=True)
        self.img_dir.mkdir(parents=True, exist_ok=True)

        image_path = self.temp_dir / "source.png"
        Image.new("RGB", (80, 50), (20, 120, 220)).save(image_path)

        doc = Document()
        for idx in range(8):
            doc.add_paragraph(f"源文档测试段落 {idx + 1}，用于验证文本新增删除修改和段落调序。")
        table = doc.add_table(rows=3, cols=3)
        for r, row in enumerate(table.rows):
            for c, cell in enumerate(row.cells):
                cell.text = f"R{r + 1}C{c + 1}"
        doc.add_paragraph("源文档图片段落")
        doc.add_picture(str(image_path), width=Inches(1.2))
        doc.save(self.source)

    def tearDown(self):
        shutil.rmtree(self.temp_dir, ignore_errors=True)

    def test_variants_include_image_delete_replace_resize_samples(self):
        manifests = [
            make_variant(self.source, self.out_dir, self.img_dir, variant)
            for variant in range(1, 4)
        ]

        operation_types = {
            op["type"]
            for manifest in manifests
            for op in manifest["operations"]
        }

        self.assertIn("image_insert", operation_types)
        self.assertIn("image_delete", operation_types)
        self.assertIn("image_replace", operation_types)
        self.assertIn("image_resize", operation_types)

    def test_verify_checks_required_diff_categories(self):
        for variant in range(1, 3):
            make_variant(self.source, self.out_dir, self.img_dir, variant)

        report = verify(self.source, self.out_dir)

        self.assertEqual(report["failed"], 0)
        self.assertEqual(report["tested"], 2)
        for item in report["results"]:
            self.assertTrue(item["ok"])
            self.assertTrue(item["checks"]["text_added"])
            self.assertTrue(item["checks"]["text_deleted"])
            self.assertTrue(item["checks"]["text_modified"])
            self.assertTrue(item["checks"]["table_changed"])
            self.assertTrue(item["checks"]["image_added"])


if __name__ == "__main__":
    unittest.main()
