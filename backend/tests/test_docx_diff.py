"""
DOCX差异引擎测试模块

测试DOCX文档差异引擎的所有功能，包括：
- 对比两个DOCX文件
- 段落级差异检测
- 表格对比
- 大文档处理
- 字符级差异分析

作者: Test Team
创建日期: 2026-05-28
"""

import os
import sys
import unittest
import tempfile
import shutil
from pathlib import Path
from unittest.mock import Mock, patch, MagicMock
from docx import Document
from docx.shared import Inches
from PIL import Image

# 添加项目根目录到路径
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from app.diff_engine.docx_diff import (
    DocxDiffEngine,
    ChangeType,
    CharDiff,
    ParagraphDiff,
    TableCellDiff,
    TableDiff,
)


class MockParagraph:
    """模拟DOCX段落对象"""
    def __init__(self, text, runs=None):
        self.text = text
        self.runs = runs or []


class MockCell:
    """模拟表格单元格"""
    def __init__(self, text):
        self.text = text


class MockRow:
    """模拟表格行"""
    def __init__(self, cells):
        self.cells = [MockCell(c) for c in cells]


class MockTable:
    """模拟表格"""
    def __init__(self, rows):
        self.rows = [MockRow(r) for r in rows]


class MockDocument:
    """模拟DOCX文档对象"""
    def __init__(self, paragraphs=None, tables=None):
        self.paragraphs = paragraphs or []
        self.tables = tables or []


class TestCharDiff(unittest.TestCase):
    """测试字符级差异类"""
    
    def test_to_dict(self):
        """测试转换为字典"""
        char_diff = CharDiff(type="insert", text="新增文本")
        result = char_diff.to_dict()
        
        self.assertEqual(result["type"], "insert")
        self.assertEqual(result["text"], "新增文本")
        
    def test_to_dict_delete(self):
        """测试删除类型"""
        char_diff = CharDiff(type="delete", text="删除文本")
        result = char_diff.to_dict()
        
        self.assertEqual(result["type"], "delete")
        self.assertEqual(result["text"], "删除文本")
        
    def test_to_dict_equal(self):
        """测试相等类型"""
        char_diff = CharDiff(type="equal", text="相同文本")
        result = char_diff.to_dict()
        
        self.assertEqual(result["type"], "equal")


class TestParagraphDiff(unittest.TestCase):
    """测试段落级差异类"""
    
    def test_to_dict_insert(self):
        """测试新增段落"""
        para_diff = ParagraphDiff(
            index=0,
            change_type=ChangeType.INSERT,
            new_text="新增段落"
        )
        result = para_diff.to_dict()
        
        self.assertEqual(result["index"], 0)
        self.assertEqual(result["change_type"], "insert")
        self.assertEqual(result["new_text"], "新增段落")
        self.assertIsNone(result["old_text"])
        
    def test_to_dict_delete(self):
        """测试删除段落"""
        para_diff = ParagraphDiff(
            index=1,
            change_type=ChangeType.DELETE,
            old_text="删除段落"
        )
        result = para_diff.to_dict()
        
        self.assertEqual(result["change_type"], "delete")
        self.assertEqual(result["old_text"], "删除段落")
        self.assertIsNone(result["new_text"])
        
    def test_to_dict_replace(self):
        """测试替换段落"""
        char_diffs = [CharDiff(type="delete", text="旧")]
        para_diff = ParagraphDiff(
            index=2,
            change_type=ChangeType.REPLACE,
            old_text="旧文本",
            new_text="新文本",
            char_diffs=char_diffs
        )
        result = para_diff.to_dict()
        
        self.assertEqual(result["change_type"], "replace")
        self.assertEqual(result["old_text"], "旧文本")
        self.assertEqual(result["new_text"], "新文本")
        self.assertEqual(len(result["char_diffs"]), 1)
        
    def test_to_dict_equal(self):
        """测试未变化段落"""
        para_diff = ParagraphDiff(
            index=3,
            change_type=ChangeType.EQUAL,
            old_text="相同文本",
            new_text="相同文本"
        )
        result = para_diff.to_dict()
        
        self.assertEqual(result["change_type"], "equal")


class TestTableCellDiff(unittest.TestCase):
    """测试表格单元格差异类"""
    
    def test_to_dict_insert(self):
        """测试新增单元格"""
        cell_diff = TableCellDiff(
            row=0,
            col=0,
            old_value=None,
            new_value="新值",
            change_type=ChangeType.INSERT
        )
        result = cell_diff.to_dict()
        
        self.assertEqual(result["row"], 0)
        self.assertEqual(result["col"], 0)
        self.assertIsNone(result["old_value"])
        self.assertEqual(result["new_value"], "新值")
        self.assertEqual(result["change_type"], "insert")
        
    def test_to_dict_replace(self):
        """测试修改单元格"""
        cell_diff = TableCellDiff(
            row=1,
            col=2,
            old_value="旧值",
            new_value="新值",
            change_type=ChangeType.REPLACE
        )
        result = cell_diff.to_dict()
        
        self.assertEqual(result["change_type"], "replace")


class TestTableDiff(unittest.TestCase):
    """测试表格差异类"""
    
    def test_to_dict_structure_changed(self):
        """测试表格结构变化"""
        table_diff = TableDiff(
            table_index=0,
            old_shape=(2, 3),
            new_shape=(3, 4),
            structure_changed=True
        )
        result = table_diff.to_dict()
        
        self.assertEqual(result["table_index"], 0)
        self.assertEqual(result["old_shape"], (2, 3))
        self.assertEqual(result["new_shape"], (3, 4))
        self.assertTrue(result["structure_changed"])
        
    def test_to_dict_with_cell_changes(self):
        """测试包含单元格变化"""
        cell_changes = [
            TableCellDiff(0, 0, "A", "B", ChangeType.REPLACE)
        ]
        table_diff = TableDiff(
            table_index=1,
            old_shape=(2, 2),
            new_shape=(2, 2),
            cell_changes=cell_changes,
            structure_changed=False
        )
        result = table_diff.to_dict()
        
        self.assertEqual(len(result["cell_changes"]), 1)
        self.assertFalse(result["structure_changed"])


class TestDocxDiffEngineInit(unittest.TestCase):
    """测试DOCX差异引擎初始化"""
    
    def test_init(self):
        """测试初始化"""
        engine = DocxDiffEngine()
        
        self.assertEqual(engine.max_paragraphs, 10000)
        self.assertEqual(engine.chunk_size, 1000)
        self.assertIsNotNone(engine.dmp)


class TestDocxDiffEngineExtractParagraphs(unittest.TestCase):
    """测试提取段落功能"""
    
    def test_extract_paragraphs_normal(self):
        """测试正常提取段落"""
        engine = DocxDiffEngine()
        
        paragraphs = [
            MockParagraph("第一段"),
            MockParagraph("第二段"),
            MockParagraph("  第三段  "),  # 测试空格清理
        ]
        doc = MockDocument(paragraphs=paragraphs)
        
        result = engine._extract_paragraphs(doc)
        
        self.assertEqual(len(result), 3)
        self.assertEqual(result[0], "第一段")
        self.assertEqual(result[2], "第三段")  # 空格被清理
        
    def test_extract_paragraphs_with_empty(self):
        """测试包含空段落"""
        engine = DocxDiffEngine()
        
        paragraphs = [
            MockParagraph("第一段"),
            MockParagraph(""),  # 空段落
            MockParagraph("第三段"),
        ]
        doc = MockDocument(paragraphs=paragraphs)
        
        result = engine._extract_paragraphs(doc)
        
        # 空段落应该被过滤
        self.assertEqual(len(result), 2)
        
    def test_extract_paragraphs_with_runs(self):
        """测试包含runs的空文本段落"""
        engine = DocxDiffEngine()
        
        paragraphs = [
            MockParagraph("", runs=["run1"]),  # 空文本但有runs
            MockParagraph("第二段"),
        ]
        doc = MockDocument(paragraphs=paragraphs)
        
        result = engine._extract_paragraphs(doc)
        
        # 有runs的空段落应该被保留
        self.assertEqual(len(result), 2)
        self.assertEqual(result[0], "")


class TestDocxDiffEngineExtractTables(unittest.TestCase):
    """测试提取表格功能"""
    
    def test_extract_tables_normal(self):
        """测试正常提取表格"""
        engine = DocxDiffEngine()
        
        tables = [
            MockTable([["A1", "B1"], ["A2", "B2"]]),
            MockTable([["C1", "D1"]]),
        ]
        doc = MockDocument(tables=tables)
        
        result = engine._extract_tables(doc)
        
        self.assertEqual(len(result), 2)
        self.assertEqual(result[0], [["A1", "B1"], ["A2", "B2"]])
        self.assertEqual(result[1], [["C1", "D1"]])
        
    def test_extract_tables_empty(self):
        """测试无表格"""
        engine = DocxDiffEngine()
        doc = MockDocument(tables=[])
        
        result = engine._extract_tables(doc)
        
        self.assertEqual(len(result), 0)


class TestDocxDiffEngineCompareParagraphs(unittest.TestCase):
    """测试段落对比功能"""
    
    def test_compare_equal_paragraphs(self):
        """测试相同段落"""
        engine = DocxDiffEngine()
        
        old = ["第一段", "第二段"]
        new = ["第一段", "第二段"]
        
        result = engine._compare_paragraphs(old, new)
        
        self.assertEqual(len(result), 2)
        self.assertEqual(result[0].change_type, ChangeType.EQUAL)
        self.assertEqual(result[1].change_type, ChangeType.EQUAL)
        
    def test_compare_insert_paragraph(self):
        """测试新增段落"""
        engine = DocxDiffEngine()
        
        old = ["第一段"]
        new = ["第一段", "新增段落"]
        
        result = engine._compare_paragraphs(old, new)
        
        self.assertEqual(len(result), 2)
        self.assertEqual(result[0].change_type, ChangeType.EQUAL)
        self.assertEqual(result[1].change_type, ChangeType.INSERT)
        self.assertEqual(result[1].new_text, "新增段落")
        
    def test_compare_delete_paragraph(self):
        """测试删除段落"""
        engine = DocxDiffEngine()
        
        old = ["第一段", "删除段落"]
        new = ["第一段"]
        
        result = engine._compare_paragraphs(old, new)
        
        self.assertEqual(len(result), 2)
        self.assertEqual(result[0].change_type, ChangeType.EQUAL)
        self.assertEqual(result[1].change_type, ChangeType.DELETE)
        self.assertEqual(result[1].old_text, "删除段落")
        
    def test_compare_replace_paragraph(self):
        """测试替换段落"""
        engine = DocxDiffEngine()
        
        old = ["旧段落"]
        new = ["新段落"]
        
        result = engine._compare_paragraphs(old, new)
        
        self.assertEqual(len(result), 1)
        self.assertEqual(result[0].change_type, ChangeType.REPLACE)
        self.assertEqual(result[0].old_text, "旧段落")
        self.assertEqual(result[0].new_text, "新段落")
        self.assertTrue(len(result[0].char_diffs) > 0)

    def test_compare_moved_paragraph(self):
        """测试段落调序识别为 move，而不是纯新增/删除"""
        engine = DocxDiffEngine()

        old = ["第一段内容", "需要移动的完整段落内容", "第三段内容"]
        new = ["第一段内容", "第三段内容", "需要移动的完整段落内容"]

        result = engine._compare_paragraphs(old, new)
        moved = [item for item in result if item.change_type == ChangeType.MOVE]

        self.assertGreaterEqual(len(moved), 1)
        self.assertEqual(engine._count_move_pairs(result), 1)


class TestDocxDiffEngineComputeCharDiffs(unittest.TestCase):
    """测试字符级差异计算"""
    
    def test_compute_char_diffs_equal(self):
        """测试相同文本"""
        engine = DocxDiffEngine()
        
        result = engine._compute_char_diffs("相同文本", "相同文本")
        
        # 应该只有一个equal类型的差异
        self.assertTrue(all(d.type == "equal" for d in result))
        
    def test_compute_char_diffs_insert(self):
        """测试插入文本"""
        engine = DocxDiffEngine()
        
        result = engine._compute_char_diffs("原文", "原文新增")
        
        # 应该有equal和insert两种类型
        types = [d.type for d in result]
        self.assertIn("equal", types)
        self.assertIn("insert", types)
        
    def test_compute_char_diffs_delete(self):
        """测试删除文本"""
        engine = DocxDiffEngine()
        
        result = engine._compute_char_diffs("原文删除", "原文")
        
        # 应该有equal和delete两种类型
        types = [d.type for d in result]
        self.assertIn("equal", types)
        self.assertIn("delete", types)
        
    def test_compute_char_diffs_replace(self):
        """测试替换文本"""
        engine = DocxDiffEngine()
        
        result = engine._compute_char_diffs("旧文本", "新文本")
        
        # 应该有delete和insert类型
        types = [d.type for d in result]
        self.assertIn("delete", types)
        self.assertIn("insert", types)


class TestDocxDiffEngineCompareTables(unittest.TestCase):
    """测试表格对比功能"""
    
    def test_compare_tables_no_change(self):
        """测试无变化的表格"""
        engine = DocxDiffEngine()
        
        old = [[["A1", "B1"], ["A2", "B2"]]]
        new = [[["A1", "B1"], ["A2", "B2"]]]
        
        result = engine._compare_tables(old, new)
        
        self.assertEqual(len(result), 0)  # 无变化返回空列表
        
    def test_compare_tables_added(self):
        """测试新增表格"""
        engine = DocxDiffEngine()
        
        old = []
        new = [[["A1", "B1"]]]
        
        result = engine._compare_tables(old, new)
        
        self.assertEqual(len(result), 1)
        self.assertTrue(result[0].structure_changed)
        
    def test_compare_tables_deleted(self):
        """测试删除表格"""
        engine = DocxDiffEngine()
        
        old = [[["A1", "B1"]]]
        new = []
        
        result = engine._compare_tables(old, new)
        
        self.assertEqual(len(result), 1)
        self.assertTrue(result[0].structure_changed)
        
    def test_compare_tables_cell_change(self):
        """测试单元格内容变化"""
        engine = DocxDiffEngine()
        
        old = [[["A1", "B1"], ["A2", "B2"]]]
        new = [[["A1", "B1"], ["A2", "修改"]]]
        
        result = engine._compare_tables(old, new)
        
        self.assertEqual(len(result), 1)
        self.assertEqual(len(result[0].cell_changes), 1)
        self.assertEqual(result[0].cell_changes[0].old_value, "B2")
        self.assertEqual(result[0].cell_changes[0].new_value, "修改")

    def test_compare_tables_row_move(self):
        """测试表格行调序识别"""
        engine = DocxDiffEngine()

        old = [[["A1", "B1"], ["A2", "B2"], ["A3", "B3"]]]
        new = [[["A1", "B1"], ["A3", "B3"], ["A2", "B2"]]]

        result = engine._compare_tables(old, new)

        self.assertEqual(len(result), 1)
        self.assertTrue(result[0].row_moves)
        self.assertEqual(len(result[0].cell_changes), 0)

    def test_compare_tables_col_move_without_false_cell_changes(self):
        """纯列调序应标记为 col_moves，不应误报整列单元格替换"""
        engine = DocxDiffEngine()

        old = [[["A1", "B1", "C1"], ["A2", "B2", "C2"]]]
        new = [[["B1", "A1", "C1"], ["B2", "A2", "C2"]]]

        result = engine._compare_tables(old, new)

        self.assertEqual(len(result), 1)
        self.assertTrue(result[0].col_moves)
        self.assertEqual(len(result[0].cell_changes), 0)

    def test_compare_tables_middle_row_insert_without_shift_replaces(self):
        """中间插入行应标记新增行和新增单元格，不应把后续行误报为修改/移动"""
        engine = DocxDiffEngine()

        old = [[["A1"], ["B1"]]]
        new = [[["A1"], ["X1"], ["B1"]]]

        result = engine._compare_tables(old, new)

        self.assertEqual(len(result), 1)
        table = result[0]
        self.assertEqual(table.added_rows, [1])
        self.assertEqual(table.deleted_rows, [])
        self.assertEqual(table.row_moves, [])
        self.assertEqual(len(table.cell_changes), 1)
        self.assertEqual(table.cell_changes[0].change_type, ChangeType.INSERT)
        self.assertEqual(table.cell_changes[0].row, 1)
        self.assertEqual(table.cell_changes[0].new_value, "X1")

    def test_compare_tables_middle_col_insert_without_shift_replaces(self):
        """中间插入列应标记新增列和新增单元格，不应把后续列误报为修改/移动"""
        engine = DocxDiffEngine()

        old = [[["A1", "C1"], ["A2", "C2"]]]
        new = [[["A1", "B1", "C1"], ["A2", "B2", "C2"]]]

        result = engine._compare_tables(old, new)

        self.assertEqual(len(result), 1)
        table = result[0]
        self.assertEqual(table.added_cols, [1])
        self.assertEqual(table.deleted_cols, [])
        self.assertEqual(table.col_moves, [])
        inserted_values = [cell.new_value for cell in table.cell_changes]
        self.assertEqual([cell.change_type for cell in table.cell_changes], [ChangeType.INSERT, ChangeType.INSERT])
        self.assertEqual(inserted_values, ["B1", "B2"])


class TestDocxDiffEngineCompareSingleTable(unittest.TestCase):
    """测试单个表格对比"""
    
    def test_compare_single_table_no_change(self):
        """测试无变化"""
        engine = DocxDiffEngine()
        
        old = [["A1", "B1"], ["A2", "B2"]]
        new = [["A1", "B1"], ["A2", "B2"]]
        
        result = engine._compare_single_table(0, old, new)
        
        self.assertIsNone(result)
        
    def test_compare_single_table_structure_change(self):
        """测试结构变化"""
        engine = DocxDiffEngine()
        
        old = [["A1", "B1"]]
        new = [["A1", "B1"], ["A2", "B2"]]  # 增加一行
        
        result = engine._compare_single_table(0, old, new)
        
        self.assertIsNotNone(result)
        self.assertTrue(result.structure_changed)
        
    def test_compare_single_table_cell_insert(self):
        """测试新增单元格"""
        engine = DocxDiffEngine()
        
        old = [["A1", None]]
        new = [["A1", "B1"]]
        
        result = engine._compare_single_table(0, old, new)
        
        self.assertIsNotNone(result)
        self.assertEqual(result.cell_changes[0].change_type, ChangeType.INSERT)


class TestDocxDiffEngineCompareImages(unittest.TestCase):
    """测试 DOCX 图片新增、删除、替换识别"""

    def setUp(self):
        self.temp_dir = tempfile.mkdtemp()

    def tearDown(self):
        shutil.rmtree(self.temp_dir, ignore_errors=True)

    def _image(self, name, color):
        path = Path(self.temp_dir) / name
        Image.new("RGB", (32, 32), color).save(path)
        return path

    def _doc(self, name, image_path=None):
        path = Path(self.temp_dir) / name
        doc = Document()
        doc.add_paragraph("图片测试文档")
        if image_path:
            doc.add_picture(str(image_path))
        doc.save(path)
        return path

    def test_compare_images_added_deleted_replaced(self):
        engine = DocxDiffEngine()
        red = self._image("red.png", (255, 0, 0))
        blue = self._image("blue.png", (0, 0, 255))

        empty_doc = self._doc("empty.docx")
        red_doc = self._doc("red.docx", red)
        blue_doc = self._doc("blue.docx", blue)

        added = engine.compare(str(empty_doc), str(red_doc))
        deleted = engine.compare(str(red_doc), str(empty_doc))
        replaced = engine.compare(str(red_doc), str(blue_doc))

        self.assertEqual(added["stats"]["images_added"], 1)
        self.assertEqual(deleted["stats"]["images_deleted"], 1)
        self.assertEqual(replaced["stats"]["images_replaced"], 1)

    def test_extract_images_includes_position_hash_dimensions_and_thumbnail(self):
        """提取图片时应保留 hash、段落位置、尺寸和前端缩略图 data URI"""
        engine = DocxDiffEngine()
        red = self._image("red-meta.png", (255, 0, 0))
        doc_path = self._doc("red-meta.docx", red)

        images = engine._extract_images(Document(str(doc_path)))

        self.assertEqual(len(images), 1)
        image = images[0]
        self.assertEqual(image["paragraph_index"], 1)
        self.assertEqual(image["image_index"], 0)
        self.assertEqual(image["position_key"], "p1:img0")
        self.assertEqual(len(image["sha256"]), 64)
        self.assertTrue(image["short_hash"])
        self.assertGreater(image["width_cm"], 0)
        self.assertGreater(image["height_cm"], 0)
        self.assertEqual(image["pixel_width"], 32)
        self.assertEqual(image["pixel_height"], 32)
        self.assertTrue(image["data_uri"].startswith("data:image/"))

    def test_compare_images_resize_is_reported_with_summary_and_items(self):
        """同一位置同一图片仅尺寸变化时，应识别为 resized 而不是新增/删除"""
        engine = DocxDiffEngine()
        red = self._image("red-resize.png", (255, 0, 0))

        old_path = Path(self.temp_dir) / "old-resize.docx"
        old_doc = Document()
        old_doc.add_paragraph("图片尺寸测试")
        old_doc.add_picture(str(red), width=Inches(1.0))
        old_doc.save(old_path)

        new_path = Path(self.temp_dir) / "new-resize.docx"
        new_doc = Document()
        new_doc.add_paragraph("图片尺寸测试")
        new_doc.add_picture(str(red), width=Inches(2.0))
        new_doc.save(new_path)

        result = engine.compare(str(old_path), str(new_path))

        self.assertEqual(result["stats"]["images_added"], 0)
        self.assertEqual(result["stats"]["images_deleted"], 0)
        self.assertEqual(result["stats"]["images_replaced"], 0)
        self.assertEqual(result["stats"]["images_resized"], 1)
        self.assertEqual(len(result["images"]["resized_list"]), 1)
        resized = result["images"]["resized_list"][0]
        self.assertNotEqual(resized["old_width_cm"], resized["new_width_cm"])
        self.assertIn("尺寸", result["summary"])

    def test_compare_images_same_media_name_far_apart_is_add_and_delete(self):
        """media/image1.png 这类自动命名在远位置复用时，不应误判为图片替换"""
        engine = DocxDiffEngine()
        old_images = [{
            "rId": "rId1",
            "filename": "word/media/image1.png",
            "display_name": "image1.png",
            "sha256": "a" * 64,
            "short_hash": "a" * 12,
            "size": 100,
            "position_key": "p1:img0",
            "paragraph_index": 1,
            "image_index": 0,
        }]
        new_images = [{
            "rId": "rId1",
            "filename": "word/media/image1.png",
            "display_name": "image1.png",
            "sha256": "b" * 64,
            "short_hash": "b" * 12,
            "size": 120,
            "position_key": "p40:img0",
            "paragraph_index": 40,
            "image_index": 0,
        }]

        result = engine._compare_images(old_images, new_images)

        self.assertEqual(result["replaced"], 0)
        self.assertEqual(result["deleted"], 1)
        self.assertEqual(result["added"], 1)
        self.assertEqual(result["deleted_items"][0]["paragraph_index"], 1)
        self.assertEqual(result["added_items"][0]["paragraph_index"], 40)

    def test_compare_images_counts_duplicate_occurrences_by_position(self):
        """同一张图片出现多次时，应按出现位置计数，删除一个位置就报告一次删除"""
        engine = DocxDiffEngine()
        image_hash = "c" * 64
        old_images = [
            {
                "rId": "rId1",
                "filename": "word/media/image1.png",
                "display_name": "image1.png",
                "sha256": image_hash,
                "short_hash": image_hash[:12],
                "size": 100,
                "position_key": "p1:img0",
                "paragraph_index": 1,
                "image_index": 0,
            },
            {
                "rId": "rId1",
                "filename": "word/media/image1.png",
                "display_name": "image1.png",
                "sha256": image_hash,
                "short_hash": image_hash[:12],
                "size": 100,
                "position_key": "p5:img0",
                "paragraph_index": 5,
                "image_index": 0,
            },
        ]
        new_images = [old_images[0].copy()]

        result = engine._compare_images(old_images, new_images)

        self.assertEqual(result["added"], 0)
        self.assertEqual(result["deleted"], 1)
        self.assertEqual(result["replaced"], 0)
        self.assertEqual(result["deleted_items"][0]["position_key"], "p5:img0")


class TestDocxDiffEngineCompareLarge(unittest.TestCase):
    """测试大文档对比"""
    
    def test_compare_large(self):
        """测试大文档简化对比"""
        engine = DocxDiffEngine()
        
        old = ["段落" + str(i) for i in range(100)]
        new = ["段落" + str(i) for i in range(150)]
        
        result = engine._compare_large(old, new)
        
        self.assertEqual(result["type"], "docx_diff")
        self.assertTrue(result["stats"]["is_large_document"])
        self.assertIn("text", result)
        self.assertIn("metadata", result)
        self.assertIn("summary", result["changes"])
        self.assertIn("stats", result["changes"])
        self.assertEqual(len(result["paragraphs"]), 0)  # 大文档不返回详细差异
        self.assertEqual(result["stats"]["paragraphs_added"], 50)


class TestDocxDiffEngineGenerateSummary(unittest.TestCase):
    """测试生成摘要功能"""
    
    def test_generate_summary_no_change(self):
        """测试无变化"""
        engine = DocxDiffEngine()
        
        para_diffs = [
            ParagraphDiff(0, ChangeType.EQUAL, "文本", "文本")
        ]
        table_diffs = []
        
        result = engine._generate_summary(para_diffs, table_diffs)
        
        self.assertEqual(result, "文档内容无变化")
        
    def test_generate_summary_all_changes(self):
        """测试所有类型的变化"""
        engine = DocxDiffEngine()
        
        para_diffs = [
            ParagraphDiff(0, ChangeType.INSERT, None, "新增"),
            ParagraphDiff(1, ChangeType.DELETE, "删除", None),
            ParagraphDiff(2, ChangeType.REPLACE, "旧", "新"),
        ]
        table_diffs = [TableDiff(0, (2, 2), (2, 2))]
        
        result = engine._generate_summary(para_diffs, table_diffs)
        
        self.assertIn("新增", result)
        self.assertIn("删除", result)
        self.assertIn("修改", result)
        self.assertIn("表格", result)
        
    def test_generate_summary_interface(self):
        """测试接口要求的generate_summary方法"""
        engine = DocxDiffEngine()
        
        diff_data = {"summary": "测试摘要"}
        result = engine.generate_summary(diff_data)
        
        self.assertEqual(result, "测试摘要")
        
    def test_generate_summary_interface_no_summary(self):
        """测试无摘要时的接口方法"""
        engine = DocxDiffEngine()
        
        diff_data = {}
        result = engine.generate_summary(diff_data)
        
        self.assertEqual(result, "")


class TestDocxDiffEngineCompare(unittest.TestCase):
    """测试主对比功能"""
    
    def setUp(self):
        """创建临时目录"""
        self.temp_dir = tempfile.mkdtemp()
        
    def tearDown(self):
        """清理"""
        shutil.rmtree(self.temp_dir, ignore_errors=True)
        
    @patch('app.diff_engine.docx_diff.Document')
    def test_compare_success(self, mock_document_class):
        """测试正常对比"""
        engine = DocxDiffEngine()
        
        # 模拟文档
        old_doc = MockDocument(
            paragraphs=[MockParagraph("第一段"), MockParagraph("第二段")],
            tables=[]
        )
        new_doc = MockDocument(
            paragraphs=[MockParagraph("第一段"), MockParagraph("修改段")],
            tables=[]
        )
        
        mock_document_class.side_effect = [old_doc, new_doc]
        
        old_path = os.path.join(self.temp_dir, "old.docx")
        new_path = os.path.join(self.temp_dir, "new.docx")
        
        # 创建空文件（实际不会被读取，因为Document被mock了）
        Path(old_path).touch()
        Path(new_path).touch()
        
        result = engine.compare(old_path, new_path)
        
        self.assertEqual(result["type"], "docx_diff")
        self.assertIn("text", result)
        self.assertIn("paragraphs", result)
        self.assertIn("tables", result)
        self.assertIn("images", result)
        self.assertIn("metadata", result)
        self.assertIn("summary", result)
        self.assertIn("stats", result)
        self.assertIn("summary", result["changes"])
        self.assertIn("stats", result["changes"])
        self.assertIn("metadata", result["changes"])
        
    @patch('app.diff_engine.docx_diff.Document')
    def test_compare_file_not_found(self, mock_document_class):
        """测试文件不存在"""
        engine = DocxDiffEngine()
        
        mock_document_class.side_effect = FileNotFoundError("文件不存在")
        
        with self.assertRaises(FileNotFoundError):
            engine.compare("/nonexistent/old.docx", "/nonexistent/new.docx")
            
    @patch('app.diff_engine.docx_diff.Document')
    def test_compare_parse_error(self, mock_document_class):
        """测试文档解析错误"""
        engine = DocxDiffEngine()
        
        mock_document_class.side_effect = Exception("解析失败")
        
        with self.assertRaises(Exception):
            engine.compare("old.docx", "new.docx")
            
    @patch('app.diff_engine.docx_diff.Document')
    def test_compare_large_document(self, mock_document_class):
        """测试大文档对比"""
        engine = DocxDiffEngine()
        engine.max_paragraphs = 10  # 设置较小的阈值
        
        # 创建超过阈值的段落
        old_doc = MockDocument(
            paragraphs=[MockParagraph(f"段落{i}") for i in range(20)],
            tables=[]
        )
        new_doc = MockDocument(
            paragraphs=[MockParagraph(f"段落{i}") for i in range(25)],
            tables=[]
        )
        
        mock_document_class.side_effect = [old_doc, new_doc]
        
        old_path = os.path.join(self.temp_dir, "old.docx")
        new_path = os.path.join(self.temp_dir, "new.docx")
        Path(old_path).touch()
        Path(new_path).touch()
        
        result = engine.compare(old_path, new_path)
        
        self.assertTrue(result["stats"]["is_large_document"])


if __name__ == "__main__":
    unittest.main()
